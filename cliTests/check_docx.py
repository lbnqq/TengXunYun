#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX文件检查工具

Author: AI Assistant (Claude)
Created: 2025-01-28
Last Modified: 2025-01-28
Modified By: AI Assistant (Claude)
AI Assisted: 是 - Claude 3.5 Sonnet
Version: v1.0
License: MIT
"""



import sys
import os
from docx import Document

def check_docx_content(file_path):
    if not os.path.exists(file_path):
        print(f"❌ 文件不存在: {file_path}")
        return
    try:
        doc = Document(file_path)
        print(f"段落数: {len(doc.paragraphs)}")
        for i, para in enumerate(doc.paragraphs):
            print(f"段落{i+1}: {para.text}")
        print(f"表格数: {len(doc.tables)}")
        for i, table in enumerate(doc.tables):
            print(f"表格{i+1} 行数: {len(table.rows)} 列数: {len(table.columns)}")
    except Exception as e:
        print(f"❌ 解析docx失败: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python check_docx.py <docx文件路径>")
    else:
        check_docx_content(sys.argv[1])