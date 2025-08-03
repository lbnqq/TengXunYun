#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文风对齐协调器

负责协调文风统一功能的各个组件，管理处理会话和任务状态。
类似于format_alignment_coordinator的设计模式。

Author: AI Assistant
Created: 2025-01-16
License: MIT
"""

import uuid
import json
import os
from typing import Dict, List, Any, Optional
from datetime import datetime
import logging
import threading
import time
from io import BytesIO

# 导入核心模块
from .style_transfer import StyleTransferEngine

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class StyleAlignmentCoordinator:
    """文风对齐协调器"""
    
    def __init__(self, spark_x1_client=None):
        """
        初始化文风对齐协调器
        
        Args:
            spark_x1_client: 星火X1客户端实例
        """
        self.spark_x1_client = spark_x1_client
        self.style_engine = StyleTransferEngine(spark_x1_client)
        
        # 任务管理
        self.active_tasks = {}  # 活跃任务
        self.task_history = {}  # 任务历史
        self.sessions = {}      # 会话管理
        
        # 进度管理
        self.task_progress = {}  # 任务进度
        self.progress_lock = threading.Lock()
        
        logger.info("✅ StyleAlignmentCoordinator初始化成功")
    
    def create_session(self, user_id: str = "default_user") -> str:
        """
        创建新的处理会话
        
        Args:
            user_id: 用户ID
            
        Returns:
            会话ID
        """
        session_id = f"style_session_{uuid.uuid4().hex[:8]}_{int(time.time())}"
        
        self.sessions[session_id] = {
            'session_id': session_id,
            'user_id': user_id,
            'created_at': datetime.now().isoformat(),
            'tasks': [],
            'status': 'active'
        }
        
        logger.info(f"📝 创建文风处理会话: {session_id}")
        return session_id
    
    def process_preset_style_generation(self, session_id: str, content: str, 
                                      style_id: str, action: str = "重写",
                                      language: str = 'auto', 
                                      temperature: Optional[float] = None) -> Dict[str, Any]:
        """
        处理预设风格生成请求
        
        Args:
            session_id: 会话ID
            content: 原始内容
            style_id: 风格ID
            action: 操作类型
            language: 语言代码
            temperature: 温度参数
            
        Returns:
            处理结果字典
        """
        try:
            # 创建任务
            task_id = f"preset_style_{uuid.uuid4().hex[:8]}"
            
            # 初始化任务状态
            self._init_task_progress(task_id, "预设风格生成")
            
            # 更新进度：开始处理
            self._update_progress(task_id, 10, "开始处理预设风格生成...")
            
            # 验证会话
            if session_id not in self.sessions:
                raise ValueError(f"无效的会话ID: {session_id}")
            
            # 更新进度：验证完成
            self._update_progress(task_id, 20, "会话验证完成")
            
            # 调用风格引擎
            self._update_progress(task_id, 30, "调用AI风格引擎...")
            
            result = self.style_engine.generate_with_style(
                content=content,
                style_id=style_id,
                action=action,
                language=language,
                temperature=temperature
            )
            
            if result.get('success'):
                # 更新进度：生成完成
                self._update_progress(task_id, 80, "风格生成完成")
                
                # 生成对比分析
                self._update_progress(task_id, 90, "生成对比分析...")
                comparison = self.style_engine.get_style_comparison(
                    content, result['generated_content'], style_id
                )
                
                # 保存任务结果
                task_result = {
                    'task_id': task_id,
                    'session_id': session_id,
                    'type': 'preset_style_generation',
                    'status': 'completed',
                    'original_content': content,
                    'generated_content': result['generated_content'],
                    'style_id': style_id,
                    'style_name': result['style_name'],
                    'language': result['language'],
                    'temperature': result['temperature'],
                    'comparison': comparison,
                    'created_at': datetime.now().isoformat(),
                    'completed_at': datetime.now().isoformat()
                }
                
                self.active_tasks[task_id] = task_result
                self.sessions[session_id]['tasks'].append(task_id)
                
                # 完成进度
                self._update_progress(task_id, 100, "处理完成")
                
                logger.info(f"✅ 预设风格生成完成: {task_id}")
                
                return {
                    'success': True,
                    'task_id': task_id,
                    'generated_content': result['generated_content'],
                    'style_name': result['style_name'],
                    'comparison': comparison,
                    'language': result['language']
                }
            else:
                # 处理失败
                self._update_progress(task_id, -1, f"处理失败: {result.get('error')}")
                
                return {
                    'success': False,
                    'task_id': task_id,
                    'error': result.get('error')
                }
                
        except Exception as e:
            logger.error(f"预设风格生成失败: {str(e)}")
            if 'task_id' in locals():
                self._update_progress(task_id, -1, f"异常错误: {str(e)}")
            
            return {
                'success': False,
                'error': str(e)
            }
    
    def process_few_shot_transfer(self, session_id: str, content: str,
                                reference_document: str, target_description: str = "",
                                language: str = 'auto', temperature: float = 0.7) -> Dict[str, Any]:
        """
        处理Few-Shot风格迁移请求
        
        Args:
            session_id: 会话ID
            content: 要转换的内容
            reference_document: 参考文档内容
            target_description: 目标风格描述
            language: 语言代码
            temperature: 温度参数
            
        Returns:
            处理结果字典
        """
        try:
            # 创建任务
            task_id = f"few_shot_{uuid.uuid4().hex[:8]}"
            
            # 初始化任务状态
            self._init_task_progress(task_id, "Few-Shot风格迁移")
            
            # 更新进度：开始处理
            self._update_progress(task_id, 10, "开始Few-Shot风格迁移...")
            
            # 验证会话
            if session_id not in self.sessions:
                raise ValueError(f"无效的会话ID: {session_id}")
            
            # 更新进度：提取风格特征
            self._update_progress(task_id, 20, "分析参考文档风格...")
            
            # 从参考文档提取风格特征
            style_analysis = self.style_engine.extract_style_from_document(
                reference_document, language
            )
            
            if not style_analysis.get('success'):
                raise Exception(f"风格分析失败: {style_analysis.get('error')}")
            
            # 更新进度：准备示例
            self._update_progress(task_id, 40, "准备风格示例...")
            
            # 将参考文档分割为示例（简单分割）
            reference_examples = self._split_document_to_examples(reference_document)
            
            # 更新进度：执行风格迁移
            self._update_progress(task_id, 60, "执行风格迁移...")
            
            # 调用Few-Shot风格迁移
            result = self.style_engine.few_shot_style_transfer(
                content=content,
                reference_examples=reference_examples,
                target_description=target_description,
                language=language,
                temperature=temperature
            )
            
            if result.get('success'):
                # 更新进度：生成完成
                self._update_progress(task_id, 90, "风格迁移完成")
                
                # 保存任务结果
                task_result = {
                    'task_id': task_id,
                    'session_id': session_id,
                    'type': 'few_shot_transfer',
                    'status': 'completed',
                    'original_content': content,
                    'generated_content': result['generated_content'],
                    'reference_document': reference_document,
                    'style_analysis': style_analysis['style_analysis'],
                    'reference_count': result['reference_count'],
                    'language': result['language'],
                    'temperature': result['temperature'],
                    'created_at': datetime.now().isoformat(),
                    'completed_at': datetime.now().isoformat()
                }
                
                self.active_tasks[task_id] = task_result
                self.sessions[session_id]['tasks'].append(task_id)
                
                # 完成进度
                self._update_progress(task_id, 100, "处理完成")
                
                logger.info(f"✅ Few-Shot风格迁移完成: {task_id}")
                
                return {
                    'success': True,
                    'task_id': task_id,
                    'generated_content': result['generated_content'],
                    'style_analysis': style_analysis['style_analysis'],
                    'reference_count': result['reference_count'],
                    'language': result['language']
                }
            else:
                # 处理失败
                self._update_progress(task_id, -1, f"处理失败: {result.get('error')}")
                
                return {
                    'success': False,
                    'task_id': task_id,
                    'error': result.get('error')
                }
                
        except Exception as e:
            logger.error(f"Few-Shot风格迁移失败: {str(e)}")
            if 'task_id' in locals():
                self._update_progress(task_id, -1, f"异常错误: {str(e)}")
            
            return {
                'success': False,
                'error': str(e)
            }

    def get_task_result(self, task_id: str) -> Dict[str, Any]:
        """
        获取任务结果

        Args:
            task_id: 任务ID

        Returns:
            任务结果字典
        """
        try:
            if task_id in self.active_tasks:
                task_result = self.active_tasks[task_id].copy()

                # 添加进度信息
                if task_id in self.task_progress:
                    task_result['progress'] = self.task_progress[task_id]

                return {
                    'success': True,
                    'data': task_result
                }
            else:
                return {
                    'success': False,
                    'error': f'任务不存在: {task_id}'
                }

        except Exception as e:
            logger.error(f"获取任务结果失败: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }

    def get_task_progress(self, task_id: str) -> Dict[str, Any]:
        """
        获取任务进度

        Args:
            task_id: 任务ID

        Returns:
            进度信息字典
        """
        with self.progress_lock:
            if task_id in self.task_progress:
                return self.task_progress[task_id].copy()
            else:
                return {
                    'task_id': task_id,
                    'progress': 0,
                    'status': 'not_found',
                    'message': '任务不存在'
                }

    def get_preset_styles(self, language: str = 'auto') -> Dict[str, Any]:
        """
        获取预设风格模板库

        Args:
            language: 语言代码

        Returns:
            预设风格字典
        """
        try:
            styles = self.style_engine.get_preset_styles(language)
            return {
                'success': True,
                'styles': styles,
                'count': len(styles)
            }
        except Exception as e:
            logger.error(f"获取预设风格失败: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }

    def export_result(self, task_id: str, format_type: str = 'txt') -> Dict[str, Any]:
        """
        导出任务结果

        Args:
            task_id: 任务ID
            format_type: 导出格式 ('txt', 'docx', 'pdf')

        Returns:
            Dict[str, Any]: 导出结果
        """
        try:
            # 获取任务结果
            task_result = self.get_task_result(task_id)

            if not task_result.get('success'):
                logger.error(f"❌ 任务结果获取失败: {task_result.get('error')}")
                return {
                    'success': False,
                    'error': '任务结果不存在或获取失败'
                }

            result_data = task_result.get('data', {})
            content = result_data.get('generated_content', '') or result_data.get('generated', '')

            if not content:
                logger.error(f"❌ 没有可导出的内容，任务ID: {task_id}")
                return {
                    'success': False,
                    'error': '没有可导出的内容'
                }

            # 生成文件名
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"style_result_{timestamp}.{format_type}"

            # 根据格式类型处理导出
            if format_type.lower() == 'txt':
                return self._export_txt(content, filename)
            elif format_type.lower() == 'docx':
                return self._export_docx(content, filename)
            elif format_type.lower() == 'pdf':
                return self._export_pdf(content, filename)
            else:
                return {
                    'success': False,
                    'error': f'不支持的导出格式: {format_type}'
                }

        except Exception as e:
            logger.error(f"❌ 导出失败: {str(e)}")
            return {
                'success': False,
                'error': f'导出失败: {str(e)}'
            }

    def _init_task_progress(self, task_id: str, task_type: str):
        """初始化任务进度"""
        with self.progress_lock:
            self.task_progress[task_id] = {
                'task_id': task_id,
                'task_type': task_type,
                'progress': 0,
                'status': 'started',
                'message': '任务已创建',
                'started_at': datetime.now().isoformat(),
                'updated_at': datetime.now().isoformat()
            }

    def _update_progress(self, task_id: str, progress: int, message: str):
        """更新任务进度"""
        with self.progress_lock:
            if task_id in self.task_progress:
                self.task_progress[task_id].update({
                    'progress': progress,
                    'message': message,
                    'updated_at': datetime.now().isoformat()
                })

                # 设置状态
                if progress == 100:
                    self.task_progress[task_id]['status'] = 'completed'
                elif progress == -1:
                    self.task_progress[task_id]['status'] = 'failed'
                else:
                    self.task_progress[task_id]['status'] = 'processing'

                logger.info(f"📊 任务进度更新 {task_id}: {progress}% - {message}")

    def _export_txt(self, content: str, filename: str) -> Dict[str, Any]:
        """导出TXT格式"""
        try:
            # 确保上传目录存在
            upload_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), 'uploads')
            os.makedirs(upload_dir, exist_ok=True)

            file_path = os.path.join(upload_dir, filename)

            # 写入文件
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)

            return {
                'success': True,
                'filename': filename,
                'file_path': file_path,
                'download_url': f'/uploads/{filename}',
                'format': 'txt'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'TXT导出失败: {str(e)}'
            }

    def _export_docx(self, content: str, filename: str) -> Dict[str, Any]:
        """导出DOCX格式"""
        try:
            # 尝试导入python-docx
            try:
                from docx import Document
                from docx.shared import Inches
            except ImportError:
                return {
                    'success': False,
                    'error': 'python-docx库未安装，无法导出DOCX格式'
                }

            # 创建文档
            doc = Document()

            # 添加标题
            doc.add_heading('文风统一处理结果', 0)

            # 添加内容
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            # 确保上传目录存在
            upload_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), 'uploads')
            os.makedirs(upload_dir, exist_ok=True)

            file_path = os.path.join(upload_dir, filename)

            # 保存文档
            doc.save(file_path)

            return {
                'success': True,
                'filename': filename,
                'file_path': file_path,
                'download_url': f'/uploads/{filename}',
                'format': 'docx'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'DOCX导出失败: {str(e)}'
            }

    def _export_pdf(self, content: str, filename: str) -> Dict[str, Any]:
        """导出PDF格式"""
        try:
            # 尝试导入reportlab
            try:
                from reportlab.lib.pagesizes import letter, A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
            except ImportError:
                return {
                    'success': False,
                    'error': 'reportlab库未安装，无法导出PDF格式'
                }

            # 确保上传目录存在
            upload_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), 'uploads')
            os.makedirs(upload_dir, exist_ok=True)

            file_path = os.path.join(upload_dir, filename)

            # 创建PDF文档
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            story = []

            # 获取样式
            styles = getSampleStyleSheet()
            title_style = styles['Title']
            normal_style = styles['Normal']

            # 添加标题
            title = Paragraph("文风统一处理结果", title_style)
            story.append(title)
            story.append(Spacer(1, 12))

            # 添加内容
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = Paragraph(paragraph.strip(), normal_style)
                    story.append(p)
                    story.append(Spacer(1, 12))

            # 构建PDF
            doc.build(story)

            return {
                'success': True,
                'filename': filename,
                'file_path': file_path,
                'download_url': f'/uploads/{filename}',
                'format': 'pdf'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'PDF导出失败: {str(e)}'
            }

    def _split_document_to_examples(self, document: str, max_examples: int = 3) -> List[str]:
        """
        将文档分割为示例

        Args:
            document: 文档内容
            max_examples: 最大示例数量

        Returns:
            示例列表
        """
        # 简单的分割逻辑：按段落分割
        paragraphs = [p.strip() for p in document.split('\n\n') if p.strip()]

        # 如果段落太少，按句子分割
        if len(paragraphs) < max_examples:
            sentences = [s.strip() for s in document.split('。') if s.strip()]
            # 合并短句子
            examples = []
            current_example = ""
            for sentence in sentences:
                if len(current_example + sentence) < 200:  # 控制示例长度
                    current_example += sentence + "。"
                else:
                    if current_example:
                        examples.append(current_example.strip())
                    current_example = sentence + "。"

                if len(examples) >= max_examples:
                    break

            if current_example and len(examples) < max_examples:
                examples.append(current_example.strip())

            return examples[:max_examples]
        else:
            return paragraphs[:max_examples]
