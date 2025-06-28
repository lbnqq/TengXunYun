"""
增强的文档填充器
集成专利分析、图片处理、AI智能填写等功能
"""

import re
import json
import os
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from .patent_document_analyzer import PatentDocumentAnalyzer
from .intelligent_image_processor import IntelligentImageProcessor
from .complex_document_filler import ComplexDocumentFiller

class EnhancedDocumentFiller:
    """增强的文档填充器"""
    
    def __init__(self, llm_client=None):
        self.llm_client = llm_client
        self.tool_name = "增强的文档填充器"
        self.description = "集成专利分析、图片处理、AI智能填写的完整文档处理系统"
        
        # 初始化各个组件
        self.patent_analyzer = PatentDocumentAnalyzer(llm_client)
        self.image_processor = IntelligentImageProcessor()
        self.document_filler = ComplexDocumentFiller(llm_client)
        
        # 文档类型识别器
        self.document_type_patterns = {
            "patent": [
                r"专利.*申请.*书",
                r"发明.*申请.*书",
                r"实用新型.*申请.*书",
                r"外观设计.*申请.*书"
            ],
            "project": [
                r"项目.*申请.*表",
                r"课题.*申请.*书",
                r"基金.*申请.*书"
            ],
            "contract": [
                r"合同.*书",
                r"协议.*书",
                r"意向.*书"
            ],
            "report": [
                r"报告.*书",
                r"总结.*报告",
                r"分析.*报告"
            ]
        }
    
    def analyze_document_structure(self, document_content: str, document_name: str = None) -> Dict[str, Any]:
        """
        分析文档结构（增强版）
        
        Args:
            document_content: 文档内容
            document_name: 文档名称
            
        Returns:
            增强的分析结果
        """
        try:
            # 1. 识别文档类型
            document_type = self._identify_document_type(document_content, document_name)
            
            # 2. 根据文档类型选择分析器
            if document_type == "patent":
                analysis_result = self.patent_analyzer.analyze_patent_document(document_content, document_name)
            else:
                # 使用通用文档分析器
                analysis_result = self.document_filler.analyze_document_structure(document_content, document_name)
            
            # 3. 增强分析结果
            enhanced_result = self._enhance_analysis_result(analysis_result, document_type, document_content)
            
            return enhanced_result
            
        except Exception as e:
            return {"error": f"文档结构分析失败: {str(e)}"}
    
    def _identify_document_type(self, content: str, document_name: str = None) -> str:
        """识别文档类型"""
        # 检查文档名称
        if document_name:
            for doc_type, patterns in self.document_type_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, document_name, re.IGNORECASE):
                        return doc_type
        
        # 检查文档内容
        for doc_type, patterns in self.document_type_patterns.items():
            for pattern in patterns:
                if re.search(pattern, content, re.IGNORECASE):
                    return doc_type
        
        return "general"
    
    def _enhance_analysis_result(self, analysis_result: Dict[str, Any], 
                               document_type: str, content: str) -> Dict[str, Any]:
        """增强分析结果"""
        if "error" in analysis_result:
            return analysis_result
        
        # 添加文档类型信息
        analysis_result["document_type"] = document_type
        
        # 添加AI填写建议
        if document_type == "patent":
            ai_suggestions = self.patent_analyzer.generate_ai_fill_suggestions(analysis_result)
            analysis_result["ai_suggestions"] = ai_suggestions
        
        # 添加图片处理信息
        image_positions = analysis_result.get("image_positions", [])
        if image_positions:
            analysis_result["image_processing_required"] = True
            analysis_result["image_count"] = len(image_positions)
        else:
            analysis_result["image_processing_required"] = False
            analysis_result["image_count"] = 0
        
        # 添加智能填写策略
        analysis_result["fill_strategy"] = self._generate_fill_strategy(analysis_result, document_type)
        
        return analysis_result
    
    def _generate_fill_strategy(self, analysis_result: Dict[str, Any], document_type: str) -> Dict[str, Any]:
        """生成填写策略"""
        strategy = {
            "document_type": document_type,
            "auto_fill_enabled": True,
            "ai_assisted_fill": True,
            "image_integration": analysis_result.get("image_processing_required", False),
            "validation_rules": [],
            "consistency_checks": []
        }
        
        if document_type == "patent":
            strategy["validation_rules"] = [
                "发明名称不能超过100字符",
                "摘要长度在100-500字符之间",
                "技术领域必须从预定义选项中选择",
                "申请日期格式为YYYY-MM-DD"
            ]
            strategy["consistency_checks"] = [
                "发明名称与摘要内容一致性",
                "技术领域与发明内容匹配",
                "附图说明与权利要求一致"
            ]
        
        return strategy
    
    def intelligent_fill_document(self, analysis_result: Dict[str, Any], 
                                user_data: Dict[str, Any] = None,
                                image_files: List[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        智能填充文档（增强版）
        
        Args:
            analysis_result: 文档分析结果
            user_data: 用户提供的数据
            image_files: 图片文件列表
            
        Returns:
            填充结果
        """
        try:
            document_type = analysis_result.get("document_type", "general")
            document_content = analysis_result.get("original_content", "")
            
            # 1. 处理图片文件
            if image_files and analysis_result.get("image_processing_required", False):
                image_result = self._process_document_images(image_files, analysis_result)
                if "error" not in image_result:
                    document_content = image_result["updated_document"]
            
            # 2. 生成AI填写内容
            ai_filled_data = self._generate_ai_filled_content(analysis_result, user_data)
            
            # 3. 合并用户数据和AI数据
            combined_data = self._merge_user_and_ai_data(user_data, ai_filled_data)
            
            # 4. 填充文档
            if document_type == "patent":
                fill_result = self._fill_patent_document(analysis_result, combined_data, document_content)
            else:
                fill_result = self.document_filler.fill_document(analysis_result, combined_data)
            
            # 5. 增强结果
            enhanced_result = self._enhance_fill_result(fill_result, analysis_result, combined_data)
            
            return enhanced_result
            
        except Exception as e:
            return {"error": f"智能文档填充失败: {str(e)}"}
    
    def generate_fill_preview(self, analysis_result: Dict[str, Any], 
                            user_data: Dict[str, Any] = None,
                            image_files: List[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        生成填充预览
        
        Args:
            analysis_result: 文档分析结果
            user_data: 用户提供的数据
            image_files: 图片文件列表
            
        Returns:
            预览结果
        """
        try:
            # 1. 验证输入参数
            if not analysis_result or "error" in analysis_result:
                return {
                    "success": False,
                    "error": "文档分析结果无效"
                }
            
            document_type = analysis_result.get("document_type", "general")
            document_content = analysis_result.get("original_content", "")
            
            # 2. 生成AI填写内容（仅预览，不保存）
            ai_filled_data = self._generate_ai_filled_content(analysis_result, user_data)
            
            # 3. 合并用户数据和AI数据
            combined_data = self._merge_user_and_ai_data(user_data, ai_filled_data)
            
            # 4. 生成预览内容
            preview_content = self._generate_preview_content(analysis_result, combined_data, document_content)
            
            # 5. 生成预览报告
            preview_report = self._generate_preview_report(analysis_result, combined_data, ai_filled_data)
            
            # 6. 生成字段映射信息
            field_mapping = self._generate_field_mapping(analysis_result, combined_data)
            
            return {
                "success": True,
                "preview_id": f"preview_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                "document_type": document_type,
                "preview_content": preview_content,
                "preview_report": preview_report,
                "field_mapping": field_mapping,
                "data_summary": {
                    "total_fields": len(analysis_result.get("fields", [])),
                    "filled_fields": len([f for f in combined_data.values() if f]),
                    "ai_generated_fields": len([f for f in ai_filled_data.values() if f]),
                    "user_provided_fields": len([f for f in (user_data or {}).values() if f]),
                    "unfilled_fields": len([f for f in analysis_result.get("fields", []) if not combined_data.get(f.get("name"))])
                },
                "quality_metrics": self._calculate_preview_quality(combined_data, analysis_result),
                "recommendations": self._generate_preview_recommendations(combined_data, analysis_result),
                "generated_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"生成填充预览失败: {str(e)}"
            }
    
    def _generate_preview_content(self, analysis_result: Dict[str, Any], 
                                combined_data: Dict[str, Any], 
                                original_content: str) -> str:
        """生成预览内容"""
        try:
            # 创建预览内容
            preview_lines = []
            preview_lines.append("# 文档填充预览")
            preview_lines.append(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            preview_lines.append(f"文档类型: {analysis_result.get('document_type', 'general')}")
            preview_lines.append("")
            
            # 添加字段填充信息
            fields = analysis_result.get("fields", [])
            for field in fields:
                field_name = field.get("name", "")
                field_value = combined_data.get(field_name, "")
                field_type = field.get("type", "text")
                
                preview_lines.append(f"## {field_name}")
                preview_lines.append(f"类型: {field_type}")
                preview_lines.append(f"值: {field_value}")
                preview_lines.append("")
            
            # 添加原始内容（如果较短）
            if len(original_content) < 1000:
                preview_lines.append("## 原始内容")
                preview_lines.append(original_content)
            
            return "\n".join(preview_lines)
            
        except Exception as e:
            return f"预览内容生成失败: {str(e)}"
    
    def _generate_preview_report(self, analysis_result: Dict[str, Any], 
                               combined_data: Dict[str, Any], 
                               ai_filled_data: Dict[str, Any]) -> Dict[str, Any]:
        """生成预览报告"""
        try:
            fields = analysis_result.get("fields", [])
            
            # 统计信息
            total_fields = len(fields)
            filled_fields = len([f for f in combined_data.values() if f])
            ai_generated = len([f for f in ai_filled_data.values() if f])
            empty_fields = total_fields - filled_fields
            
            # 字段状态
            field_status = []
            for field in fields:
                field_name = field.get("name", "")
                field_value = combined_data.get(field_name, "")
                is_ai_generated = field_name in ai_filled_data and ai_filled_data[field_name]
                
                field_status.append({
                    "name": field_name,
                    "type": field.get("type", "text"),
                    "filled": bool(field_value),
                    "ai_generated": is_ai_generated,
                    "value_preview": field_value[:100] + "..." if len(field_value) > 100 else field_value
                })
            
            return {
                "summary": {
                    "total_fields": total_fields,
                    "filled_fields": filled_fields,
                    "empty_fields": empty_fields,
                    "ai_generated_fields": ai_generated,
                    "fill_rate": filled_fields / total_fields if total_fields > 0 else 0
                },
                "field_status": field_status,
                "quality_assessment": self._assess_preview_quality(combined_data, analysis_result)
            }
            
        except Exception as e:
            return {"error": f"预览报告生成失败: {str(e)}"}
    
    def _generate_field_mapping(self, analysis_result: Dict[str, Any], 
                              combined_data: Dict[str, Any]) -> Dict[str, Any]:
        """生成字段映射信息"""
        try:
            fields = analysis_result.get("fields", [])
            mapping = {
                "exact_matches": [],
                "partial_matches": [],
                "unmatched_fields": [],
                "confidence_scores": {}
            }
            
            for field in fields:
                field_name = field.get("name", "")
                field_value = combined_data.get(field_name, "")
                
                if field_value:
                    # 计算匹配度
                    confidence = self._calculate_field_confidence(field, field_value)
                    mapping["confidence_scores"][field_name] = confidence
                    
                    if confidence >= 0.8:
                        mapping["exact_matches"].append(field_name)
                    elif confidence >= 0.5:
                        mapping["partial_matches"].append(field_name)
                    else:
                        mapping["unmatched_fields"].append(field_name)
                else:
                    mapping["unmatched_fields"].append(field_name)
                    mapping["confidence_scores"][field_name] = 0.0
            
            return mapping
            
        except Exception as e:
            return {"error": f"字段映射生成失败: {str(e)}"}
    
    def _calculate_field_confidence(self, field: Dict[str, Any], value: str) -> float:
        """计算字段匹配置信度"""
        try:
            confidence = 0.5  # 基础置信度
            
            # 根据字段类型调整置信度
            field_type = field.get("type", "text")
            if field_type == "date" and self._validate_date_format(value):
                confidence += 0.3
            elif field_type == "email" and self._validate_email_format(value):
                confidence += 0.3
            elif field_type == "phone" and self._validate_phone_format(value):
                confidence += 0.3
            
            # 根据内容长度调整置信度
            if len(value) >= field.get("min_length", 0):
                confidence += 0.1
            
            if len(value) <= field.get("max_length", 1000):
                confidence += 0.1
            
            return min(1.0, confidence)
            
        except Exception:
            return 0.0
    
    def _calculate_preview_quality(self, combined_data: Dict[str, Any], 
                                 analysis_result: Dict[str, Any]) -> Dict[str, Any]:
        """计算预览质量指标"""
        try:
            total_fields = len(analysis_result.get("fields", []))
            filled_fields = len([v for v in combined_data.values() if v])
            
            quality_metrics = {
                "completeness": filled_fields / total_fields if total_fields > 0 else 0,
                "average_length": sum(len(str(v)) for v in combined_data.values() if v) / max(filled_fields, 1),
                "validation_score": self._calculate_validation_score(combined_data, analysis_result),
                "consistency_score": self._calculate_consistency_score(combined_data, analysis_result)
            }
            
            return quality_metrics
            
        except Exception as e:
            return {"error": f"质量计算失败: {str(e)}"}
    
    def _calculate_validation_score(self, combined_data: Dict[str, Any], 
                                  analysis_result: Dict[str, Any]) -> float:
        """计算验证分数"""
        try:
            fields = analysis_result.get("fields", [])
            valid_fields = 0
            
            for field in fields:
                field_name = field.get("name", "")
                field_value = combined_data.get(field_name, "")
                
                if field_value and self._validate_field_type(field.get("type", "text"), field_value):
                    valid_fields += 1
            
            return valid_fields / len(fields) if fields else 0.0
            
        except Exception:
            return 0.0
    
    def _calculate_consistency_score(self, combined_data: Dict[str, Any], 
                                   analysis_result: Dict[str, Any]) -> float:
        """计算一致性分数"""
        try:
            # 简化的相关性检查
            score = 0.5  # 基础分数
            
            # 检查日期格式一致性
            date_fields = [f for f in analysis_result.get("fields", []) if f.get("type") == "date"]
            if len(date_fields) > 1:
                date_values = [combined_data.get(f.get("name", "")) for f in date_fields]
                if all(self._validate_date_format(v) for v in date_values if v):
                    score += 0.2
            
            # 检查必填字段完整性
            required_fields = [f for f in analysis_result.get("fields", []) if f.get("required", False)]
            filled_required = sum(1 for f in required_fields if combined_data.get(f.get("name", "")))
            if required_fields:
                score += 0.3 * (filled_required / len(required_fields))
            
            return min(1.0, score)
            
        except Exception:
            return 0.0
    
    def _assess_preview_quality(self, combined_data: Dict[str, Any], 
                              analysis_result: Dict[str, Any]) -> Dict[str, Any]:
        """评估预览质量"""
        try:
            quality_metrics = self._calculate_preview_quality(combined_data, analysis_result)
            
            # 质量等级评估
            completeness = quality_metrics.get("completeness", 0)
            validation_score = quality_metrics.get("validation_score", 0)
            consistency_score = quality_metrics.get("consistency_score", 0)
            
            overall_score = (completeness + validation_score + consistency_score) / 3
            
            if overall_score >= 0.8:
                quality_level = "excellent"
                assessment = "预览质量优秀，可以立即应用"
            elif overall_score >= 0.6:
                quality_level = "good"
                assessment = "预览质量良好，建议小幅调整"
            elif overall_score >= 0.4:
                quality_level = "fair"
                assessment = "预览质量一般，需要中等程度修改"
            else:
                quality_level = "poor"
                assessment = "预览质量较差，需要大量修改"
            
            return {
                "overall_score": overall_score,
                "quality_level": quality_level,
                "assessment": assessment,
                "metrics": quality_metrics,
                "improvement_suggestions": self._generate_quality_suggestions(quality_metrics)
            }
            
        except Exception as e:
            return {"error": f"质量评估失败: {str(e)}"}
    
    def _generate_quality_suggestions(self, quality_metrics: Dict[str, Any]) -> List[str]:
        """生成质量改进建议"""
        suggestions = []
        
        completeness = quality_metrics.get("completeness", 0)
        if completeness < 0.8:
            suggestions.append("建议填写更多字段以提高完整性")
        
        validation_score = quality_metrics.get("validation_score", 0)
        if validation_score < 0.9:
            suggestions.append("建议检查字段格式和验证规则")
        
        consistency_score = quality_metrics.get("consistency_score", 0)
        if consistency_score < 0.8:
            suggestions.append("建议检查字段间的一致性和相关性")
        
        if not suggestions:
            suggestions.append("预览质量良好，可以继续")
        
        return suggestions
    
    def _generate_preview_recommendations(self, combined_data: Dict[str, Any], 
                                        analysis_result: Dict[str, Any]) -> List[str]:
        """生成预览建议"""
        recommendations = []
        
        # 检查必填字段
        required_fields = [f for f in analysis_result.get("fields", []) if f.get("required", False)]
        missing_required = [f.get("name") for f in required_fields if not combined_data.get(f.get("name", ""))]
        
        if missing_required:
            recommendations.append(f"需要填写必填字段: {', '.join(missing_required)}")
        
        # 检查字段长度
        for field in analysis_result.get("fields", []):
            field_name = field.get("name", "")
            field_value = combined_data.get(field_name, "")
            
            if field_value:
                min_length = field.get("min_length", 0)
                max_length = field.get("max_length", 1000)
                
                if len(field_value) < min_length:
                    recommendations.append(f"字段 '{field_name}' 内容过短，建议补充")
                elif len(field_value) > max_length:
                    recommendations.append(f"字段 '{field_name}' 内容过长，建议精简")
        
        # 检查文档类型特定建议
        document_type = analysis_result.get("document_type", "general")
        if document_type == "patent":
            recommendations.append("专利文档建议检查技术术语的准确性")
        elif document_type == "contract":
            recommendations.append("合同文档建议检查法律条款的完整性")
        
        if not recommendations:
            recommendations.append("预览内容符合要求，可以应用")
        
        return recommendations
    
    def _process_document_images(self, image_files: List[Dict[str, Any]], 
                               analysis_result: Dict[str, Any]) -> Dict[str, Any]:
        """处理文档图片"""
        try:
            image_positions = analysis_result.get("image_positions", [])
            document_content = analysis_result.get("original_content", "")
            
            # 批量处理图片
            batch_result = self.image_processor.batch_process_images(image_files, document_content)
            
            if "error" in batch_result:
                return batch_result
            
            return {
                "success": True,
                "processed_images": batch_result["processed_images"],
                "updated_document": batch_result["updated_document"],
                "image_count": len(batch_result["processed_images"])
            }
            
        except Exception as e:
            return {"error": f"图片处理失败: {str(e)}"}
    
    def _generate_ai_filled_content(self, analysis_result: Dict[str, Any], 
                                  user_data: Dict[str, Any] = None) -> Dict[str, Any]:
        """生成AI填写内容"""
        try:
            ai_filled_data = {}
            fields = analysis_result.get("fields", [])
            total_fields = len(fields)
            completed_fields = 0
            
            # 获取文档类型和用户角色信息
            document_type = analysis_result.get("document_type", "general")
            user_role = self._identify_user_role(analysis_result, user_data)
            
            print(f"🤖 开始AI内容生成，共{total_fields}个字段需要处理...")
            print(f"📝 文档类型: {document_type}, 用户角色: {user_role}")
            
            for field in fields:
                field_id = field["field_id"]
                field_name = field["field_name"]
                field_type = field["field_type"]
                ai_prompt = field.get("ai_fill_prompt", "")
                
                # 如果用户没有提供数据，使用AI生成
                if not user_data or field_id not in user_data:
                    print(f"🧠 正在为字段'{field_name}'生成内容...")
                    
                    if self.llm_client and ai_prompt:
                        # 构建增强的AI提示词
                        enhanced_prompt = self._build_enhanced_ai_prompt(
                            field, analysis_result, user_data, user_role
                        )
                        
                        # 带重试机制的AI内容生成
                        ai_content = self._generate_ai_content_with_retry(enhanced_prompt, field)
                        
                        # 内容质量验证
                        validation_result = self._validate_ai_generated_content(ai_content, field, document_type)
                        
                        if validation_result["is_valid"]:
                            ai_filled_data[field_id] = {
                                "content": ai_content.strip(),
                                "source": "ai_generated",
                                "confidence": validation_result["confidence"],
                                "quality_score": validation_result["quality_score"],
                                "validation_notes": validation_result.get("notes", [])
                            }
                            print(f"✅ 字段'{field_name}'内容生成成功，质量评分: {validation_result['quality_score']:.1f}")
                        else:
                            # 如果验证失败，使用改进后的内容或默认值
                            improved_content = self._improve_ai_content(ai_content, field, validation_result)
                            ai_filled_data[field_id] = {
                                "content": improved_content,
                                "source": "ai_generated_improved",
                                "confidence": validation_result["confidence"] * 0.8,
                                "quality_score": validation_result["quality_score"],
                                "validation_notes": validation_result.get("notes", [])
                            }
                            print(f"⚠️ 字段'{field_name}'内容已优化，质量评分: {validation_result['quality_score']:.1f}")
                    else:
                        # 使用默认值
                        default_content = self._get_default_field_value(field)
                        ai_filled_data[field_id] = {
                            "content": default_content,
                            "source": "default",
                            "confidence": 0.5,
                            "quality_score": 0.5,
                            "validation_notes": ["使用默认值"]
                        }
                        print(f"📋 字段'{field_name}'使用默认值")
                    
                    completed_fields += 1
                    print(f"📊 进度: {completed_fields}/{total_fields} ({completed_fields/total_fields*100:.1f}%)")
                else:
                    # 用户已提供数据，跳过AI生成
                    print(f"👤 字段'{field_name}'用户已提供数据，跳过AI生成")
            
            print(f"🎉 AI内容生成完成！共生成{len(ai_filled_data)}个字段的内容")
            
            # 生成整体质量报告
            quality_report = self._generate_quality_report(ai_filled_data, analysis_result)
            
            return {
                "ai_filled_data": ai_filled_data,
                "quality_report": quality_report,
                "generation_summary": {
                    "total_fields": total_fields,
                    "ai_generated": len([v for v in ai_filled_data.values() if v["source"] == "ai_generated"]),
                    "ai_generated_improved": len([v for v in ai_filled_data.values() if v["source"] == "ai_generated_improved"]),
                    "default_values": len([v for v in ai_filled_data.values() if v["source"] == "default"]),
                    "average_quality_score": sum(v["quality_score"] for v in ai_filled_data.values()) / len(ai_filled_data) if ai_filled_data else 0
                }
            }
            
        except Exception as e:
            print(f"❌ AI内容生成失败: {str(e)}")
            return {"error": f"AI内容生成失败: {str(e)}"}
    
    def _identify_user_role(self, analysis_result: Dict[str, Any], user_data: Dict[str, Any] = None) -> str:
        """识别用户角色"""
        document_type = analysis_result.get("document_type", "general")
        
        # 根据文档类型和内容推断用户角色
        if document_type == "patent":
            return "专利申请人"
        elif document_type == "project":
            return "项目负责人"
        elif document_type == "contract":
            return "合同起草人"
        elif document_type == "report":
            return "报告撰写人"
        else:
            # 从用户数据中推断角色
            if user_data:
                if "applicant" in user_data or "申请人" in str(user_data):
                    return "申请人"
                elif "author" in user_data or "作者" in str(user_data):
                    return "文档作者"
                elif "manager" in user_data or "负责人" in str(user_data):
                    return "项目负责人"
            
            return "文档填报人"
    
    def _build_enhanced_ai_prompt(self, field: Dict[str, Any], 
                                analysis_result: Dict[str, Any], 
                                user_data: Dict[str, Any] = None,
                                user_role: str = "文档填报人") -> str:
        """构建增强的AI提示词"""
        document_objective = analysis_result.get("total_objective", "")
        core_theme = analysis_result.get("core_theme", "")
        field_name = field["field_name"]
        field_type = field["field_type"]
        ai_prompt = field.get("ai_fill_prompt", "")
        document_type = analysis_result.get("document_type", "general")
        
        # 构建上下文信息
        context_info = f"""
        文档目标：{document_objective}
        核心主题：{core_theme}
        文档类型：{document_type}
        字段名称：{field_name}
        字段类型：{field_type}
        用户角色：{user_role}
        """
        
        if user_data:
            context_info += f"\n用户已提供数据：{json.dumps(user_data, ensure_ascii=False)}"
        
        # 构建约束信息
        constraints = field.get("constraints", {})
        constraint_info = ""
        if constraints.get("min_length"):
            constraint_info += f"\n最小长度：{constraints['min_length']}字符"
        if constraints.get("max_length"):
            constraint_info += f"\n最大长度：{constraints['max_length']}字符"
        if constraints.get("required"):
            constraint_info += f"\n必填字段"
        if constraints.get("options"):
            constraint_info += f"\n可选值：{', '.join(constraints['options'])}"
        
        # 角色特定的写作要求
        role_requirements = self._get_role_specific_requirements(user_role, document_type)
        
        # 完整的提示词
        full_prompt = f"""
你是一位专业的{user_role}，正在起草一份{document_type}文档。请根据以下信息为字段"{field_name}"生成合适的内容：

{context_info}

填写要求：{ai_prompt}

约束条件：{constraint_info}

{role_requirements}

【专业写作要求】
1. 身份定位：以{user_role}的身份和视角进行写作
2. 专业术语：使用该领域常用的专业术语和表达方式
3. 逻辑清晰：内容结构合理，前后呼应
4. 表达准确：用词精准，避免模糊表述

【去除AIGC痕迹要求】
- 避免使用"首先、其次、最后"等机械化过渡词
- 不要使用"值得注意的是"、"需要强调的是"等AI常用表达
- 减少使用"进行"、"实施"、"开展"等动词
- 避免过度使用"的"字结构
- 不要出现"综上所述"、"总而言之"等总结性套话
- 使用自然流畅的表达方式，像人类专业写作

【内容质量要求】
- 信息准确：确保所有事实和数据的准确性
- 针对性强：根据具体场景和用户需求定制内容
- 专业度适中：既要专业又要通俗易懂
- 符合规范：遵循该类型文档的写作规范

请直接返回填写内容，不要包含解释。内容应该自然、专业、符合{user_role}的写作风格。
"""
        
        return full_prompt
    
    def _get_role_specific_requirements(self, user_role: str, document_type: str) -> str:
        """获取角色特定的写作要求"""
        requirements = {
            "专利申请人": """
【专利申请人写作要求】
- 技术描述准确：使用精确的技术术语
- 创新点突出：强调发明的创新性和实用性
- 逻辑严密：技术方案描述逻辑清晰
- 客观陈述：避免主观评价，客观描述技术特征
""",
            "项目负责人": """
【项目负责人写作要求】
- 管理视角：从项目管理角度描述
- 目标明确：突出项目目标和预期成果
- 计划周密：体现项目管理的系统性
- 责任清晰：明确各方职责和分工
""",
            "合同起草人": """
【合同起草人写作要求】
- 条款严谨：使用准确的法律术语
- 权利义务明确：清晰界定各方权利义务
- 风险控制：体现风险防范意识
- 可执行性强：条款具有可操作性
""",
            "报告撰写人": """
【报告撰写人写作要求】
- 数据支撑：基于事实和数据进行分析
- 结论明确：提出明确的结论和建议
- 逻辑清晰：报告结构合理，层次分明
- 客观公正：保持客观中立的立场
""",
            "文档填报人": """
【文档填报人写作要求】
- 信息完整：确保填写信息完整准确
- 格式规范：符合文档格式要求
- 内容真实：基于实际情况填写
- 表达清晰：使用简洁明了的语言
"""
        }
        
        return requirements.get(user_role, requirements["文档填报人"])
    
    def _generate_ai_content_with_retry(self, prompt: str, field: Dict[str, Any], max_retries: int = 3) -> str:
        """带重试机制的AI内容生成"""
        field_name = field.get("field_name", "未知字段")
        
        for attempt in range(max_retries):
            try:
                print(f"🧠 第{attempt + 1}次尝试生成字段'{field_name}'的内容...")
                content = self.llm_client.generate(prompt)
                
                if content and content.strip():
                    print(f"✅ 字段'{field_name}'内容生成成功")
                    return content.strip()
                else:
                    print(f"⚠️ 字段'{field_name}'生成内容为空，重试中...")
                    
            except Exception as e:
                print(f"❌ 字段'{field_name}'第{attempt + 1}次生成失败: {str(e)}")
                if attempt < max_retries - 1:
                    import time
                    time.sleep(1)  # 重试前等待
                else:
                    print(f"💡 字段'{field_name}'所有重试失败，使用默认值")
        
        # 所有重试失败，返回默认值
        return self._get_default_field_value(field)
    
    def _validate_ai_generated_content(self, content: str, field: Dict[str, Any], document_type: str) -> Dict[str, Any]:
        """验证AI生成内容的质量"""
        validation_result = {
            "is_valid": True,
            "confidence": 0.8,
            "quality_score": 0.8,
            "notes": [],
            "issues": []
        }
        
        field_name = field.get("field_name", "")
        field_type = field.get("field_type", "")
        constraints = field.get("constraints", {})
        
        # 1. 基础验证
        if not content or not content.strip():
            validation_result["is_valid"] = False
            validation_result["confidence"] = 0.0
            validation_result["quality_score"] = 0.0
            validation_result["issues"].append("内容为空")
            validation_result["notes"].append("生成内容为空，需要重新生成")
            return validation_result
        
        # 2. 长度验证
        content_length = len(content.strip())
        min_length = constraints.get("min_length", 0)
        max_length = constraints.get("max_length", 1000)
        
        if content_length < min_length:
            validation_result["issues"].append(f"内容过短({content_length}字符，要求至少{min_length}字符)")
            validation_result["quality_score"] *= 0.7
        
        if content_length > max_length:
            validation_result["issues"].append(f"内容过长({content_length}字符，要求最多{max_length}字符)")
            validation_result["quality_score"] *= 0.8
        
        # 3. 格式验证
        if field_type == "date":
            if not self._validate_date_format(content):
                validation_result["issues"].append("日期格式不正确")
                validation_result["quality_score"] *= 0.6
        
        elif field_type == "email":
            if not self._validate_email_format(content):
                validation_result["issues"].append("邮箱格式不正确")
                validation_result["quality_score"] *= 0.6
        
        elif field_type == "phone":
            if not self._validate_phone_format(content):
                validation_result["issues"].append("电话格式不正确")
                validation_result["quality_score"] *= 0.6
        
        # 4. AIGC痕迹检测
        aigc_score = self._detect_aigc_traces(content)
        if aigc_score > 0.7:
            validation_result["issues"].append("存在明显的AI生成痕迹")
            validation_result["quality_score"] *= (1 - aigc_score * 0.3)
            validation_result["notes"].append("建议优化表达方式，去除AI痕迹")
        
        # 5. 专业度评估
        professionalism_score = self._assess_professionalism(content, document_type)
        validation_result["quality_score"] *= professionalism_score
        
        # 6. 相关性检查
        relevance_score = self._check_content_relevance(content, field)
        validation_result["quality_score"] *= relevance_score
        
        # 7. 最终质量评分
        if validation_result["quality_score"] < 0.6:
            validation_result["is_valid"] = False
            validation_result["confidence"] *= 0.8
        
        # 8. 生成改进建议
        if validation_result["issues"]:
            validation_result["notes"].extend([
                f"发现{len(validation_result['issues'])}个问题需要改进",
                "建议重新生成或手动优化内容"
            ])
        
        return validation_result
    
    def _validate_date_format(self, content: str) -> bool:
        """验证日期格式"""
        import re
        date_patterns = [
            r'\d{4}-\d{2}-\d{2}',  # YYYY-MM-DD
            r'\d{4}年\d{1,2}月\d{1,2}日',  # YYYY年MM月DD日
            r'\d{1,2}/\d{1,2}/\d{4}',  # MM/DD/YYYY
        ]
        
        for pattern in date_patterns:
            if re.match(pattern, content.strip()):
                return True
        return False
    
    def _validate_email_format(self, content: str) -> bool:
        """验证邮箱格式"""
        import re
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return bool(re.match(email_pattern, content.strip()))
    
    def _validate_phone_format(self, content: str) -> bool:
        """验证电话格式"""
        import re
        phone_patterns = [
            r'^1[3-9]\d{9}$',  # 手机号
            r'^0\d{2,3}-\d{7,8}$',  # 座机号
            r'^\+86-1[3-9]\d{9}$',  # 国际手机号
        ]
        
        content_clean = re.sub(r'[\s\-\(\)]', '', content.strip())
        for pattern in phone_patterns:
            if re.match(pattern, content_clean):
                return True
        return False
    
    def _detect_aigc_traces(self, content: str) -> float:
        """检测AIGC痕迹"""
        aigc_indicators = [
            "首先", "其次", "最后", "第一", "第二", "第三",
            "值得注意的是", "需要强调的是", "特别指出",
            "进行", "实施", "开展", "推进", "推动",
            "综上所述", "总而言之", "总的来说",
            "的的", "非常非常", "特别特别",
            "让我们", "我们来看", "我们可以看到",
            "根据", "基于", "鉴于", "由于",
            "因此", "所以", "故而", "从而"
        ]
        
        content_lower = content.lower()
        aigc_count = 0
        
        for indicator in aigc_indicators:
            if indicator in content_lower:
                aigc_count += 1
        
        # 计算AIGC痕迹分数
        aigc_score = min(aigc_count / 5.0, 1.0)  # 最多5个指标，分数范围0-1
        
        return aigc_score
    
    def _assess_professionalism(self, content: str, document_type: str) -> float:
        """评估专业度"""
        # 根据文档类型定义专业词汇
        professional_terms = {
            "patent": ["发明", "技术方案", "技术特征", "权利要求", "实施例", "背景技术"],
            "project": ["项目", "目标", "计划", "实施", "评估", "成果"],
            "contract": ["合同", "协议", "条款", "义务", "权利", "违约责任"],
            "report": ["报告", "分析", "结论", "建议", "数据", "评估"]
        }
        
        terms = professional_terms.get(document_type, [])
        if not terms:
            return 0.8  # 默认专业度
        
        content_lower = content.lower()
        term_count = sum(1 for term in terms if term in content_lower)
        
        # 专业度评分：0.6-1.0
        professionalism_score = 0.6 + (term_count / len(terms)) * 0.4
        return min(professionalism_score, 1.0)
    
    def _check_content_relevance(self, content: str, field: Dict[str, Any]) -> float:
        """检查内容相关性"""
        field_name = field.get("field_name", "").lower()
        content_lower = content.lower()
        
        # 根据字段名称关键词检查相关性
        relevance_keywords = {
            "名称": ["名称", "名字", "标题", "题目"],
            "日期": ["日期", "时间", "年月日"],
            "地址": ["地址", "地点", "位置", "住址"],
            "电话": ["电话", "手机", "联系方式"],
            "邮箱": ["邮箱", "邮件", "email"],
            "金额": ["金额", "费用", "价格", "成本"],
            "描述": ["描述", "说明", "介绍", "阐述"]
        }
        
        for field_type, keywords in relevance_keywords.items():
            if any(keyword in field_name for keyword in keywords):
                # 检查内容是否包含相关词汇
                relevance_count = sum(1 for keyword in keywords if keyword in content_lower)
                if relevance_count > 0:
                    return 0.9  # 高相关性
                else:
                    return 0.6  # 中等相关性
        
        return 0.8  # 默认相关性
    
    def _improve_ai_content(self, content: str, field: Dict[str, Any], validation_result: Dict[str, Any]) -> str:
        """改进AI生成的内容"""
        improved_content = content
        
        # 根据验证结果进行改进
        issues = validation_result.get("issues", [])
        
        for issue in issues:
            if "AIGC痕迹" in issue:
                improved_content = self._remove_aigc_traces(improved_content)
            elif "内容过短" in issue:
                improved_content = self._expand_content(improved_content, field)
            elif "内容过长" in issue:
                improved_content = self._shorten_content(improved_content, field)
        
        return improved_content
    
    def _remove_aigc_traces(self, content: str) -> str:
        """去除AIGC痕迹"""
        import re
        
        # 替换AIGC常用表达
        replacements = {
            r'首先': '开始',
            r'其次': '接着',
            r'最后': '最终',
            r'值得注意的是': '重要的是',
            r'需要强调的是': '要强调的是',
            r'特别指出': '特别说明',
            r'进行': '做',
            r'实施': '执行',
            r'开展': '进行',
            r'综上所述': '总之',
            r'总而言之': '总的来说',
            r'总的来说': '总体而言',
            r'让我们': '我们',
            r'我们来看': '我们看',
            r'我们可以看到': '我们看',
        }
        
        improved_content = content
        for pattern, replacement in replacements.items():
            improved_content = re.sub(pattern, replacement, improved_content)
        
        # 减少重复的"的"字
        improved_content = re.sub(r'的的+', '的', improved_content)
        
        # 减少过度修饰
        improved_content = re.sub(r'非常非常', '非常', improved_content)
        improved_content = re.sub(r'特别特别', '特别', improved_content)
        
        return improved_content
    
    def _expand_content(self, content: str, field: Dict[str, Any]) -> str:
        """扩展内容"""
        field_name = field.get("field_name", "")
        
        # 根据字段类型添加详细信息
        if "描述" in field_name or "说明" in field_name:
            if len(content) < 50:
                content += "。具体情况需要根据实际项目要求进行详细说明。"
        elif "地址" in field_name:
            if len(content) < 20:
                content += "，具体门牌号待补充"
        
        return content
    
    def _shorten_content(self, content: str, field: Dict[str, Any]) -> str:
        """缩短内容"""
        max_length = field.get("constraints", {}).get("max_length", 1000)
        
        if len(content) > max_length:
            # 保留前max_length个字符，确保句子完整
            shortened = content[:max_length]
            last_period = shortened.rfind('。')
            if last_period > max_length * 0.8:  # 如果句号位置合理
                return shortened[:last_period + 1]
            else:
                return shortened + "..."
        
        return content
    
    def _generate_quality_report(self, ai_filled_data: Dict[str, Any], analysis_result: Dict[str, Any]) -> Dict[str, Any]:
        """生成质量报告"""
        total_fields = len(ai_filled_data)
        if total_fields == 0:
            return {"error": "没有生成任何内容"}
        
        quality_scores = [v["quality_score"] for v in ai_filled_data.values()]
        avg_quality = sum(quality_scores) / len(quality_scores)
        
        # 统计各质量等级的数量
        excellent_count = len([s for s in quality_scores if s >= 0.9])
        good_count = len([s for s in quality_scores if 0.7 <= s < 0.9])
        fair_count = len([s for s in quality_scores if 0.5 <= s < 0.7])
        poor_count = len([s for s in quality_scores if s < 0.5])
        
        # 统计来源分布
        source_stats = {}
        for data in ai_filled_data.values():
            source = data["source"]
            source_stats[source] = source_stats.get(source, 0) + 1
        
        return {
            "overall_quality_score": avg_quality,
            "quality_distribution": {
                "excellent": excellent_count,
                "good": good_count,
                "fair": fair_count,
                "poor": poor_count
            },
            "source_distribution": source_stats,
            "total_fields": total_fields,
            "recommendations": self._generate_quality_recommendations(avg_quality, source_stats)
        }
    
    def _generate_quality_recommendations(self, avg_quality: float, source_stats: Dict[str, int]) -> List[str]:
        """生成质量改进建议"""
        recommendations = []
        
        if avg_quality < 0.7:
            recommendations.append("整体内容质量有待提升，建议人工审核和优化")
        
        if source_stats.get("default", 0) > 0:
            recommendations.append("部分字段使用了默认值，建议补充具体信息")
        
        if source_stats.get("ai_generated_improved", 0) > 0:
            recommendations.append("部分内容经过优化，建议进一步人工完善")
        
        if avg_quality >= 0.8:
            recommendations.append("内容质量良好，可以直接使用")
        
        return recommendations
    
    def _merge_user_and_ai_data(self, user_data: Dict[str, Any], 
                               ai_data: Dict[str, Any]) -> Dict[str, Any]:
        """合并用户数据和AI数据"""
        merged_data = {}
        
        # 添加AI生成的数据
        for field_id, ai_content in ai_data.items():
            if isinstance(ai_content, dict):
                merged_data[field_id] = ai_content["content"]
            else:
                merged_data[field_id] = ai_content
        
        # 用户数据优先覆盖AI数据
        if user_data:
            for field_id, user_content in user_data.items():
                merged_data[field_id] = user_content
        
        return merged_data
    
    def _fill_patent_document(self, analysis_result: Dict[str, Any], 
                            fill_data: Dict[str, Any], 
                            document_content: str) -> Dict[str, Any]:
        """填充专利文档"""
        try:
            # 使用专利分析器的特殊填充逻辑
            filled_content = document_content
            fields = analysis_result.get("fields", [])
            
            # 按字段填充
            for field in fields:
                field_id = field["field_id"]
                if field_id in fill_data:
                    field_value = fill_data[field_id]
                    match_text = field["match_text"]
                    
                    # 替换字段内容
                    if match_text in filled_content:
                        replacement = match_text + str(field_value)
                        filled_content = filled_content.replace(match_text, replacement, 1)
            
            # 生成HTML输出
            html_content = self._generate_patent_html_output(filled_content, analysis_result)
            
            return {
                "success": True,
                "filled_content": filled_content,
                "html_content": html_content,
                "fill_summary": self._generate_patent_fill_summary(analysis_result, fill_data),
                "download_ready": True
            }
            
        except Exception as e:
            return {"error": f"专利文档填充失败: {str(e)}"}
    
    def _generate_patent_html_output(self, content: str, analysis_result: Dict[str, Any]) -> str:
        """生成专利文档HTML输出"""
        document_name = analysis_result.get("document_name", "专利申请书")
        
        html_template = f"""
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{document_name}</title>
            <style>
                body {{
                    font-family: 'SimSun', 'Microsoft YaHei', serif;
                    line-height: 1.8;
                    margin: 40px;
                    color: #333;
                    background-color: #fff;
                }}
                .document-container {{
                    max-width: 800px;
                    margin: 0 auto;
                    background-color: #fff;
                    padding: 40px;
                    box-shadow: 0 0 20px rgba(0,0,0,0.1);
                    border-radius: 8px;
                }}
                .document-title {{
                    text-align: center;
                    font-size: 24px;
                    font-weight: bold;
                    margin-bottom: 30px;
                    color: #2c3e50;
                    border-bottom: 2px solid #3498db;
                    padding-bottom: 15px;
                }}
                .document-content {{
                    white-space: pre-wrap;
                    font-size: 16px;
                    text-align: justify;
                }}
                .filled-field {{
                    background-color: #e8f5e8;
                    padding: 2px 4px;
                    border-radius: 3px;
                    border-left: 3px solid #27ae60;
                }}
                .image-placeholder {{
                    background-color: #f8f9fa;
                    border: 2px dashed #dee2e6;
                    padding: 20px;
                    text-align: center;
                    margin: 10px 0;
                    border-radius: 5px;
                }}
                .download-btn {{
                    display: inline-block;
                    margin-top: 20px;
                    padding: 10px 20px;
                    background-color: #3498db;
                    color: white;
                    text-decoration: none;
                    border-radius: 5px;
                    font-weight: bold;
                }}
                .download-btn:hover {{
                    background-color: #2980b9;
                }}
            </style>
        </head>
        <body>
            <div class="document-container">
                <div class="document-title">{document_name}</div>
                <div class="document-content">{content}</div>
                <a href="#" class="download-btn" onclick="downloadDocument()">下载文档</a>
            </div>
            
            <script>
                function downloadDocument() {{
                    const content = document.querySelector('.document-container').innerHTML;
                    const blob = new Blob([content], {{type: 'text/html'}});
                    const url = URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = '{document_name}.html';
                    a.click();
                    URL.revokeObjectURL(url);
                }}
            </script>
        </body>
        </html>
        """
        
        return html_template
    
    def _generate_patent_fill_summary(self, analysis_result: Dict[str, Any], 
                                    fill_data: Dict[str, Any]) -> Dict[str, Any]:
        """生成专利文档填充摘要"""
        fields = analysis_result.get("fields", [])
        total_fields = len(fields)
        filled_fields = len(fill_data)
        
        return {
            "total_fields": total_fields,
            "filled_fields": filled_fields,
            "completion_rate": (filled_fields / total_fields * 100) if total_fields > 0 else 0,
            "document_type": "patent_application",
            "confidence_score": analysis_result.get("confidence_score", 0.0),
            "ai_generated_count": len([v for v in fill_data.values() if isinstance(v, str) and v.startswith("[") and v.endswith("]")]),
            "user_provided_count": filled_fields - len([v for v in fill_data.values() if isinstance(v, str) and v.startswith("[") and v.endswith("]")])
        }
    
    def _enhance_fill_result(self, fill_result: Dict[str, Any], 
                           analysis_result: Dict[str, Any], 
                           fill_data: Dict[str, Any]) -> Dict[str, Any]:
        """增强填充结果"""
        if "error" in fill_result:
            return fill_result
        
        # 添加元数据
        fill_result["metadata"] = {
            "fill_time": datetime.now().isoformat(),
            "document_type": analysis_result.get("document_type", "general"),
            "total_objective": analysis_result.get("total_objective", ""),
            "core_theme": analysis_result.get("core_theme", ""),
            "image_processed": analysis_result.get("image_processing_required", False),
            "ai_assisted": True
        }
        
        # 添加质量评估
        fill_result["quality_assessment"] = self._assess_fill_quality(analysis_result, fill_data)
        
        return fill_result
    
    def _assess_fill_quality(self, analysis_result: Dict[str, Any], 
                           fill_data: Dict[str, Any]) -> Dict[str, Any]:
        """评估填充质量"""
        fields = analysis_result.get("fields", [])
        total_fields = len(fields)
        filled_fields = len(fill_data)
        
        # 计算完成度
        completion_rate = (filled_fields / total_fields * 100) if total_fields > 0 else 0
        
        # 计算质量分数
        quality_score = 0.0
        
        # 完成度权重
        quality_score += completion_rate * 0.4
        
        # 字段类型匹配度
        type_match_score = 0.0
        for field in fields:
            field_id = field["field_id"]
            if field_id in fill_data:
                field_type = field["field_type"]
                field_value = fill_data[field_id]
                
                # 检查类型匹配
                if self._validate_field_type(field_type, field_value):
                    type_match_score += 1
        
        type_match_rate = (type_match_score / filled_fields * 100) if filled_fields > 0 else 0
        quality_score += type_match_rate * 0.3
        
        # 约束满足度
        constraint_score = 0.0
        for field in fields:
            field_id = field["field_id"]
            if field_id in fill_data:
                field_value = fill_data[field_id]
                constraints = field.get("constraints", {})
                
                if self._validate_field_constraints(field_value, constraints):
                    constraint_score += 1
        
        constraint_rate = (constraint_score / filled_fields * 100) if filled_fields > 0 else 0
        quality_score += constraint_rate * 0.3
        
        return {
            "overall_score": min(100, quality_score),
            "completion_rate": completion_rate,
            "type_match_rate": type_match_rate,
            "constraint_satisfaction_rate": constraint_rate,
            "assessment_time": datetime.now().isoformat()
        }
    
    def _validate_field_type(self, field_type: str, field_value: str) -> bool:
        """验证字段类型"""
        try:
            if field_type == "date":
                # 检查日期格式
                import re
                date_patterns = [
                    r'\d{4}-\d{2}-\d{2}',
                    r'\d{4}年\d{1,2}月\d{1,2}日',
                    r'\d{1,2}/\d{1,2}/\d{4}'
                ]
                return any(re.search(pattern, field_value) for pattern in date_patterns)
            elif field_type == "number":
                # 检查数字格式
                return field_value.replace('.', '').replace(',', '').isdigit()
            elif field_type == "email":
                # 检查邮箱格式
                import re
                email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
                return re.match(email_pattern, field_value) is not None
            else:
                return True
        except:
            return False
    
    def _validate_field_constraints(self, field_value: str, constraints: Dict[str, Any]) -> bool:
        """验证字段约束"""
        try:
            # 长度约束
            if "min_length" in constraints and len(field_value) < constraints["min_length"]:
                return False
            if "max_length" in constraints and len(field_value) > constraints["max_length"]:
                return False
            
            # 模式约束
            if "pattern" in constraints:
                import re
                if not re.match(constraints["pattern"], field_value):
                    return False
            
            # 选项约束
            if "options" in constraints and field_value not in constraints["options"]:
                return False
            
            return True
        except:
            return False

    def apply_fill_changes(self, analysis_result: Dict[str, Any], 
                          fill_data: Dict[str, Any],
                          image_files: List[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        应用填充变化到文档
        
        Args:
            analysis_result: 文档分析结果
            fill_data: 填充数据
            image_files: 图片文件列表
            
        Returns:
            应用结果
        """
        try:
            # 1. 验证输入参数
            if not analysis_result or "error" in analysis_result:
                return {
                    "success": False,
                    "error": "文档分析结果无效"
                }
            
            if not fill_data:
                return {
                    "success": False,
                    "error": "没有提供填充数据"
                }
            
            document_type = analysis_result.get("document_type", "general")
            document_content = analysis_result.get("original_content", "")
            
            # 2. 验证填充数据
            validation_result = self._validate_fill_data(fill_data, analysis_result)
            if not validation_result["valid"]:
                return {
                    "success": False,
                    "error": "填充数据验证失败",
                    "validation_errors": validation_result["errors"]
                }
            
            # 3. 处理图片文件
            processed_content = document_content
            image_processing_result = None
            if image_files and analysis_result.get("image_processing_required", False):
                image_processing_result = self._process_document_images(image_files, analysis_result)
                if "error" not in image_processing_result:
                    processed_content = image_processing_result["updated_document"]
            
            # 4. 应用填充数据
            if document_type == "patent":
                fill_result = self._apply_patent_fill_changes(analysis_result, fill_data, processed_content)
            else:
                fill_result = self._apply_general_fill_changes(analysis_result, fill_data, processed_content)
            
            # 5. 生成最终文档
            final_document = self._generate_final_document(fill_result, analysis_result, fill_data)
            
            # 6. 生成应用报告
            application_report = self._generate_fill_application_report(fill_result, analysis_result, fill_data)
            
            return {
                "success": True,
                "document_type": document_type,
                "final_document": final_document,
                "application_report": application_report,
                "fill_statistics": {
                    "total_fields": len(analysis_result.get("fields", [])),
                    "filled_fields": len([f for f in fill_data.values() if f]),
                    "fill_rate": len([f for f in fill_data.values() if f]) / len(analysis_result.get("fields", [])) if analysis_result.get("fields") else 0,
                    "image_count": len(image_files) if image_files else 0
                },
                "quality_assessment": self._assess_final_quality(fill_result, analysis_result, fill_data),
                "applied_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"应用填充变化失败: {str(e)}"
            }
    
    def _validate_fill_data(self, fill_data: Dict[str, Any], analysis_result: Dict[str, Any]) -> Dict[str, Any]:
        """验证填充数据"""
        try:
            errors = []
            fields = analysis_result.get("fields", [])
            
            for field in fields:
                field_name = field.get("name", "")
                field_value = fill_data.get(field_name, "")
                field_type = field.get("type", "text")
                required = field.get("required", False)
                
                # 检查必填字段
                if required and not field_value:
                    errors.append(f"必填字段 '{field_name}' 不能为空")
                    continue
                
                # 检查字段类型
                if field_value and not self._validate_field_type(field_type, field_value):
                    errors.append(f"字段 '{field_name}' 的值不符合类型要求: {field_type}")
                
                # 检查字段约束
                constraints = field.get("constraints", {})
                if field_value and not self._validate_field_constraints(field_value, constraints):
                    errors.append(f"字段 '{field_name}' 的值不符合约束要求")
            
            return {
                "valid": len(errors) == 0,
                "errors": errors
            }
            
        except Exception as e:
            return {
                "valid": False,
                "errors": [f"验证过程发生错误: {str(e)}"]
            }
    
    def _apply_patent_fill_changes(self, analysis_result: Dict[str, Any], 
                                 fill_data: Dict[str, Any], 
                                 document_content: str) -> Dict[str, Any]:
        """应用专利文档填充变化"""
        try:
            # 使用专利分析器的填充方法
            fill_result = self.patent_analyzer.fill_patent_document(analysis_result, fill_data, document_content)
            
            if "error" in fill_result:
                return fill_result
            
            # 增强结果
            enhanced_result = {
                "success": True,
                "filled_content": fill_result.get("filled_content", document_content),
                "fill_summary": fill_result.get("fill_summary", {}),
                "html_output": fill_result.get("html_output", ""),
                "document_type": "patent"
            }
            
            return enhanced_result
            
        except Exception as e:
            return {
                "success": False,
                "error": f"专利文档填充失败: {str(e)}"
            }
    
    def _apply_general_fill_changes(self, analysis_result: Dict[str, Any], 
                                  fill_data: Dict[str, Any], 
                                  document_content: str) -> Dict[str, Any]:
        """应用通用文档填充变化"""
        try:
            # 使用通用文档填充器
            fill_result = self.document_filler.fill_document(analysis_result, fill_data)
            
            if "error" in fill_result:
                return fill_result
            
            # 增强结果
            enhanced_result = {
                "success": True,
                "filled_content": fill_result.get("filled_content", document_content),
                "fill_summary": fill_result.get("fill_summary", {}),
                "document_type": "general"
            }
            
            return enhanced_result
            
        except Exception as e:
            return {
                "success": False,
                "error": f"通用文档填充失败: {str(e)}"
            }
    
    def _generate_final_document(self, fill_result: Dict[str, Any], 
                               analysis_result: Dict[str, Any], 
                               fill_data: Dict[str, Any]) -> Dict[str, Any]:
        """生成最终文档"""
        try:
            document_type = analysis_result.get("document_type", "general")
            filled_content = fill_result.get("filled_content", "")
            
            final_document = {
                "content": filled_content,
                "document_type": document_type,
                "metadata": {
                    "title": analysis_result.get("title", "未命名文档"),
                    "author": fill_data.get("author", "未知"),
                    "created_at": datetime.now().isoformat(),
                    "version": "1.0"
                },
                "format": "text"
            }
            
            # 根据文档类型添加特定格式
            if document_type == "patent":
                final_document["html_content"] = fill_result.get("html_output", "")
                final_document["format"] = "html"
            
            return final_document
            
        except Exception as e:
            return {
                "error": f"生成最终文档失败: {str(e)}"
            }
    
    def _generate_fill_application_report(self, fill_result: Dict[str, Any], 
                                        analysis_result: Dict[str, Any], 
                                        fill_data: Dict[str, Any]) -> Dict[str, Any]:
        """生成填充应用报告"""
        try:
            fields = analysis_result.get("fields", [])
            
            # 统计信息
            total_fields = len(fields)
            filled_fields = len([f for f in fill_data.values() if f])
            empty_fields = total_fields - filled_fields
            
            # 字段状态
            field_status = []
            for field in fields:
                field_name = field.get("name", "")
                field_value = fill_data.get(field_name, "")
                
                field_status.append({
                    "name": field_name,
                    "type": field.get("type", "text"),
                    "filled": bool(field_value),
                    "value_preview": field_value[:100] + "..." if len(field_value) > 100 else field_value,
                    "validation_status": "valid" if self._validate_field_type(field.get("type", "text"), field_value) else "invalid"
                })
            
            return {
                "summary": {
                    "total_fields": total_fields,
                    "filled_fields": filled_fields,
                    "empty_fields": empty_fields,
                    "fill_rate": filled_fields / total_fields if total_fields > 0 else 0
                },
                "field_status": field_status,
                "quality_metrics": self._calculate_fill_quality(fill_data, analysis_result),
                "recommendations": self._generate_fill_recommendations(fill_data, analysis_result)
            }
            
        except Exception as e:
            return {"error": f"生成填充应用报告失败: {str(e)}"}
    
    def _calculate_fill_quality(self, fill_data: Dict[str, Any], analysis_result: Dict[str, Any]) -> Dict[str, Any]:
        """计算填充质量"""
        try:
            fields = analysis_result.get("fields", [])
            
            # 完整性
            total_fields = len(fields)
            filled_fields = len([f for f in fill_data.values() if f])
            completeness = filled_fields / total_fields if total_fields > 0 else 0
            
            # 准确性
            valid_fields = 0
            for field in fields:
                field_name = field.get("name", "")
                field_value = fill_data.get(field_name, "")
                if field_value and self._validate_field_type(field.get("type", "text"), field_value):
                    valid_fields += 1
            
            accuracy = valid_fields / total_fields if total_fields > 0 else 0
            
            # 一致性
            consistency_score = self._calculate_consistency_score(fill_data, analysis_result)
            
            return {
                "completeness": completeness,
                "accuracy": accuracy,
                "consistency": consistency_score,
                "overall_score": (completeness + accuracy + consistency_score) / 3
            }
            
        except Exception as e:
            return {"error": f"质量计算失败: {str(e)}"}
    
    def _generate_fill_recommendations(self, fill_data: Dict[str, Any], analysis_result: Dict[str, Any]) -> List[str]:
        """生成填充建议"""
        recommendations = []
        
        # 检查必填字段
        required_fields = [f for f in analysis_result.get("fields", []) if f.get("required", False)]
        missing_required = [f.get("name") for f in required_fields if not fill_data.get(f.get("name", ""))]
        
        if missing_required:
            recommendations.append(f"需要填写必填字段: {', '.join(missing_required)}")
        
        # 检查字段质量
        for field in analysis_result.get("fields", []):
            field_name = field.get("name", "")
            field_value = fill_data.get(field_name, "")
            
            if field_value:
                if not self._validate_field_type(field.get("type", "text"), field_value):
                    recommendations.append(f"字段 '{field_name}' 的值格式不正确")
        
        if not recommendations:
            recommendations.append("填充质量良好，可以导出文档")
        
        return recommendations
    
    def _assess_final_quality(self, fill_result: Dict[str, Any], 
                            analysis_result: Dict[str, Any], 
                            fill_data: Dict[str, Any]) -> Dict[str, Any]:
        """评估最终质量"""
        try:
            quality_metrics = self._calculate_fill_quality(fill_data, analysis_result)
            
            overall_score = quality_metrics.get("overall_score", 0)
            
            if overall_score >= 0.9:
                quality_level = "excellent"
                assessment = "文档质量优秀，可以直接使用"
            elif overall_score >= 0.7:
                quality_level = "good"
                assessment = "文档质量良好，建议小幅调整"
            elif overall_score >= 0.5:
                quality_level = "fair"
                assessment = "文档质量一般，需要中等程度修改"
            else:
                quality_level = "poor"
                assessment = "文档质量较差，需要大量修改"
            
            return {
                "overall_score": overall_score,
                "quality_level": quality_level,
                "assessment": assessment,
                "metrics": quality_metrics
            }
            
        except Exception as e:
            return {"error": f"质量评估失败: {str(e)}"}
    
    def export_document(self, final_document: Dict[str, Any], 
                       export_format: str = "docx",
                       export_options: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        导出文档
        
        Args:
            final_document: 最终文档数据
            export_format: 导出格式 (docx, pdf, html, txt)
            export_options: 导出选项
            
        Returns:
            导出结果
        """
        try:
            # 1. 验证输入参数
            if not final_document or "error" in final_document:
                return {
                    "success": False,
                    "error": "最终文档数据无效"
                }
            
            if export_format not in ["docx", "pdf", "html", "txt"]:
                return {
                    "success": False,
                    "error": f"不支持的导出格式: {export_format}"
                }
            
            # 2. 准备导出选项
            if export_options is None:
                export_options = {}
            
            default_options = {
                "include_metadata": True,
                "include_watermark": False,
                "page_break": True,
                "font_size": 12,
                "line_spacing": 1.5
            }
            export_options = {**default_options, **export_options}
            
            # 3. 根据格式导出文档
            if export_format == "docx":
                export_result = self._export_to_docx(final_document, export_options)
            elif export_format == "pdf":
                export_result = self._export_to_pdf(final_document, export_options)
            elif export_format == "html":
                export_result = self._export_to_html(final_document, export_options)
            elif export_format == "txt":
                export_result = self._export_to_txt(final_document, export_options)
            
            # 4. 生成导出报告
            export_report = self._generate_export_report(final_document, export_format, export_options, export_result)
            
            return {
                "success": True,
                "export_format": export_format,
                "export_result": export_result,
                "export_report": export_report,
                "exported_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"导出文档失败: {str(e)}"
            }
    
    def _export_to_docx(self, final_document: Dict[str, Any], export_options: Dict[str, Any]) -> Dict[str, Any]:
        """导出为DOCX格式"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 创建文档
            doc = Document()
            
            # 设置文档属性
            metadata = final_document.get("metadata", {})
            if metadata.get("title"):
                doc.core_properties.title = metadata["title"]
            if metadata.get("author"):
                doc.core_properties.author = metadata["author"]
            
            # 添加标题
            title = metadata.get("title", "未命名文档")
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(title)
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加元数据
            if export_options.get("include_metadata", True):
                doc.add_paragraph(f"作者: {metadata.get('author', '未知')}")
                doc.add_paragraph(f"创建时间: {metadata.get('created_at', '未知')}")
                doc.add_paragraph(f"版本: {metadata.get('version', '1.0')}")
                doc.add_paragraph("")  # 空行
            
            # 添加内容
            content = final_document.get("content", "")
            if content:
                # 按段落分割
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        paragraph = doc.add_paragraph(para_text.strip())
                        paragraph.paragraph_format.line_spacing = export_options.get("line_spacing", 1.5)
            
            # 添加HTML内容（如果有）
            html_content = final_document.get("html_content", "")
            if html_content:
                doc.add_paragraph("HTML版本:")
                doc.add_paragraph(html_content)
            
            # 保存文档
            output_path = f"output/document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            
            return {
                "file_path": output_path,
                "file_size": os.path.getsize(output_path),
                "format": "docx"
            }
            
        except ImportError:
            return {
                "success": False,
                "error": "缺少python-docx库，无法导出DOCX格式"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"导出DOCX失败: {str(e)}"
            }
    
    def _export_to_pdf(self, final_document: Dict[str, Any], export_options: Dict[str, Any]) -> Dict[str, Any]:
        """导出为PDF格式"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            
            # 创建PDF文档
            output_path = f"output/document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # 创建自定义样式
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1  # 居中
            )
            
            content_style = ParagraphStyle(
                'CustomContent',
                parent=styles['Normal'],
                fontSize=export_options.get("font_size", 12),
                spaceAfter=12,
                leading=export_options.get("line_spacing", 1.5) * 12
            )
            
            # 构建内容
            story = []
            
            # 添加标题
            metadata = final_document.get("metadata", {})
            title = metadata.get("title", "未命名文档")
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 20))
            
            # 添加元数据
            if export_options.get("include_metadata", True):
                story.append(Paragraph(f"作者: {metadata.get('author', '未知')}", content_style))
                story.append(Paragraph(f"创建时间: {metadata.get('created_at', '未知')}", content_style))
                story.append(Paragraph(f"版本: {metadata.get('version', '1.0')}", content_style))
                story.append(Spacer(1, 20))
            
            # 添加内容
            content = final_document.get("content", "")
            if content:
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        story.append(Paragraph(para_text.strip(), content_style))
            
            # 构建PDF
            doc.build(story)
            
            return {
                "file_path": output_path,
                "file_size": os.path.getsize(output_path),
                "format": "pdf"
            }
            
        except ImportError:
            return {
                "success": False,
                "error": "缺少reportlab库，无法导出PDF格式"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"导出PDF失败: {str(e)}"
            }
    
    def _export_to_html(self, final_document: Dict[str, Any], export_options: Dict[str, Any]) -> Dict[str, Any]:
        """导出为HTML格式"""
        try:
            metadata = final_document.get("metadata", {})
            content = final_document.get("content", "")
            html_content = final_document.get("html_content", "")
            
            # 生成HTML
            html_template = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{metadata.get('title', '未命名文档')}</title>
    <style>
        body {{
            font-family: 'Microsoft YaHei', Arial, sans-serif;
            line-height: {export_options.get('line_spacing', 1.5)};
            font-size: {export_options.get('font_size', 12)}px;
            margin: 40px;
            color: #333;
        }}
        .title {{
            text-align: center;
            font-size: 24px;
            font-weight: bold;
            margin-bottom: 30px;
            color: #2c3e50;
        }}
        .metadata {{
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 30px;
            border-left: 4px solid #007bff;
        }}
        .content {{
            text-align: justify;
        }}
        .paragraph {{
            margin-bottom: 15px;
        }}
        .watermark {{
            position: fixed;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%) rotate(-45deg);
            font-size: 48px;
            color: rgba(0,0,0,0.1);
            z-index: -1;
        }}
    </style>
</head>
<body>
"""
            
            # 添加水印
            if export_options.get("include_watermark", False):
                html_template += '<div class="watermark">DRAFT</div>'
            
            # 添加标题
            html_template += f'<div class="title">{metadata.get("title", "未命名文档")}</div>'
            
            # 添加元数据
            if export_options.get("include_metadata", True):
                html_template += f'''
<div class="metadata">
    <p><strong>作者:</strong> {metadata.get('author', '未知')}</p>
    <p><strong>创建时间:</strong> {metadata.get('created_at', '未知')}</p>
    <p><strong>版本:</strong> {metadata.get('version', '1.0')}</p>
</div>
'''
            
            # 添加内容
            html_template += '<div class="content">'
            if html_content:
                html_template += html_content
            else:
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        html_template += f'<div class="paragraph">{para_text.strip()}</div>'
            
            html_template += '''
</div>
</body>
</html>
'''
            
            # 保存HTML文件
            output_path = f"output/document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_template)
            
            return {
                "file_path": output_path,
                "file_size": os.path.getsize(output_path),
                "format": "html"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"导出HTML失败: {str(e)}"
            }
    
    def _export_to_txt(self, final_document: Dict[str, Any], export_options: Dict[str, Any]) -> Dict[str, Any]:
        """导出为TXT格式"""
        try:
            metadata = final_document.get("metadata", {})
            content = final_document.get("content", "")
            
            # 生成文本内容
            text_content = []
            
            # 添加标题
            text_content.append(metadata.get("title", "未命名文档"))
            text_content.append("=" * 50)
            text_content.append("")
            
            # 添加元数据
            if export_options.get("include_metadata", True):
                text_content.append(f"作者: {metadata.get('author', '未知')}")
                text_content.append(f"创建时间: {metadata.get('created_at', '未知')}")
                text_content.append(f"版本: {metadata.get('version', '1.0')}")
                text_content.append("")
            
            # 添加内容
            text_content.append(content)
            
            # 保存文本文件
            output_path = f"output/document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(text_content))
            
            return {
                "file_path": output_path,
                "file_size": os.path.getsize(output_path),
                "format": "txt"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"导出TXT失败: {str(e)}"
            }
    
    def _generate_export_report(self, final_document: Dict[str, Any], 
                              export_format: str, 
                              export_options: Dict[str, Any], 
                              export_result: Dict[str, Any]) -> Dict[str, Any]:
        """生成导出报告"""
        try:
            metadata = final_document.get("metadata", {})
            content = final_document.get("content", "")
            
            report = {
                "export_summary": {
                    "format": export_format,
                    "file_path": export_result.get("file_path", ""),
                    "file_size": export_result.get("file_size", 0),
                    "export_time": datetime.now().isoformat()
                },
                "document_info": {
                    "title": metadata.get("title", "未命名文档"),
                    "author": metadata.get("author", "未知"),
                    "content_length": len(content),
                    "word_count": len(content.split()),
                    "paragraph_count": len([p for p in content.split('\n\n') if p.strip()])
                },
                "export_options": export_options,
                "quality_metrics": {
                    "completeness": 1.0 if content else 0.0,
                    "formatting": 1.0 if export_format in ["docx", "pdf", "html"] else 0.8,
                    "accessibility": 1.0 if export_format in ["txt", "html"] else 0.9
                }
            }
            
            return report
            
        except Exception as e:
            return {"error": f"生成导出报告失败: {str(e)}"}