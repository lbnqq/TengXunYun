import PyPDF2
import re
import json
from typing import Dict, List, Any, Optional, Tuple
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .base_tool import BaseTool

class EnhancedDocumentParserTool(BaseTool):
    """
    Enhanced document parser with deep structural analysis capabilities.
    Features:
    - Deep structural analysis (paragraphs, headings, lists)
    - Key information extraction (entities, keywords, summaries)
    - Data table recognition and parsing
    - Document structure tree generation
    - Metadata analysis
    """

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.heading_patterns = [
            r'^#{1,6}\s+(.+)$',  # Markdown headings
            r'^(\d+\.?\s+.+)$',  # Numbered headings
            r'^([A-Z][A-Z\s]+)$',  # ALL CAPS headings
            r'^(.+)[:：]\s*$',   # Colon-ended headings
        ]
        self.list_patterns = [
            r'^\s*[-*+]\s+(.+)$',  # Bullet lists
            r'^\s*\d+[.)]\s+(.+)$',  # Numbered lists
            r'^\s*[a-zA-Z][.)]\s+(.+)$',  # Lettered lists
        ]
        self.table_patterns = [
            r'\|.*\|',  # Markdown tables
            r'\t.*\t',  # Tab-separated
        ]

    def execute(self, file_path: str, file_type: str = None, analysis_depth: str = "deep") -> dict:
        """
        Execute enhanced document parsing with configurable analysis depth.

        Args:
            file_path: Path to the document file
            file_type: Document type (auto-detected if None)
            analysis_depth: "basic", "standard", or "deep"
        """
        if file_type is None:
            file_type = self._detect_file_type(file_path)
            if not file_type:
                return {"error": "Unsupported file type or cannot infer from path."}

        try:
            # Basic content extraction
            extraction_result = self._extract_content(file_path, file_type)
            if "error" in extraction_result:
                return extraction_result

            text_content = extraction_result["text_content"]
            basic_structure = extraction_result["structure_info"]

            # Enhanced analysis based on depth
            if analysis_depth in ["standard", "deep"]:
                structural_analysis = self._analyze_document_structure(text_content)
                key_info = self._extract_key_information(text_content)

                result = {
                    "text_content": text_content,
                    "basic_structure": basic_structure,
                    "structural_analysis": structural_analysis,
                    "key_information": key_info,
                    "analysis_depth": analysis_depth
                }

                if analysis_depth == "deep":
                    deep_analysis = self._perform_deep_analysis(text_content, structural_analysis)
                    result.update(deep_analysis)

                return result
            else:
                return {
                    "text_content": text_content,
                    "structure_info": basic_structure,
                    "analysis_depth": "basic"
                }

        except FileNotFoundError:
            return {"error": f"File not found: {file_path}"}
        except Exception as e:
            return {"error": f"Error parsing document {file_path}: {e}"}

    def _detect_file_type(self, file_path: str) -> Optional[str]:
        """Detect file type from extension."""
        if file_path.lower().endswith(".docx"):
            return "docx"
        elif file_path.lower().endswith(".pdf"):
            return "pdf"
        elif file_path.lower().endswith(".txt"):
            return "txt"
        elif file_path.lower().endswith(".md"):
            return "markdown"
        return None

    def _extract_content(self, file_path: str, file_type: str) -> dict:
        """Extract basic content from document."""
        text_content = ""
        structure_info = {}

        if file_type == "docx":
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal",
                        "alignment": str(para.alignment) if para.alignment else "LEFT"
                    })
                    text_content += para.text + "\n"

            structure_info = {
                "paragraphs": len(paragraphs),
                "paragraph_details": paragraphs,
                "has_tables": len(doc.tables) > 0,
                "table_count": len(doc.tables)
            }

        elif file_type == "pdf":
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                pages_content = []
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    page_text = page.extract_text()
                    pages_content.append(page_text)
                    text_content += page_text + "\n"

                structure_info = {
                    "pages": len(reader.pages),
                    "pages_content": pages_content
                }

        elif file_type in ["txt", "markdown"]:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()

            lines = text_content.split('\n')
            structure_info = {
                "lines": len(lines),
                "non_empty_lines": len([line for line in lines if line.strip()]),
                "file_type": file_type
            }
        else:
            return {"error": f"Unsupported file type: {file_type}"}

        return {"text_content": text_content.strip(), "structure_info": structure_info}

    def _analyze_document_structure(self, text_content: str) -> dict:
        """Analyze document structure including headings, lists, and paragraphs."""
        lines = text_content.split('\n')
        structure = {
            "headings": [],
            "lists": [],
            "paragraphs": [],
            "tables": [],
            "document_tree": []
        }

        current_section = None
        paragraph_buffer = []

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                if paragraph_buffer:
                    structure["paragraphs"].append({
                        "content": " ".join(paragraph_buffer),
                        "line_start": i - len(paragraph_buffer),
                        "line_end": i,
                        "section": current_section
                    })
                    paragraph_buffer = []
                continue

            # Check for headings
            heading_match = self._match_heading(line)
            if heading_match:
                if paragraph_buffer:
                    structure["paragraphs"].append({
                        "content": " ".join(paragraph_buffer),
                        "line_start": i - len(paragraph_buffer),
                        "line_end": i,
                        "section": current_section
                    })
                    paragraph_buffer = []

                heading_info = {
                    "text": heading_match["text"],
                    "level": heading_match["level"],
                    "line_number": i,
                    "type": heading_match["type"]
                }
                structure["headings"].append(heading_info)
                current_section = heading_match["text"]
                continue

            # Check for lists
            list_match = self._match_list_item(line)
            if list_match:
                if paragraph_buffer:
                    structure["paragraphs"].append({
                        "content": " ".join(paragraph_buffer),
                        "line_start": i - len(paragraph_buffer),
                        "line_end": i,
                        "section": current_section
                    })
                    paragraph_buffer = []

                structure["lists"].append({
                    "content": list_match["content"],
                    "type": list_match["type"],
                    "line_number": i,
                    "section": current_section
                })
                continue

            # Check for tables
            if self._is_table_row(line):
                structure["tables"].append({
                    "content": line,
                    "line_number": i,
                    "section": current_section
                })
                continue

            # Regular paragraph content
            paragraph_buffer.append(line)

        # Handle remaining paragraph buffer
        if paragraph_buffer:
            structure["paragraphs"].append({
                "content": " ".join(paragraph_buffer),
                "line_start": len(lines) - len(paragraph_buffer),
                "line_end": len(lines),
                "section": current_section
            })

        # Generate document tree
        structure["document_tree"] = self._generate_document_tree(structure["headings"])

        return structure

    def _match_heading(self, line: str) -> Optional[dict]:
        """Match line against heading patterns."""
        # Markdown style headings
        markdown_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if markdown_match:
            return {
                "text": markdown_match.group(2),
                "level": len(markdown_match.group(1)),
                "type": "markdown"
            }

        # Numbered headings
        numbered_match = re.match(r'^(\d+\.?\s+)(.+)$', line)
        if numbered_match and len(line) < 100:  # Likely a heading, not a long sentence
            return {
                "text": numbered_match.group(2),
                "level": 1,
                "type": "numbered"
            }

        # ALL CAPS headings (short lines only)
        if line.isupper() and len(line) < 50 and len(line.split()) <= 8:
            return {
                "text": line,
                "level": 1,
                "type": "caps"
            }

        # Colon-ended headings
        colon_match = re.match(r'^(.+)[:：]\s*$', line)
        if colon_match and len(line) < 80:
            return {
                "text": colon_match.group(1),
                "level": 2,
                "type": "colon"
            }

        return None

    def _match_list_item(self, line: str) -> Optional[dict]:
        """Match line against list item patterns."""
        # Bullet lists
        bullet_match = re.match(r'^\s*[-*+]\s+(.+)$', line)
        if bullet_match:
            return {
                "content": bullet_match.group(1),
                "type": "bullet"
            }

        # Numbered lists
        numbered_match = re.match(r'^\s*(\d+)[.)]\s+(.+)$', line)
        if numbered_match:
            return {
                "content": numbered_match.group(2),
                "type": "numbered",
                "number": numbered_match.group(1)
            }

        # Lettered lists
        lettered_match = re.match(r'^\s*([a-zA-Z])[.)]\s+(.+)$', line)
        if lettered_match:
            return {
                "content": lettered_match.group(2),
                "type": "lettered",
                "letter": lettered_match.group(1)
            }

        return None

    def _is_table_row(self, line: str) -> bool:
        """Check if line appears to be part of a table."""
        # Markdown table format
        if '|' in line and line.count('|') >= 2:
            return True

        # Tab-separated format
        if '\t' in line and line.count('\t') >= 2:
            return True

        # Multiple spaces (potential column alignment)
        if re.search(r'\s{3,}', line) and len(line.split()) >= 3:
            return True

        return False

    def _generate_document_tree(self, headings: List[dict]) -> List[dict]:
        """Generate hierarchical document tree from headings."""
        if not headings:
            return []

        tree = []
        stack = []

        for heading in headings:
            node = {
                "text": heading["text"],
                "level": heading["level"],
                "line_number": heading["line_number"],
                "children": []
            }

            # Find the correct parent level
            while stack and stack[-1]["level"] >= heading["level"]:
                stack.pop()

            if stack:
                stack[-1]["children"].append(node)
            else:
                tree.append(node)

            stack.append(node)

        return tree

    def _extract_key_information(self, text_content: str) -> dict:
        """Extract key information including entities, keywords, and summary."""
        # Basic keyword extraction using frequency analysis
        words = re.findall(r'\b[a-zA-Z\u4e00-\u9fff]{3,}\b', text_content.lower())
        word_freq = {}
        for word in words:
            if len(word) > 2 and word not in ['the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 'her', 'was', 'one', 'our', 'out', 'day', 'get', 'has', 'him', 'his', 'how', 'its', 'may', 'new', 'now', 'old', 'see', 'two', 'who', 'boy', 'did', 'man', 'men', 'put', 'say', 'she', 'too', 'use']:
                word_freq[word] = word_freq.get(word, 0) + 1

        # Get top keywords
        top_keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:20]

        # Extract potential entities (capitalized words/phrases)
        entities = re.findall(r'\b[A-Z][a-zA-Z\u4e00-\u9fff]*(?:\s+[A-Z][a-zA-Z\u4e00-\u9fff]*)*\b', text_content)
        entity_freq = {}
        for entity in entities:
            if len(entity) > 2:
                entity_freq[entity] = entity_freq.get(entity, 0) + 1

        top_entities = sorted(entity_freq.items(), key=lambda x: x[1], reverse=True)[:15]

        # Extract numbers and dates
        numbers = re.findall(r'\b\d+(?:\.\d+)?%?\b', text_content)
        dates = re.findall(r'\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b|\b\d{1,2}[-/]\d{1,2}[-/]\d{4}\b', text_content)

        # Generate basic summary (first and last sentences of each paragraph)
        paragraphs = [p.strip() for p in text_content.split('\n\n') if p.strip()]
        summary_sentences = []
        for para in paragraphs[:3]:  # First 3 paragraphs
            sentences = re.split(r'[.!?。！？]', para)
            if sentences:
                summary_sentences.append(sentences[0].strip())

        return {
            "keywords": [{"word": word, "frequency": freq} for word, freq in top_keywords],
            "entities": [{"entity": entity, "frequency": freq} for entity, freq in top_entities],
            "numbers": list(set(numbers)),
            "dates": list(set(dates)),
            "summary_sentences": summary_sentences,
            "statistics": {
                "total_words": len(words),
                "unique_words": len(set(words)),
                "total_characters": len(text_content),
                "total_paragraphs": len(paragraphs)
            }
        }

    def _perform_deep_analysis(self, text_content: str, structural_analysis: dict) -> dict:
        """Perform deep analysis including content patterns and document characteristics."""

        # Analyze content patterns
        content_patterns = self._analyze_content_patterns(text_content)

        # Analyze document characteristics
        doc_characteristics = self._analyze_document_characteristics(text_content, structural_analysis)

        # Extract semantic sections
        semantic_sections = self._extract_semantic_sections(text_content, structural_analysis)

        # Analyze writing style
        writing_style = self._analyze_writing_style(text_content)

        return {
            "content_patterns": content_patterns,
            "document_characteristics": doc_characteristics,
            "semantic_sections": semantic_sections,
            "writing_style": writing_style
        }

    def _analyze_content_patterns(self, text_content: str) -> dict:
        """Analyze patterns in content structure and organization."""
        patterns = {
            "has_introduction": False,
            "has_conclusion": False,
            "has_methodology": False,
            "has_results": False,
            "has_references": False,
            "question_count": len(re.findall(r'[?？]', text_content)),
            "exclamation_count": len(re.findall(r'[!！]', text_content)),
            "bullet_points": len(re.findall(r'^\s*[-*+]\s+', text_content, re.MULTILINE)),
            "numbered_items": len(re.findall(r'^\s*\d+[.)]\s+', text_content, re.MULTILINE))
        }

        # Check for common section indicators
        text_lower = text_content.lower()
        intro_indicators = ['introduction', 'overview', '引言', '概述', '背景']
        conclusion_indicators = ['conclusion', 'summary', '结论', '总结', '小结']
        method_indicators = ['methodology', 'method', 'approach', '方法', '方案']
        result_indicators = ['results', 'findings', '结果', '发现', '成果']
        ref_indicators = ['references', 'bibliography', '参考文献', '引用']

        patterns["has_introduction"] = any(indicator in text_lower for indicator in intro_indicators)
        patterns["has_conclusion"] = any(indicator in text_lower for indicator in conclusion_indicators)
        patterns["has_methodology"] = any(indicator in text_lower for indicator in method_indicators)
        patterns["has_results"] = any(indicator in text_lower for indicator in result_indicators)
        patterns["has_references"] = any(indicator in text_lower for indicator in ref_indicators)

        return patterns

    def _analyze_document_characteristics(self, text_content: str, structural_analysis: dict) -> dict:
        """Analyze overall document characteristics."""
        characteristics = {
            "document_length": "short",  # short, medium, long
            "structure_complexity": "simple",  # simple, moderate, complex
            "content_density": "low",  # low, medium, high
            "formality_level": "informal",  # informal, semi-formal, formal
            "technical_level": "basic"  # basic, intermediate, advanced
        }

        word_count = len(text_content.split())

        # Document length
        if word_count < 500:
            characteristics["document_length"] = "short"
        elif word_count < 2000:
            characteristics["document_length"] = "medium"
        else:
            characteristics["document_length"] = "long"

        # Structure complexity
        heading_count = len(structural_analysis.get("headings", []))
        if heading_count < 3:
            characteristics["structure_complexity"] = "simple"
        elif heading_count < 8:
            characteristics["structure_complexity"] = "moderate"
        else:
            characteristics["structure_complexity"] = "complex"

        # Content density (average words per paragraph)
        paragraph_count = len(structural_analysis.get("paragraphs", []))
        if paragraph_count > 0:
            avg_words_per_para = word_count / paragraph_count
            if avg_words_per_para < 30:
                characteristics["content_density"] = "low"
            elif avg_words_per_para < 80:
                characteristics["content_density"] = "medium"
            else:
                characteristics["content_density"] = "high"

        # Technical level (based on technical terms and jargon)
        technical_indicators = ['algorithm', 'implementation', 'framework', 'architecture', 'methodology', 'analysis', 'optimization', 'performance', 'system', 'technical', '算法', '实现', '框架', '架构', '方法论', '分析', '优化', '性能', '系统', '技术']
        technical_count = sum(1 for indicator in technical_indicators if indicator in text_content.lower())

        if technical_count < 3:
            characteristics["technical_level"] = "basic"
        elif technical_count < 10:
            characteristics["technical_level"] = "intermediate"
        else:
            characteristics["technical_level"] = "advanced"

        return characteristics

    def _extract_semantic_sections(self, text_content: str, structural_analysis: dict) -> List[dict]:
        """Extract semantic sections based on content and structure."""
        sections = []
        paragraphs = structural_analysis.get("paragraphs", [])
        headings = structural_analysis.get("headings", [])

        current_section = {
            "title": "Introduction",
            "content": "",
            "type": "introduction",
            "start_line": 0,
            "end_line": 0
        }

        for i, paragraph in enumerate(paragraphs):
            # Check if this paragraph starts a new section
            section_type = self._classify_paragraph_type(paragraph["content"])

            if section_type != current_section["type"] and current_section["content"]:
                # Save current section and start new one
                sections.append(current_section.copy())
                current_section = {
                    "title": section_type.title(),
                    "content": paragraph["content"],
                    "type": section_type,
                    "start_line": paragraph.get("line_start", i),
                    "end_line": paragraph.get("line_end", i)
                }
            else:
                # Continue current section
                current_section["content"] += " " + paragraph["content"]
                current_section["end_line"] = paragraph.get("line_end", i)

        # Add the last section
        if current_section["content"]:
            sections.append(current_section)

        return sections

    def _classify_paragraph_type(self, content: str) -> str:
        """Classify paragraph type based on content patterns."""
        content_lower = content.lower()

        # Introduction indicators
        if any(word in content_lower for word in ['introduction', 'overview', 'background', '引言', '概述', '背景']):
            return "introduction"

        # Methodology indicators
        if any(word in content_lower for word in ['method', 'approach', 'procedure', '方法', '方案', '步骤']):
            return "methodology"

        # Results indicators
        if any(word in content_lower for word in ['result', 'finding', 'outcome', '结果', '发现', '成果']):
            return "results"

        # Discussion indicators
        if any(word in content_lower for word in ['discussion', 'analysis', 'implication', '讨论', '分析', '影响']):
            return "discussion"

        # Conclusion indicators
        if any(word in content_lower for word in ['conclusion', 'summary', 'final', '结论', '总结', '最终']):
            return "conclusion"

        return "content"

    def _analyze_writing_style(self, text_content: str) -> dict:
        """Analyze writing style characteristics."""
        sentences = re.split(r'[.!?。！？]', text_content)
        sentences = [s.strip() for s in sentences if s.strip()]

        if not sentences:
            return {"error": "No sentences found for analysis"}

        # Calculate sentence length statistics
        sentence_lengths = [len(s.split()) for s in sentences]
        avg_sentence_length = sum(sentence_lengths) / len(sentence_lengths)

        # Analyze vocabulary complexity
        words = re.findall(r'\b[a-zA-Z\u4e00-\u9fff]+\b', text_content.lower())
        long_words = [w for w in words if len(w) > 6]
        vocabulary_complexity = len(long_words) / len(words) if words else 0

        # Analyze sentence structure variety
        question_ratio = len([s for s in sentences if '?' in s or '？' in s]) / len(sentences)
        exclamation_ratio = len([s for s in sentences if '!' in s or '！' in s]) / len(sentences)

        # Determine writing style characteristics
        style = {
            "average_sentence_length": round(avg_sentence_length, 2),
            "vocabulary_complexity": round(vocabulary_complexity, 3),
            "question_ratio": round(question_ratio, 3),
            "exclamation_ratio": round(exclamation_ratio, 3),
            "total_sentences": len(sentences),
            "style_indicators": {
                "concise": avg_sentence_length < 15,
                "complex": vocabulary_complexity > 0.2,
                "interactive": question_ratio > 0.1,
                "emphatic": exclamation_ratio > 0.05
            }
        }

        return style


# Maintain backward compatibility
class DocumentParserTool(EnhancedDocumentParserTool):
    """Backward compatibility wrapper for the enhanced document parser."""

    def execute(self, file_path: str, file_type: str = None) -> dict:
        """Execute with basic analysis for backward compatibility."""
        result = super().execute(file_path, file_type, analysis_depth="basic")

        # Convert to old format for compatibility
        if "structure_info" in result:
            return {
                "text_content": result["text_content"],
                "structure_info": result["structure_info"]
            }

        return result