#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Writing Style Analyzer - 核心模块

Author: AI Assistant (Claude)
Created: 2025-01-28
Last Modified: 2025-01-28
Modified By: AI Assistant (Claude)
AI Assisted: 是 - Claude 3.5 Sonnet
Version: v1.0
License: MIT
"""

import re
import json
import os
from typing import Dict, Any, List, Tuple, Optional
from datetime import datetime
import hashlib
import difflib

# 导入增强的文风分析组件
try:
    from .comprehensive_style_processor import ComprehensiveStyleProcessor
    from .enhanced_style_extractor import EnhancedStyleExtractor
    from .style_alignment_engine import StyleAlignmentEngine
    ENHANCED_FEATURES_AVAILABLE = True
except ImportError:
    ENHANCED_FEATURES_AVAILABLE = False
    print("Warning: Enhanced style analysis features not available. Using basic functionality.")

class WritingStyleAnalyzer:

    def __init__(self, storage_path: str = "src/core/knowledge_base/writing_style_templates", llm_client=None):
        self.tool_name = "文风分析器"
        self.description = "分析文档写作风格，生成文风模板，支持文风对齐和润色功能"
        self.storage_path = storage_path
        self.llm_client = llm_client

        # 确保存储目录存在
        os.makedirs(storage_path, exist_ok=True)

        # 初始化增强功能组件
        if ENHANCED_FEATURES_AVAILABLE and llm_client:
            self.enhanced_processor = ComprehensiveStyleProcessor(
                llm_client=llm_client,
                storage_path=os.path.join(storage_path, "enhanced_analysis")
            )
            self.use_enhanced_features = True
            print("✅ 增强文风分析功能已启用")
        else:
            self.enhanced_processor = None
            self.use_enhanced_features = False
            print("⚠️ 使用基础文风分析功能")
        
        # 文风特征分析维度
        self.style_dimensions = {
            "sentence_structure": {
                "name": "句式结构",
                "features": ["平均句长", "长短句比例", "复合句使用", "并列句使用"]
            },
            "vocabulary_choice": {
                "name": "词汇选择", 
                "features": ["正式程度", "专业术语", "修饰词使用", "动词类型偏好"]
            },
            "expression_style": {
                "name": "表达方式",
                "features": ["主被动语态", "人称使用", "语气强度", "情感色彩"]
            },
            "text_organization": {
                "name": "文本组织",
                "features": ["段落结构", "逻辑连接", "过渡方式", "总结习惯"]
            },
            "language_habits": {
                "name": "语言习惯",
                "features": ["口语化程度", "书面语规范", "地域特色", "行业特色"]
            }
        }
        
        # 常见文风类型
        self.style_types = {
            "formal_official": {
                "name": "正式公文风格",
                "characteristics": ["严谨规范", "用词准确", "逻辑清晰", "格式标准"],
                "typical_patterns": ["根据", "按照", "现将", "特此", "务必"]
            },
            "business_professional": {
                "name": "商务专业风格", 
                "characteristics": ["简洁明了", "重点突出", "数据导向", "结果导向"],
                "typical_patterns": ["提升", "优化", "实现", "达成", "推进"]
            },
            "academic_research": {
                "name": "学术研究风格",
                "characteristics": ["客观严谨", "逻辑严密", "论证充分", "引用规范"],
                "typical_patterns": ["研究表明", "分析发现", "综合考虑", "深入探讨"]
            },
            "narrative_descriptive": {
                "name": "叙述描述风格",
                "characteristics": ["生动形象", "细节丰富", "情感真实", "故事性强"],
                "typical_patterns": ["生动地", "详细地", "深刻地", "真实地"]
            },
            "concise_practical": {
                "name": "简洁实用风格",
                "characteristics": ["言简意赅", "直接有效", "操作性强", "易于理解"],
                "typical_patterns": ["直接", "立即", "马上", "简单", "快速"]
            }
        }
    
        # 初始化风格行为缓存/导出目录
        self.semantic_behavior_dir = os.path.join(self.storage_path, "semantic_behavior")
        os.makedirs(os.path.join(self.semantic_behavior_dir, "profiles"), exist_ok=True)

    def analyze_writing_style(self, document_content: str, document_name: str = None,
                             use_enhanced: bool = None) -> Dict[str, Any]:
        """
        分析文档的写作风格

        Args:
            document_content: 文档内容
            document_name: 文档名称
            use_enhanced: 是否使用增强功能（None时自动判断）

        Returns:
            文风分析结果
        """
        # 决定是否使用增强功能
        if use_enhanced is None:
            use_enhanced = self.use_enhanced_features

        if use_enhanced and self.enhanced_processor:
            return self._analyze_with_enhanced_features(document_content, document_name)
        else:
            return self._analyze_with_basic_features(document_content, document_name)

    def _analyze_with_enhanced_features(self, document_content: str, document_name: str = None) -> Dict[str, Any]:
        """使用增强功能进行文风分析"""
        try:
            print(f"🔍 使用增强功能分析文档: {document_name or '未命名文档'}")

            # 使用综合文风处理器进行分析
            enhanced_result = self.enhanced_processor.extract_comprehensive_style_features(
                document_content, document_name, include_advanced_analysis=True
            )

            # 转换为兼容的格式
            analysis_result = {
                "document_name": document_name or "未命名文档",
                "analysis_time": datetime.now().isoformat(),
                "analysis_method": "enhanced",
                "document_stats": self._get_document_statistics(document_content),
                "enhanced_analysis": enhanced_result,
                "style_features": self._extract_style_features_from_enhanced(enhanced_result),
                "style_type": self._determine_style_type_from_enhanced(enhanced_result),
                "confidence_score": self._calculate_confidence_from_enhanced(enhanced_result),
                "style_prompt": self._generate_style_prompt_from_enhanced(enhanced_result),
                "template_id": self._generate_template_id(document_name, {}),
                "detailed_analysis": enhanced_result.get("advanced_features", {}),
                "writing_recommendations": self._generate_recommendations_from_enhanced(enhanced_result),
                "style_comparison": {}
            }

            print("✅ 增强文风分析完成")
            return analysis_result

        except Exception as e:
            print(f"❌ 增强分析失败，回退到基础分析: {str(e)}")
            return self._analyze_with_basic_features(document_content, document_name)

    def _analyze_with_basic_features(self, document_content: str, document_name: str = None) -> Dict[str, Any]:
        """使用基础功能进行文风分析"""
        try:
            analysis_result = {
                "document_name": document_name or "未命名文档",
                "analysis_time": datetime.now().isoformat(),
                "analysis_method": "basic",
                "document_stats": self._get_document_statistics(document_content),
                "style_features": {},
                "style_type": None,
                "confidence_score": 0.0,
                "style_prompt": "",
                "template_id": None,
                "detailed_analysis": {},
                "writing_recommendations": [],
                "style_comparison": {}
            }

            # 分析各个维度的文风特征
            style_features = self._analyze_style_features(document_content)
            analysis_result["style_features"] = style_features

            # 识别文风类型
            style_type, confidence = self._identify_style_type(document_content, style_features)
            analysis_result["style_type"] = style_type
            analysis_result["confidence_score"] = confidence

            # 生成详细分析报告
            detailed_analysis = self._generate_detailed_analysis(document_content, style_features, style_type)
            analysis_result["detailed_analysis"] = detailed_analysis

            # 生成写作建议
            recommendations = self._generate_writing_recommendations(style_features, style_type)
            analysis_result["writing_recommendations"] = recommendations

            # 生成风格对比
            style_comparison = self._generate_style_comparison(style_features)
            analysis_result["style_comparison"] = style_comparison

            # 生成文风提示词
            style_prompt = self._generate_enhanced_style_prompt(style_features, style_type, detailed_analysis)
            analysis_result["style_prompt"] = style_prompt

            # 生成模板ID
            template_id = self._generate_template_id(document_name, style_features)
            analysis_result["template_id"] = template_id

            return analysis_result

        except Exception as e:
            return {"error": f"文风分析失败: {str(e)}"}

    def _extract_style_features_from_enhanced(self, enhanced_result: Dict[str, Any]) -> Dict[str, Any]:
        """从增强分析结果中提取风格特征"""
        style_features = {}

        try:
            basic_features = enhanced_result.get("basic_features", {})

            # 提取量化特征
            quant_features = basic_features.get("quantitative_features", {})
            if quant_features:
                lexical = quant_features.get("lexical_features", {})
                syntactic = quant_features.get("syntactic_features", {})

                style_features["lexical_richness"] = lexical.get("ttr", 0)
                style_features["avg_word_length"] = lexical.get("avg_word_length", 0)
                style_features["formal_density"] = lexical.get("formal_word_density", 0)
                style_features["avg_sentence_length"] = syntactic.get("avg_sentence_length", 0)
                style_features["sentence_variety"] = syntactic.get("sentence_length_std", 0)

            # 提取LLM特征
            llm_features = basic_features.get("llm_features", {})
            if llm_features:
                evaluations = llm_features.get("evaluations", {})
                for dimension, eval_data in evaluations.items():
                    if isinstance(eval_data, dict) and "score" in eval_data:
                        style_features[f"llm_{dimension}"] = eval_data["score"]

        except Exception as e:
            style_features["extraction_error"] = str(e)

        return style_features

    def _determine_style_type_from_enhanced(self, enhanced_result: Dict[str, Any]) -> str:
        """从增强分析结果中确定风格类型"""
        try:
            advanced_features = enhanced_result.get("advanced_features", {})
            comprehensive_analysis = advanced_features.get("comprehensive_analysis", {})

            if comprehensive_analysis.get("success"):
                parsed_analysis = comprehensive_analysis.get("parsed_analysis", {})
                overall_style = parsed_analysis.get("overall_style", {})

                # 从LLM分析中提取风格类型
                style_type = overall_style.get("主要风格类型", "")
                if style_type:
                    # 映射到内部风格类型
                    style_mapping = {
                        "正式公文": "formal_official",
                        "商务专业": "business_professional",
                        "学术研究": "academic_research",
                        "叙述描述": "narrative_descriptive",
                        "简洁实用": "concise_practical"
                    }
                    return style_mapping.get(style_type, "business_professional")

            # 回退到基于特征的判断
            style_features = self._extract_style_features_from_enhanced(enhanced_result)
            formal_score = style_features.get("llm_正式程度", 3.0)

            if formal_score >= 4.0:
                return "formal_official"
            elif formal_score >= 3.5:
                return "business_professional"
            else:
                return "concise_practical"

        except Exception:
            return "business_professional"

    def _calculate_confidence_from_enhanced(self, enhanced_result: Dict[str, Any]) -> float:
        """从增强分析结果中计算置信度"""
        try:
            basic_features = enhanced_result.get("basic_features", {})

            # 基于成功提取的特征数量计算置信度
            feature_vector = basic_features.get("feature_vector", [])
            if feature_vector:
                base_confidence = min(len(feature_vector) / 20.0, 1.0)  # 假设20个特征为满分
            else:
                base_confidence = 0.5

            # 如果有LLM分析，提高置信度
            llm_features = basic_features.get("llm_features", {})
            if llm_features.get("evaluations"):
                base_confidence = min(base_confidence + 0.2, 1.0)

            # 如果有高级分析，进一步提高置信度
            advanced_features = enhanced_result.get("advanced_features", {})
            if advanced_features:
                base_confidence = min(base_confidence + 0.1, 1.0)

            return round(base_confidence, 3)

        except Exception:
            return 0.7  # 默认置信度

    def _generate_style_prompt_from_enhanced(self, enhanced_result: Dict[str, Any]) -> str:
        """从增强分析结果中生成风格提示词"""
        try:
            style_features = self._extract_style_features_from_enhanced(enhanced_result)
            style_type = self._determine_style_type_from_enhanced(enhanced_result)

            # 获取风格类型信息
            style_info = self.style_types.get(style_type, {})
            style_name = style_info.get("name", "标准风格")
            characteristics = style_info.get("characteristics", [])

            # 构建提示词
            prompt_parts = [f"请按照{style_name}进行写作"]

            if characteristics:
                prompt_parts.append(f"特点：{', '.join(characteristics)}")

            # 添加具体的风格指导
            if style_features.get("formal_density", 0) > 10:
                prompt_parts.append("使用正式词汇和表达")

            if style_features.get("avg_sentence_length", 0) > 15:
                prompt_parts.append("采用较长的复合句结构")
            elif style_features.get("avg_sentence_length", 0) < 10:
                prompt_parts.append("使用简洁明了的短句")

            return "；".join(prompt_parts)

        except Exception:
            return "请保持原有的写作风格"

    def _generate_recommendations_from_enhanced(self, enhanced_result: Dict[str, Any]) -> List[str]:
        """从增强分析结果中生成写作建议"""
        recommendations = []

        try:
            style_features = self._extract_style_features_from_enhanced(enhanced_result)

            # 基于特征给出建议
            if style_features.get("lexical_richness", 0) < 0.5:
                recommendations.append("建议增加词汇多样性，避免重复使用相同词汇")

            if style_features.get("sentence_variety", 0) < 3:
                recommendations.append("建议增加句式变化，使用长短句结合的方式")

            if style_features.get("formal_density", 0) < 5:
                recommendations.append("如需提高正式程度，可增加正式词汇的使用")

            # 从LLM分析中提取建议
            advanced_features = enhanced_result.get("advanced_features", {})
            comprehensive_analysis = advanced_features.get("comprehensive_analysis", {})

            if comprehensive_analysis.get("success"):
                parsed_analysis = comprehensive_analysis.get("parsed_analysis", {})
                style_summary = parsed_analysis.get("style_summary", {})

                improvement_suggestions = style_summary.get("改进建议", "")
                if improvement_suggestions and improvement_suggestions != "无":
                    recommendations.append(improvement_suggestions)

            return recommendations[:5]  # 最多返回5条建议

        except Exception:
            return ["建议保持当前写作风格的一致性"]

    def analyze_with_semantic_behavior(self, document_content: str, document_name: str = None) -> Dict[str, Any]:
        """
        使用语义空间行为算法进行文风分析

        Args:
            document_content: 文档内容
            document_name: 文档名称

        Returns:
            语义行为分析结果
        """
        if not self.use_enhanced_features or not self.enhanced_processor:
            return {"error": "增强功能未启用，无法进行语义行为分析"}

        try:
            print(f"🧠 开始语义空间行为分析: {document_name or '未命名文档'}")

            # 使用综合处理器的语义分析功能
            semantic_result = self.enhanced_processor.analyze_semantic_behavior(
                document_content, document_name, "comprehensive"
            )

            if semantic_result.get("success"):
                # 转换为兼容的格式
                analysis_result = {
                    "document_name": document_name or "未命名文档",
                    "analysis_time": datetime.now().isoformat(),
                    "analysis_method": "semantic_behavior",
                    "semantic_analysis": semantic_result,
                    "style_features": self._extract_semantic_style_features(semantic_result),
                    "style_type": self._determine_semantic_style_type(semantic_result),
                    "confidence_score": self._calculate_semantic_confidence(semantic_result),
                    "style_prompt": self._generate_semantic_style_prompt(semantic_result),
                    "template_id": self._generate_template_id(document_name, {}),
                    "detailed_analysis": semantic_result.get("comprehensive_insights", {}),
                    "writing_recommendations": self._generate_semantic_recommendations(semantic_result),
                    "style_comparison": {}
                }

                print("✅ 语义空间行为分析完成")
                return analysis_result
            else:
                return {"error": f"语义分析失败: {semantic_result.get('error', '未知错误')}"}

        except Exception as e:
            return {"error": f"语义空间行为分析失败: {str(e)}"}

    def _extract_semantic_style_features(self, semantic_result: Dict[str, Any]) -> Dict[str, Any]:
        """从语义分析结果中提取风格特征"""
        style_features = {}

        try:
            # 从最终画像中提取特征
            final_profile = semantic_result.get("semantic_analysis", {}).get("final_profile", {})
            if final_profile.get("success"):
                style_scores = final_profile.get("style_scores", {})

                # 映射到传统特征名称
                style_features.update({
                    "conceptual_organization": style_scores.get("conceptual_organization", 3.0),
                    "semantic_coherence": style_scores.get("semantic_coherence", 3.0),
                    "creative_association": style_scores.get("creative_association", 3.0),
                    "emotional_expression": style_scores.get("emotional_expression", 3.0),
                    "cognitive_complexity": style_scores.get("cognitive_complexity", 3.0),
                    "thematic_focus": style_scores.get("thematic_focus", 3.0)
                })

                # 添加特征向量长度
                feature_vector = final_profile.get("feature_vector", [])
                style_features["feature_vector_length"] = len(feature_vector)
                style_features["feature_vector_norm"] = final_profile.get("comparative_metrics", {}).get("feature_vector_norm", 0.0)

        except Exception as e:
            style_features["extraction_error"] = str(e)

        return style_features

    def _determine_semantic_style_type(self, semantic_result: Dict[str, Any]) -> str:
        """从语义分析结果中确定风格类型"""
        try:
            final_profile = semantic_result.get("semantic_analysis", {}).get("final_profile", {})
            if final_profile.get("success"):
                classification = final_profile.get("style_classification", {})
                primary_style = classification.get("primary_style", "")

                # 映射到内部风格类型
                style_mapping = {
                    "系统性思维型": "academic_research",
                    "逻辑连贯型": "business_professional",
                    "创新联想型": "creative_narrative",
                    "情感表达型": "narrative_descriptive",
                    "复杂思维型": "academic_research",
                    "专注聚焦型": "formal_official"
                }

                return style_mapping.get(primary_style, "business_professional")

            return "business_professional"

        except Exception:
            return "business_professional"

    def _calculate_semantic_confidence(self, semantic_result: Dict[str, Any]) -> float:
        """计算语义分析的置信度"""
        try:
            # 基于分析成功的阶段数量
            analysis_summary = semantic_result.get("semantic_analysis", {}).get("analysis_summary", {})
            stages_completed = analysis_summary.get("stages_completed", 0)
            max_stages = 4

            base_confidence = stages_completed / max_stages

            # 如果有最终画像，提高置信度
            final_profile = semantic_result.get("semantic_analysis", {}).get("final_profile", {})
            if final_profile.get("success"):
                profile_confidence = final_profile.get("comparative_metrics", {}).get("style_score_average", 3.0) / 5.0
                base_confidence = (base_confidence + profile_confidence) / 2

            return min(1.0, max(0.0, base_confidence))

        except Exception:
            return 0.7

    def _generate_semantic_style_prompt(self, semantic_result: Dict[str, Any]) -> str:
        """生成语义风格提示词"""
        try:
            final_profile = semantic_result.get("semantic_analysis", {}).get("final_profile", {})
            if final_profile.get("success"):
                classification = final_profile.get("style_classification", {})
                primary_style = classification.get("primary_style", "综合型")
                characteristics = classification.get("style_characteristics", [])

                prompt_parts = [f"请按照{primary_style}进行写作"]

                if characteristics:
                    prompt_parts.append(f"特点：{', '.join(characteristics)}")

                # 添加具体的语义指导
                style_scores = final_profile.get("style_scores", {})
                if style_scores.get("conceptual_organization", 0) > 4.0:
                    prompt_parts.append("注重概念的系统性组织")
                if style_scores.get("creative_association", 0) > 4.0:
                    prompt_parts.append("发挥创新联想能力")
                if style_scores.get("emotional_expression", 0) > 4.0:
                    prompt_parts.append("增强情感表达力")

                return "；".join(prompt_parts)

            return "请保持语义连贯和逻辑清晰的写作风格"

        except Exception:
            return "请保持原有的写作风格"

    def _generate_semantic_recommendations(self, semantic_result: Dict[str, Any]) -> List[str]:
        """生成语义分析建议"""
        recommendations = []

        try:
            # 从综合洞察中提取建议
            comprehensive_insights = semantic_result.get("comprehensive_insights", {})
            actionable_recommendations = comprehensive_insights.get("actionable_recommendations", [])
            recommendations.extend(actionable_recommendations[:3])

            # 从最终画像中提取改进建议
            final_profile = semantic_result.get("semantic_analysis", {}).get("final_profile", {})
            if final_profile.get("success"):
                profile_summary = final_profile.get("profile_summary", {})
                improvements = profile_summary.get("potential_improvements", [])
                for improvement in improvements[:2]:
                    recommendations.append(f"建议提升{improvement}")

            # 如果没有具体建议，提供通用建议
            if not recommendations:
                recommendations = [
                    "建议保持概念组织的系统性",
                    "注意语义连贯性和逻辑性",
                    "适当增加创新性表达"
                ]

            return recommendations[:5]

        except Exception:
            return ["建议保持当前的语义风格特征"]

    def compare_semantic_styles(self, document1_content: str, document2_content: str,
                              doc1_name: str = None, doc2_name: str = None) -> Dict[str, Any]:
        """
        比较两个文档的语义风格

        Args:
            document1_content: 第一个文档内容
            document2_content: 第二个文档内容
            doc1_name: 第一个文档名称
            doc2_name: 第二个文档名称

        Returns:
            语义风格比较结果
        """
        if not self.use_enhanced_features or not self.enhanced_processor:
            return {"error": "增强功能未启用，无法进行语义风格比较"}

        comparison_result = {
            "comparison_time": datetime.now().isoformat(),
            "document1_name": doc1_name or "文档1",
            "document2_name": doc2_name or "文档2",
            "document1_analysis": {},
            "document2_analysis": {},
            "semantic_comparison": {},
            "style_compatibility": "unknown",
            "comparison_summary": {},
            "success": False
        }

        try:
            print(f"🔍 开始语义风格比较: {doc1_name or '文档1'} vs {doc2_name or '文档2'}")

            # 分析第一个文档
            doc1_analysis = self.analyze_with_semantic_behavior(document1_content, doc1_name)
            comparison_result["document1_analysis"] = doc1_analysis

            # 分析第二个文档
            doc2_analysis = self.analyze_with_semantic_behavior(document2_content, doc2_name)
            comparison_result["document2_analysis"] = doc2_analysis

            # 如果两个分析都成功，进行比较
            if (not doc1_analysis.get("error") and not doc2_analysis.get("error") and
                self.enhanced_processor.semantic_analysis_enabled):

                # 使用综合处理器的语义比较功能
                semantic_comparison = self.enhanced_processor.compare_semantic_profiles(
                    document1_content, document2_content, doc1_name, doc2_name
                )
                comparison_result["semantic_comparison"] = semantic_comparison

                # 生成兼容性评估
                if semantic_comparison.get("success"):
                    profile_comparison = semantic_comparison.get("profile_comparison", {})
                    similarity_score = profile_comparison.get("similarity_score", 0.0)

                    if similarity_score > 0.8:
                        comparison_result["style_compatibility"] = "高度兼容"
                    elif similarity_score > 0.6:
                        comparison_result["style_compatibility"] = "较为兼容"
                    elif similarity_score > 0.4:
                        comparison_result["style_compatibility"] = "部分兼容"
                    else:
                        comparison_result["style_compatibility"] = "差异较大"

                # 生成比较摘要
                comparison_result["comparison_summary"] = self._generate_semantic_comparison_summary(
                    doc1_analysis, doc2_analysis, semantic_comparison
                )

                comparison_result["success"] = True
                print("✅ 语义风格比较完成")
            else:
                comparison_result["error"] = "文档分析失败，无法进行语义比较"
                print("❌ 语义风格比较失败")

        except Exception as e:
            comparison_result["error"] = str(e)
            print(f"❌ 语义风格比较失败: {str(e)}")

        return comparison_result

    def _generate_semantic_comparison_summary(self, doc1_analysis: Dict[str, Any],
                                            doc2_analysis: Dict[str, Any],
                                            semantic_comparison: Dict[str, Any]) -> Dict[str, Any]:
        """生成语义比较摘要"""
        summary = {
            "overall_similarity": 0.0,
            "style_differences": [],
            "common_characteristics": [],
            "recommendation": ""
        }

        try:
            # 整体相似度
            if semantic_comparison.get("success"):
                profile_comparison = semantic_comparison.get("profile_comparison", {})
                summary["overall_similarity"] = profile_comparison.get("similarity_score", 0.0)

                # 维度差异
                dimension_diffs = profile_comparison.get("dimension_differences", {})
                differences = []
                similarities = []

                for dimension, diff_data in dimension_diffs.items():
                    difference = diff_data.get("difference", 0)
                    if difference > 1.0:  # 差异较大
                        differences.append(f"{dimension}差异较大")
                    elif difference < 0.5:  # 相似度较高
                        similarities.append(f"{dimension}较为相似")

                summary["style_differences"] = differences[:3]
                summary["common_characteristics"] = similarities[:3]

            # 生成建议
            similarity_score = summary["overall_similarity"]
            if similarity_score > 0.7:
                summary["recommendation"] = "两个文档风格相近，可以进行风格对齐"
            elif similarity_score > 0.4:
                summary["recommendation"] = "两个文档风格有一定差异，建议重点调整差异较大的维度"
            else:
                summary["recommendation"] = "两个文档风格差异较大，需要全面的风格迁移"

        except Exception as e:
            summary["error"] = str(e)

        return summary

    def _get_document_statistics(self, content: str) -> Dict[str, Any]:
        """获取文档基础统计信息"""
        lines = content.split('\n')
        sentences = re.split(r'[。！？.!?]', content)
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        words = re.findall(r'[\u4e00-\u9fff]+', content)

        return {
            "total_characters": len(content),
            "total_lines": len(lines),
            "total_sentences": len([s for s in sentences if s.strip()]),
            "total_paragraphs": len(paragraphs),
            "total_words": len(words),
            "average_sentence_length": round(len(content) / max(len([s for s in sentences if s.strip()]), 1), 1),
            "average_paragraph_length": round(len(content) / max(len(paragraphs), 1), 1),
            "reading_time_minutes": round(len(words) / 200, 1)  # 假设每分钟200字
        }

    def _generate_detailed_analysis(self, content: str, features: Dict[str, Any], style_type: str) -> Dict[str, Any]:
        """生成详细分析报告"""
        return {
            "readability_analysis": self._analyze_readability(content, features),
            "tone_analysis": self._analyze_tone_details(content, features),
            "structure_analysis": self._analyze_structure_details(content, features),
            "vocabulary_analysis": self._analyze_vocabulary_details(content, features),
            "style_consistency": self._analyze_style_consistency(content, features)
        }

    def _generate_writing_recommendations(self, features: Dict[str, Any], style_type: str) -> List[str]:
        """生成写作建议"""
        recommendations = []

        # 基于句式结构的建议
        sentence_features = features.get("sentence_structure", {})
        avg_length = sentence_features.get("average_length", 15)

        if avg_length > 30:
            recommendations.append("句子偏长，建议适当拆分为短句以提高可读性")
        elif avg_length < 10:
            recommendations.append("句子偏短，可以适当增加句子的完整性和表达力")

        # 基于词汇选择的建议
        vocab_features = features.get("vocabulary_choice", {})
        modifier_usage = vocab_features.get("modifier_usage", 0)

        if modifier_usage > 50:
            recommendations.append("修饰词使用较多，建议精简表达，突出重点")
        elif modifier_usage < 10:
            recommendations.append("可以适当增加修饰词，丰富表达层次")

        # 基于文风类型的建议
        if style_type == "business_professional":
            recommendations.append("保持专业性，注意用词准确性和逻辑清晰")
        elif style_type == "academic_research":
            recommendations.append("增强论证严密性，注意引用和数据支撑")
        elif style_type == "concise_practical":
            recommendations.append("继续保持简洁明了，突出实用性")

        return recommendations

    def _generate_style_comparison(self, features: Dict[str, Any]) -> Dict[str, Any]:
        """生成风格对比分析"""
        return {
            "formal_vs_informal": self._compare_formality(features),
            "technical_vs_general": self._compare_technicality(features),
            "objective_vs_subjective": self._compare_objectivity(features),
            "concise_vs_elaborate": self._compare_conciseness(features)
        }

    def _analyze_style_features(self, content: str) -> Dict[str, Any]:
        """分析文风特征"""
        if not content or not content.strip():
            return self._get_empty_features()

        features = {}

        try:
            # 句式结构分析
            features["sentence_structure"] = self._analyze_sentence_structure(content)

            # 词汇选择分析
            features["vocabulary_choice"] = self._analyze_vocabulary_choice(content)

            # 表达方式分析
            features["expression_style"] = self._analyze_expression_style(content)

            # 文本组织分析
            features["text_organization"] = self._analyze_text_organization(content)

            # 语言习惯分析
            features["language_habits"] = self._analyze_language_habits(content)

            # 新增：情感色彩分析
            features["emotional_tone"] = self._analyze_emotional_tone(content)

            # 新增：专业性分析
            features["professionalism"] = self._analyze_professionalism(content)

            # 新增：修辞手法分析
            features["rhetorical_devices"] = self._analyze_rhetorical_devices(content)

        except Exception as e:
            print(f"文风特征分析出错: {str(e)}")
            features = self._get_empty_features()

        return features

    def _get_empty_features(self) -> Dict[str, Any]:
        """获取空的特征结构"""
        return {
            "sentence_structure": {"average_length": 0, "long_short_ratio": 0, "complex_ratio": 0, "total_sentences": 0},
            "vocabulary_choice": {"formality_score": 0, "technical_density": 0, "modifier_usage": 0, "action_verb_ratio": 0},
            "expression_style": {"passive_active_ratio": 0, "person_usage": {"first_person": 0, "second_person": 0, "third_person": 0}, "tone_strength": 0},
            "text_organization": {"paragraph_count": 0, "average_paragraph_length": 0, "connector_density": 0, "summary_usage": 0},
            "language_habits": {"colloquial_level": 0, "formal_structure_usage": 0, "de_structure_density": 0},
            "emotional_tone": {"positive_ratio": 0, "negative_ratio": 0, "neutral_ratio": 1.0, "intensity": 0},
            "professionalism": {"domain_specificity": 0, "authority_indicators": 0, "precision_level": 0},
            "rhetorical_devices": {"metaphor_usage": 0, "parallel_structure": 0, "question_usage": 0}
        }
    
    def _analyze_sentence_structure(self, content: str) -> Dict[str, Any]:
        """分析句式结构"""
        # 清理内容，移除多余的空白字符
        content = re.sub(r'\s+', ' ', content.strip())

        # 更精确的句子分割
        sentences = re.split(r'[。！？；\.!?;]', content)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 3]

        if not sentences:
            return {
                "average_length": 0,
                "long_short_ratio": 0,
                "complex_ratio": 0,
                "total_sentences": 0
            }

        # 计算平均句长
        total_chars = sum(len(s) for s in sentences)
        average_length = total_chars / len(sentences)

        # 长短句比例（长句定义为超过25字）
        long_sentences = [s for s in sentences if len(s) > 25]
        short_sentences = [s for s in sentences if len(s) <= 15]
        long_short_ratio = len(long_sentences) / len(sentences) if sentences else 0

        # 复合句比例（包含逗号、分号、连词等的句子）
        complex_patterns = [r'，', r'；', r'：', r'而且', r'但是', r'然而', r'因为', r'所以', r'如果', r'虽然']
        complex_sentences = []
        for sentence in sentences:
            if any(re.search(pattern, sentence) for pattern in complex_patterns):
                complex_sentences.append(sentence)
        complex_ratio = len(complex_sentences) / len(sentences) if sentences else 0

        return {
            "average_length": round(average_length, 1),
            "long_short_ratio": round(long_short_ratio, 3),
            "complex_ratio": round(complex_ratio, 3),
            "total_sentences": len(sentences),
            "long_sentences": len(long_sentences),
            "short_sentences": len(short_sentences)
        }
    
    def _analyze_vocabulary_choice(self, content: str) -> Dict[str, Any]:
        """分析词汇选择"""
        # 清理内容
        content_clean = re.sub(r'\s+', '', content)
        total_chars = len(content_clean)

        if total_chars == 0:
            return {
                "formality_score": 0,
                "technical_density": 0,
                "modifier_usage": 0,
                "action_verb_ratio": 0
            }

        # 正式词汇模式
        formal_words = ["根据", "按照", "依据", "鉴于", "基于", "关于", "针对", "为了", "通过", "采用",
                       "进行", "实施", "开展", "落实", "确保", "保证", "维护", "促进", "推动", "加强"]
        formal_count = sum(content.count(word) for word in formal_words)

        # 专业术语模式
        technical_patterns = [
            r'[\u4e00-\u9fff]+系统', r'[\u4e00-\u9fff]+技术', r'[\u4e00-\u9fff]+方案',
            r'[\u4e00-\u9fff]+标准', r'[\u4e00-\u9fff]+规范', r'[\u4e00-\u9fff]+平台',
            r'[\u4e00-\u9fff]+模式', r'[\u4e00-\u9fff]+机制', r'[\u4e00-\u9fff]+流程'
        ]
        technical_count = sum(len(re.findall(pattern, content)) for pattern in technical_patterns)

        # 修饰词使用
        modifiers = ["很", "非常", "特别", "极其", "相当", "比较", "较为", "十分", "更加", "进一步"]
        modifier_count = sum(content.count(word) for word in modifiers)

        # 动词类型（动作动词）
        action_verbs = ["实施", "执行", "开展", "推进", "落实", "完成", "达成", "实现", "提升", "优化",
                       "改进", "建设", "构建", "发展", "创新", "突破", "解决", "处理", "管理", "运营"]
        action_count = sum(content.count(word) for word in action_verbs)

        # 计算密度（每千字）
        multiplier = 1000 / total_chars if total_chars > 0 else 0

        return {
            "formality_score": round(formal_count * multiplier, 2),
            "technical_density": round(technical_count * multiplier, 2),
            "modifier_usage": round(modifier_count * multiplier, 2),
            "action_verb_ratio": round(action_count * multiplier, 2)
        }
    
    def _analyze_expression_style(self, content: str) -> Dict[str, Any]:
        """分析表达方式"""
        total_chars = len(content)

        if total_chars == 0:
            return {
                "passive_active_ratio": 0,
                "person_usage": {"first_person": 0, "second_person": 0, "third_person": 0},
                "tone_strength": 0
            }

        # 被动语态检测（更精确的模式）
        passive_patterns = [
            r'[\u4e00-\u9fff]+被[\u4e00-\u9fff]+', r'[\u4e00-\u9fff]+受到[\u4e00-\u9fff]+',
            r'[\u4e00-\u9fff]+得到[\u4e00-\u9fff]+', r'[\u4e00-\u9fff]+获得[\u4e00-\u9fff]+',
            r'[\u4e00-\u9fff]+遭到[\u4e00-\u9fff]+', r'[\u4e00-\u9fff]+受[\u4e00-\u9fff]+影响'
        ]
        passive_count = sum(len(re.findall(pattern, content)) for pattern in passive_patterns)

        # 主动语态检测
        active_patterns = [
            r'[\u4e00-\u9fff]+进行[\u4e00-\u9fff]+', r'[\u4e00-\u9fff]+开展[\u4e00-\u9fff]+',
            r'[\u4e00-\u9fff]+实施[\u4e00-\u9fff]+', r'[\u4e00-\u9fff]+推进[\u4e00-\u9fff]+',
            r'[\u4e00-\u9fff]+完成[\u4e00-\u9fff]+', r'[\u4e00-\u9fff]+建设[\u4e00-\u9fff]+'
        ]
        active_count = sum(len(re.findall(pattern, content)) for pattern in active_patterns)

        # 人称使用统计
        first_person = content.count('我') + content.count('我们') + content.count('本人') + content.count('笔者')
        second_person = content.count('你') + content.count('您') + content.count('你们') + content.count('各位')
        third_person = content.count('他') + content.count('她') + content.count('它') + content.count('他们') + content.count('她们')

        # 语气强度分析
        strong_words = ['必须', '务必', '严禁', '绝对', '一定', '坚决', '严格', '切实']
        mild_words = ['建议', '希望', '可以', '尽量', '适当', '酌情', '可能', '或许']

        strong_tone = sum(content.count(word) for word in strong_words)
        mild_tone = sum(content.count(word) for word in mild_words)

        # 计算比例
        total_voice_patterns = passive_count + active_count
        passive_ratio = passive_count / total_voice_patterns if total_voice_patterns > 0 else 0

        return {
            "passive_active_ratio": round(passive_ratio, 3),
            "person_usage": {
                "first_person": first_person,
                "second_person": second_person,
                "third_person": third_person
            },
            "tone_strength": round((strong_tone - mild_tone) / total_chars * 1000, 2) if total_chars > 0 else 0,
            "passive_count": passive_count,
            "active_count": active_count
        }
    
    def _analyze_text_organization(self, content: str) -> Dict[str, Any]:
        """分析文本组织"""
        # 段落分析（按换行符分割）
        paragraphs = re.split(r'\n\s*\n', content)
        paragraphs = [p.strip() for p in paragraphs if p.strip() and len(p.strip()) > 10]

        total_chars = len(content)

        if total_chars == 0:
            return {
                "paragraph_count": 0,
                "average_paragraph_length": 0,
                "connector_density": 0,
                "summary_usage": 0
            }

        # 逻辑连接词（更全面）
        connectors = [
            "首先", "其次", "然后", "最后", "因此", "所以", "但是", "然而", "此外", "另外",
            "同时", "而且", "并且", "不过", "虽然", "尽管", "由于", "鉴于", "基于", "根据",
            "总之", "综上", "另一方面", "与此同时", "相比之下", "换言之", "也就是说"
        ]
        connector_count = sum(content.count(word) for word in connectors)

        # 总结性表达
        summary_words = ["总之", "综上", "总的来说", "综合以上", "总而言之", "综上所述", "由此可见", "可以看出"]
        summary_count = sum(content.count(word) for word in summary_words)

        # 列举标识
        enumeration_patterns = [r'[一二三四五六七八九十]、', r'\d+[\.、]', r'[（(]\d+[）)]']
        enumeration_count = sum(len(re.findall(pattern, content)) for pattern in enumeration_patterns)

        return {
            "paragraph_count": len(paragraphs),
            "average_paragraph_length": round(sum(len(p) for p in paragraphs) / len(paragraphs), 1) if paragraphs else 0,
            "connector_density": round(connector_count / total_chars * 1000, 2),
            "summary_usage": summary_count,
            "enumeration_usage": enumeration_count
        }
    
    def _analyze_language_habits(self, content: str) -> Dict[str, Any]:
        """分析语言习惯"""
        total_chars = len(content)

        if total_chars == 0:
            return {
                "colloquial_level": 0,
                "formal_structure_usage": 0,
                "de_structure_density": 0
            }

        # 口语化词汇
        colloquial_words = [
            "挺", "蛮", "特别", "超级", "好像", "感觉", "觉得", "应该", "可能", "大概",
            "差不多", "基本上", "一般来说", "说实话", "老实说", "坦白说", "真的", "确实"
        ]
        colloquial_count = sum(content.count(word) for word in colloquial_words)

        # 书面语结构词
        formal_structures = [
            "之", "其", "所", "乃", "即", "亦", "且", "而", "于", "以", "为", "与",
            "及", "或", "若", "则", "故", "因", "由", "自", "从", "向", "至", "及其"
        ]
        formal_count = sum(content.count(word) for word in formal_structures)

        # "的"字结构密度
        de_count = content.count('的')

        # 书面语句式
        formal_patterns = [r'[\u4e00-\u9fff]+之[\u4e00-\u9fff]+', r'所[\u4e00-\u9fff]+', r'其[\u4e00-\u9fff]+']
        formal_pattern_count = sum(len(re.findall(pattern, content)) for pattern in formal_patterns)

        return {
            "colloquial_level": round(colloquial_count / total_chars * 1000, 2),
            "formal_structure_usage": round((formal_count + formal_pattern_count) / total_chars * 1000, 2),
            "de_structure_density": round(de_count / total_chars * 1000, 2)
        }

    def _analyze_emotional_tone(self, content: str) -> Dict[str, Any]:
        """分析情感色彩"""
        total_chars = len(content)
        if total_chars == 0:
            return {"positive_ratio": 0, "negative_ratio": 0, "neutral_ratio": 1.0, "intensity": 0}

        # 积极词汇
        positive_words = ["优秀", "卓越", "成功", "进步", "提升", "改善", "创新", "突破", "发展", "增长",
                         "满意", "高兴", "喜悦", "赞扬", "表彰", "肯定", "支持", "鼓励", "希望", "信心"]

        # 消极词汇
        negative_words = ["问题", "困难", "挑战", "不足", "缺陷", "错误", "失败", "下降", "减少", "损失",
                         "担心", "忧虑", "批评", "质疑", "反对", "拒绝", "否定", "警告", "风险", "危机"]

        # 强度词汇
        intensity_words = ["非常", "极其", "特别", "十分", "相当", "很", "最", "更", "进一步", "大幅"]

        positive_count = sum(content.count(word) for word in positive_words)
        negative_count = sum(content.count(word) for word in negative_words)
        intensity_count = sum(content.count(word) for word in intensity_words)

        total_emotional = positive_count + negative_count
        if total_emotional == 0:
            return {"positive_ratio": 0, "negative_ratio": 0, "neutral_ratio": 1.0, "intensity": 0}

        return {
            "positive_ratio": round(positive_count / total_emotional, 2),
            "negative_ratio": round(negative_count / total_emotional, 2),
            "neutral_ratio": round(1 - (positive_count + negative_count) / total_emotional, 2),
            "intensity": round(intensity_count / total_chars * 1000, 2)
        }

    def _analyze_professionalism(self, content: str) -> Dict[str, Any]:
        """分析专业性特征"""
        total_chars = len(content)
        if total_chars == 0:
            return {"domain_specificity": 0, "authority_indicators": 0, "precision_level": 0}

        # 权威性指标词汇
        authority_words = ["根据", "依据", "按照", "研究表明", "数据显示", "统计", "调查", "分析",
                          "专家", "学者", "权威", "官方", "正式", "法规", "标准", "规范"]

        # 精确性指标
        precision_patterns = [r'\d+%', r'\d+\.\d+', r'第\d+', r'\d+年\d+月', r'\d+万', r'\d+亿']

        # 领域特定词汇（示例）
        domain_words = ["技术", "系统", "平台", "方案", "策略", "机制", "模式", "框架", "体系", "流程"]

        authority_count = sum(content.count(word) for word in authority_words)
        domain_count = sum(content.count(word) for word in domain_words)
        precision_count = sum(len(re.findall(pattern, content)) for pattern in precision_patterns)

        return {
            "domain_specificity": round(domain_count / total_chars * 1000, 2),
            "authority_indicators": round(authority_count / total_chars * 1000, 2),
            "precision_level": round(precision_count / total_chars * 1000, 2)
        }

    def _analyze_rhetorical_devices(self, content: str) -> Dict[str, Any]:
        """分析修辞手法"""
        total_chars = len(content)
        if total_chars == 0:
            return {"metaphor_usage": 0, "parallel_structure": 0, "question_usage": 0}

        # 比喻词汇
        metaphor_words = ["如同", "好比", "犹如", "仿佛", "像", "似", "宛如", "恰似"]

        # 排比结构模式
        parallel_patterns = [r'[\u4e00-\u9fff]+，[\u4e00-\u9fff]+，[\u4e00-\u9fff]+',
                           r'不仅[\u4e00-\u9fff]+，而且[\u4e00-\u9fff]+',
                           r'既[\u4e00-\u9fff]+，又[\u4e00-\u9fff]+']

        # 疑问句
        question_count = content.count('？') + content.count('?')

        metaphor_count = sum(content.count(word) for word in metaphor_words)
        parallel_count = sum(len(re.findall(pattern, content)) for pattern in parallel_patterns)

        return {
            "metaphor_usage": round(metaphor_count / total_chars * 1000, 2),
            "parallel_structure": round(parallel_count / total_chars * 1000, 2),
            "question_usage": round(question_count / total_chars * 1000, 2)
        }

    def _analyze_readability(self, content: str, features: Dict[str, Any]) -> Dict[str, Any]:
        """分析可读性"""
        sentence_features = features.get("sentence_structure", {})
        vocab_features = features.get("vocabulary_choice", {})

        # 计算可读性分数 (简化版)
        avg_sentence_length = sentence_features.get("average_length", 15)
        technical_density = vocab_features.get("technical_density", 0)

        readability_score = max(0, min(100, 100 - (avg_sentence_length - 15) * 2 - technical_density))

        if readability_score >= 80:
            level = "很容易阅读"
        elif readability_score >= 60:
            level = "较容易阅读"
        elif readability_score >= 40:
            level = "中等难度"
        elif readability_score >= 20:
            level = "较难阅读"
        else:
            level = "很难阅读"

        return {
            "readability_score": round(readability_score, 1),
            "readability_level": level,
            "factors": {
                "sentence_complexity": "高" if avg_sentence_length > 25 else "中" if avg_sentence_length > 15 else "低",
                "vocabulary_difficulty": "高" if technical_density > 20 else "中" if technical_density > 10 else "低"
            }
        }

    def _analyze_tone_details(self, content: str, features: Dict[str, Any]) -> Dict[str, Any]:
        """分析语调详情"""
        emotional_features = features.get("emotional_tone", {})

        # 分析语调倾向
        positive_words = ["好", "优秀", "成功", "提升", "改善", "创新", "发展", "进步"]
        negative_words = ["问题", "困难", "挑战", "不足", "缺陷", "失败", "下降", "减少"]
        neutral_words = ["分析", "研究", "探讨", "考虑", "建议", "方案", "计划", "实施"]

        positive_count = sum(content.count(word) for word in positive_words)
        negative_count = sum(content.count(word) for word in negative_words)
        neutral_count = sum(content.count(word) for word in neutral_words)

        total_tone_words = positive_count + negative_count + neutral_count

        if total_tone_words > 0:
            tone_distribution = {
                "positive_ratio": round(positive_count / total_tone_words * 100, 1),
                "negative_ratio": round(negative_count / total_tone_words * 100, 1),
                "neutral_ratio": round(neutral_count / total_tone_words * 100, 1)
            }
        else:
            tone_distribution = {"positive_ratio": 33.3, "negative_ratio": 33.3, "neutral_ratio": 33.3}

        return {
            "tone_distribution": tone_distribution,
            "dominant_tone": max(tone_distribution.items(), key=lambda x: x[1])[0].replace("_ratio", ""),
            "emotional_intensity": emotional_features.get("intensity_score", 0),
            "tone_consistency": "高" if max(tone_distribution.values()) > 60 else "中" if max(tone_distribution.values()) > 40 else "低"
        }

    def _analyze_structure_details(self, content: str, features: Dict[str, Any]) -> Dict[str, Any]:
        """分析结构详情"""
        org_features = features.get("text_organization", {})

        return {
            "paragraph_structure": {
                "paragraph_count": org_features.get("paragraph_count", 0),
                "average_length": org_features.get("average_paragraph_length", 0),
                "length_consistency": "高" if org_features.get("average_paragraph_length", 0) > 100 else "中"
            },
            "logical_flow": {
                "connector_usage": org_features.get("connector_density", 0),
                "enumeration_usage": org_features.get("enumeration_usage", 0),
                "summary_usage": org_features.get("summary_usage", 0)
            },
            "organization_score": min(100, (org_features.get("connector_density", 0) * 2 +
                                           org_features.get("enumeration_usage", 0) * 5 +
                                           org_features.get("summary_usage", 0) * 10))
        }

    def _analyze_vocabulary_details(self, content: str, features: Dict[str, Any]) -> Dict[str, Any]:
        """分析词汇详情"""
        vocab_features = features.get("vocabulary_choice", {})

        return {
            "vocabulary_richness": {
                "unique_words": len(set(re.findall(r'[\u4e00-\u9fff]+', content))),
                "total_words": len(re.findall(r'[\u4e00-\u9fff]+', content)),
                "diversity_ratio": round(len(set(re.findall(r'[\u4e00-\u9fff]+', content))) /
                                       max(len(re.findall(r'[\u4e00-\u9fff]+', content)), 1) * 100, 1)
            },
            "word_complexity": {
                "technical_density": vocab_features.get("technical_density", 0),
                "formality_score": vocab_features.get("formality_score", 0),
                "modifier_usage": vocab_features.get("modifier_usage", 0)
            },
            "action_orientation": {
                "action_verb_ratio": vocab_features.get("action_verb_ratio", 0),
                "passive_voice_usage": self._count_passive_voice(content)
            }
        }

    def _analyze_style_consistency(self, content: str, features: Dict[str, Any]) -> Dict[str, Any]:
        """分析风格一致性"""
        # 分段分析风格一致性
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]

        if len(paragraphs) < 2:
            return {"consistency_score": 100, "variation_level": "无法评估"}

        # 简化的一致性分析
        consistency_factors = []

        # 句长一致性
        sentence_lengths = []
        for para in paragraphs:
            sentences = re.split(r'[。！？.!?]', para)
            avg_len = sum(len(s) for s in sentences if s.strip()) / max(len([s for s in sentences if s.strip()]), 1)
            sentence_lengths.append(avg_len)

        if sentence_lengths:
            length_variance = max(sentence_lengths) - min(sentence_lengths)
            consistency_factors.append(max(0, 100 - length_variance * 2))

        consistency_score = sum(consistency_factors) / max(len(consistency_factors), 1)

        return {
            "consistency_score": round(consistency_score, 1),
            "variation_level": "低" if consistency_score > 80 else "中" if consistency_score > 60 else "高",
            "factors_analyzed": ["句长一致性", "词汇使用", "语调变化"]
        }

    def _count_passive_voice(self, content: str) -> float:
        """统计被动语态使用频率"""
        passive_patterns = [r'被[\u4e00-\u9fff]+', r'受到[\u4e00-\u9fff]+', r'得到[\u4e00-\u9fff]+']
        passive_count = sum(len(re.findall(pattern, content)) for pattern in passive_patterns)
        total_chars = len(content)
        return round(passive_count / total_chars * 1000, 2) if total_chars > 0 else 0

    def _compare_formality(self, features: Dict[str, Any]) -> Dict[str, Any]:
        """对比正式性"""
        vocab_features = features.get("vocabulary_choice", {})
        habit_features = features.get("language_habits", {})

        formality_score = vocab_features.get("formality_score", 0)
        formal_structure = habit_features.get("formal_structure_usage", 0)
        colloquial_level = habit_features.get("colloquial_level", 0)

        overall_formality = (formality_score + formal_structure - colloquial_level) / 3

        return {
            "formality_level": "高" if overall_formality > 20 else "中" if overall_formality > 10 else "低",
            "formal_score": round(overall_formality, 1),
            "indicators": {
                "formal_vocabulary": formality_score,
                "formal_structures": formal_structure,
                "colloquial_elements": colloquial_level
            }
        }

    def _compare_technicality(self, features: Dict[str, Any]) -> Dict[str, Any]:
        """对比技术性"""
        vocab_features = features.get("vocabulary_choice", {})
        professional_features = features.get("professionalism", {})

        technical_density = vocab_features.get("technical_density", 0)
        professional_score = professional_features.get("professional_score", 0)

        return {
            "technicality_level": "高" if technical_density > 20 else "中" if technical_density > 10 else "低",
            "technical_score": round(technical_density, 1),
            "professional_score": round(professional_score, 1),
            "balance": "技术性强" if technical_density > professional_score else "通用性强"
        }

    def _compare_objectivity(self, features: Dict[str, Any]) -> Dict[str, Any]:
        """对比客观性"""
        emotional_features = features.get("emotional_tone", {})
        expression_features = features.get("expression_style", {})

        emotional_intensity = emotional_features.get("intensity_score", 0)
        assertive_score = expression_features.get("assertive_score", 0)

        objectivity_score = max(0, 100 - emotional_intensity * 10 - assertive_score)

        return {
            "objectivity_level": "高" if objectivity_score > 70 else "中" if objectivity_score > 40 else "低",
            "objectivity_score": round(objectivity_score, 1),
            "subjectivity_indicators": {
                "emotional_intensity": emotional_intensity,
                "assertive_tone": assertive_score
            }
        }

    def _compare_conciseness(self, features: Dict[str, Any]) -> Dict[str, Any]:
        """对比简洁性"""
        sentence_features = features.get("sentence_structure", {})
        vocab_features = features.get("vocabulary_choice", {})

        avg_length = sentence_features.get("average_length", 15)
        modifier_usage = vocab_features.get("modifier_usage", 0)

        conciseness_score = max(0, 100 - (avg_length - 15) * 2 - modifier_usage)

        return {
            "conciseness_level": "高" if conciseness_score > 70 else "中" if conciseness_score > 40 else "低",
            "conciseness_score": round(conciseness_score, 1),
            "verbosity_indicators": {
                "sentence_length": avg_length,
                "modifier_density": modifier_usage
            }
        }

    def _generate_enhanced_style_prompt(self, features: Dict[str, Any], style_type: str, detailed_analysis: Dict[str, Any]) -> str:
        """生成增强的文风提示词"""
        prompt_parts = []

        # 基础风格描述
        style_info = self.style_types.get(style_type, {})
        style_name = style_info.get("name", "通用风格")
        characteristics = style_info.get("characteristics", [])

        prompt_parts.append(f"请采用{style_name}进行写作，具体特征包括：{', '.join(characteristics)}。")

        # 句式要求
        sentence_features = features.get("sentence_structure", {})
        avg_length = sentence_features.get("average_length", 15)

        if avg_length > 25:
            prompt_parts.append("使用较长的复合句，注重表达的完整性和逻辑性。")
        elif avg_length < 15:
            prompt_parts.append("使用简洁明了的短句，突出重点，便于理解。")
        else:
            prompt_parts.append("句式长短适中，兼顾表达完整性和可读性。")

        # 词汇要求
        vocab_features = features.get("vocabulary_choice", {})
        formality_score = vocab_features.get("formality_score", 0)
        technical_density = vocab_features.get("technical_density", 0)

        if formality_score > 20:
            prompt_parts.append("使用正式、规范的词汇，避免口语化表达。")
        elif formality_score < 10:
            prompt_parts.append("可以使用相对轻松、自然的表达方式。")

        if technical_density > 15:
            prompt_parts.append("适当使用专业术语，体现专业性。")

        # 语调要求
        tone_analysis = detailed_analysis.get("tone_analysis", {})
        dominant_tone = tone_analysis.get("dominant_tone", "neutral")

        if dominant_tone == "positive":
            prompt_parts.append("保持积极正面的语调，突出优势和成果。")
        elif dominant_tone == "negative":
            prompt_parts.append("客观分析问题，提出建设性意见。")
        else:
            prompt_parts.append("保持客观中性的语调，注重事实陈述。")

        # 结构要求
        org_features = features.get("text_organization", {})
        connector_density = org_features.get("connector_density", 0)

        if connector_density > 10:
            prompt_parts.append("注重逻辑连接，使用适当的过渡词和连接词。")

        # 可读性要求
        readability = detailed_analysis.get("readability_analysis", {})
        readability_level = readability.get("readability_level", "中等难度")

        prompt_parts.append(f"确保文本{readability_level}，适合目标读者群体。")

        return " ".join(prompt_parts)
    
    def _identify_style_type(self, content: str, features: Dict[str, Any]) -> Tuple[str, float]:
        """识别文风类型"""
        scores = {}

        # 安全获取特征值的辅助函数
        def safe_get(feature_dict: Dict, key: str, default: float = 0.0) -> float:
            return feature_dict.get(key, default) if feature_dict else default

        for style_id, style_info in self.style_types.items():
            try:
                score = 0.0

                # 基于特征模式匹配 (权重: 0.2)
                pattern_score = 0.0
                for pattern in style_info.get("typical_patterns", []):
                    if pattern in content:
                        pattern_score += 0.04  # 每个模式0.04分，最多5个模式
                score += min(pattern_score, 0.2)

                # 基于特征分析结果 (权重: 0.8)
                sentence_features = features.get("sentence_structure", {})
                vocab_features = features.get("vocabulary_choice", {})
                expression_features = features.get("expression_style", {})
                org_features = features.get("text_organization", {})
                habit_features = features.get("language_habits", {})
                emotional_features = features.get("emotional_tone", {})
                professional_features = features.get("professionalism", {})

                if style_id == "formal_official":
                    # 正式公文风格
                    score += min(safe_get(vocab_features, "formality_score") * 0.02, 0.2)  # 正式度
                    score += min(safe_get(org_features, "connector_density") * 0.01, 0.1)  # 连接词密度
                    score += min(safe_get(professional_features, "authority_indicators") * 0.02, 0.2)  # 权威性
                    score += min((1 - safe_get(habit_features, "colloquial_level") * 0.01), 0.1)  # 非口语化
                    score += min(safe_get(expression_features, "passive_active_ratio") * 0.2, 0.2)  # 被动语态

                elif style_id == "business_professional":
                    # 商务专业风格
                    score += min(safe_get(vocab_features, "action_verb_ratio") * 0.02, 0.2)  # 动作动词
                    score += min((1 - safe_get(expression_features, "passive_active_ratio")) * 0.3, 0.3)  # 主动语态
                    score += min(safe_get(professional_features, "precision_level") * 0.02, 0.1)  # 精确性
                    score += min(safe_get(emotional_features, "neutral_ratio") * 0.2, 0.2)  # 中性语调

                elif style_id == "academic_research":
                    # 学术研究风格
                    score += min(safe_get(sentence_features, "complex_ratio") * 0.4, 0.4)  # 复杂句比例
                    score += min(safe_get(vocab_features, "technical_density") * 0.05, 0.2)  # 技术密度
                    score += min(safe_get(professional_features, "authority_indicators") * 0.02, 0.2)  # 权威性

                elif style_id == "narrative_descriptive":
                    # 叙述描述风格
                    score += min(safe_get(vocab_features, "modifier_usage") * 0.02, 0.2)  # 修饰词使用
                    score += min(safe_get(habit_features, "colloquial_level") * 0.02, 0.2)  # 口语化程度
                    score += min(safe_get(emotional_features, "intensity") * 0.02, 0.2)  # 情感强度
                    score += min((safe_get(emotional_features, "positive_ratio") +
                                safe_get(emotional_features, "negative_ratio")) * 0.2, 0.2)  # 情感色彩

                elif style_id == "concise_practical":
                    # 简洁实用风格
                    avg_length = safe_get(sentence_features, "average_length", 20)
                    if avg_length < 20:
                        score += 0.3  # 短句偏好
                    else:
                        score += max(0, 0.3 - (avg_length - 20) * 0.01)

                    score += min((1 - safe_get(habit_features, "de_structure_density") * 0.001) * 0.2, 0.2)
                    score += min((1 - safe_get(vocab_features, "modifier_usage") * 0.01) * 0.2, 0.2)
                    score += min(safe_get(vocab_features, "action_verb_ratio") * 0.01, 0.1)

                scores[style_id] = min(max(score, 0.0), 1.0)  # 确保分数在0-1之间

            except Exception as e:
                print(f"计算风格 {style_id} 分数时出错: {str(e)}")
                scores[style_id] = 0.0

        # 找到得分最高的文风类型
        if not scores:
            return "business_professional", 0.5

        best_style = max(scores.items(), key=lambda x: x[1])

        # 如果最高分太低，返回默认风格
        if best_style[1] < 0.3:
            return "business_professional", best_style[1]

        return best_style[0], best_style[1]
    
    def _generate_style_prompt(self, features: Dict[str, Any], style_type: str) -> str:
        """生成文风提示词"""
        try:
            style_info = self.style_types.get(style_type, {})
            style_name = style_info.get("name", "通用风格")
            characteristics = style_info.get("characteristics", [])

            # 基础文风描述
            prompt_parts = [
                f"请按照{style_name}进行内容生成，具体要求如下：",
                "",
                "【文风特征】"
            ]

            for char in characteristics:
                prompt_parts.append(f"- {char}")

            # 根据具体特征数据生成个性化要求
            prompt_parts.extend(self._generate_sentence_requirements(features))
            prompt_parts.extend(self._generate_vocabulary_requirements(features))
            prompt_parts.extend(self._generate_expression_requirements(features))
            prompt_parts.extend(self._generate_organization_requirements(features))
            prompt_parts.extend(self._generate_tone_requirements(features))

            prompt_parts.append("")
            prompt_parts.append("【特别注意】")
            prompt_parts.append("- 避免AI生成痕迹，让内容自然流畅")
            prompt_parts.append("- 保持与原文档风格的一致性")
            prompt_parts.append("- 确保内容准确、逻辑清晰")

            # 添加负面示例和避免事项
            prompt_parts.extend(self._generate_avoidance_guidelines(features, style_type))

            return "\n".join(prompt_parts)

        except Exception as e:
            # 如果生成失败，返回基础提示词
            return f"请按照{style_info.get('name', '专业')}风格进行内容生成，保持语言规范、逻辑清晰、表达准确。"

    def _generate_sentence_requirements(self, features: Dict[str, Any]) -> List[str]:
        """生成句式要求"""
        requirements = ["", "【句式要求】"]

        sentence_features = features.get("sentence_structure", {})
        avg_length = sentence_features.get("average_length", 15)
        complex_ratio = sentence_features.get("complex_ratio", 0.5)

        # 句长要求
        if avg_length > 30:
            requirements.append("- 保持长句的使用，注重表达的完整性和逻辑层次")
            requirements.append("- 适当使用分号和破折号来组织复杂句式")
        elif avg_length > 20:
            requirements.append("- 适当使用长句，保持表达的完整性和逻辑性")
            requirements.append("- 长短句结合，避免句式过于单调")
        elif avg_length < 12:
            requirements.append("- 多使用短句，保持表达的简洁明了")
            requirements.append("- 避免冗长复杂的句式结构")
        else:
            requirements.append("- 长短句结合，保持节奏感和可读性")

        # 复杂度要求
        if complex_ratio > 0.7:
            requirements.append("- 保持句式的复杂性和层次感")
        elif complex_ratio < 0.3:
            requirements.append("- 使用简单直接的句式结构")

        return requirements

    def _generate_vocabulary_requirements(self, features: Dict[str, Any]) -> List[str]:
        """生成词汇要求"""
        requirements = ["", "【词汇要求】"]

        vocab_features = features.get("vocabulary_choice", {})
        formality = vocab_features.get("formality_score", 0)
        technical_density = vocab_features.get("technical_density", 0)
        modifier_usage = vocab_features.get("modifier_usage", 0)
        action_verb_ratio = vocab_features.get("action_verb_ratio", 0)

        # 正式度要求
        if formality > 15:
            requirements.append("- 使用正式、规范的书面语词汇")
            requirements.append("- 适当使用专业术语和学术表达")
        elif formality > 8:
            requirements.append("- 使用标准的书面语表达")
            requirements.append("- 避免过于口语化的词汇")
        else:
            requirements.append("- 使用通俗易懂的词汇")
            requirements.append("- 避免过于正式或生僻的表达")

        # 技术性要求
        if technical_density > 5:
            requirements.append("- 保持专业术语的准确使用")

        # 修饰词要求
        if modifier_usage > 10:
            requirements.append("- 适当使用形容词和副词进行修饰")
        elif modifier_usage < 3:
            requirements.append("- 减少不必要的修饰词，保持表达简洁")

        # 动词使用要求
        if action_verb_ratio > 20:
            requirements.append("- 多使用动作性强的动词")
            requirements.append("- 让表达更加生动有力")

        return requirements

    def _generate_expression_requirements(self, features: Dict[str, Any]) -> List[str]:
        """生成表达方式要求"""
        requirements = ["", "【表达方式】"]

        expression_features = features.get("expression_style", {})
        passive_ratio = expression_features.get("passive_active_ratio", 0)
        person_usage = expression_features.get("person_usage", {})

        # 语态要求
        if passive_ratio > 0.6:
            requirements.append("- 适当使用被动语态，体现客观性和正式性")
        elif passive_ratio > 0.3:
            requirements.append("- 主被动语态结合使用，保持表达的灵活性")
        else:
            requirements.append("- 优先使用主动语态，让表达更直接有力")

        # 人称使用要求
        first_person = person_usage.get("first_person", 0)
        if first_person > 5:
            requirements.append("- 可以适当使用第一人称，体现主观态度")
        elif first_person == 0:
            requirements.append("- 避免使用第一人称，保持客观中立")

        return requirements

    def _generate_organization_requirements(self, features: Dict[str, Any]) -> List[str]:
        """生成组织结构要求"""
        requirements = ["", "【组织结构】"]

        org_features = features.get("text_organization", {})
        connector_density = org_features.get("connector_density", 0)
        summary_usage = org_features.get("summary_usage", 0)

        # 连接词要求
        if connector_density > 8:
            requirements.append("- 使用丰富的逻辑连接词")
            requirements.append("- 保持段落间的逻辑关系清晰")
        elif connector_density > 3:
            requirements.append("- 适当使用逻辑连接词")
            requirements.append("- 注意段落间的过渡自然")
        else:
            requirements.append("- 减少机械化的过渡词使用")
            requirements.append("- 通过内容逻辑自然过渡")

        # 总结性表达
        if summary_usage > 2:
            requirements.append("- 适当使用总结性表达")
            requirements.append("- 注重内容的归纳和提炼")

        return requirements

    def _generate_tone_requirements(self, features: Dict[str, Any]) -> List[str]:
        """生成语调要求"""
        requirements = ["", "【语调风格】"]

        emotional_features = features.get("emotional_tone", {})
        professional_features = features.get("professionalism", {})

        positive_ratio = emotional_features.get("positive_ratio", 0)
        negative_ratio = emotional_features.get("negative_ratio", 0)
        intensity = emotional_features.get("intensity", 0)
        authority_indicators = professional_features.get("authority_indicators", 0)

        # 情感色彩要求
        if positive_ratio > 0.6:
            requirements.append("- 保持积极正面的语调")
        elif negative_ratio > 0.4:
            requirements.append("- 可以适当表达关切和问题意识")
        else:
            requirements.append("- 保持中性客观的语调")

        # 强度要求
        if intensity > 15:
            requirements.append("- 适当使用强调词汇，体现重要性")
        elif intensity < 5:
            requirements.append("- 保持平和的表达强度")

        # 权威性要求
        if authority_indicators > 10:
            requirements.append("- 体现专业权威性，使用准确的数据和引用")

        return requirements

    def _generate_avoidance_guidelines(self, features: Dict[str, Any], style_type: str) -> List[str]:
        """生成避免事项指导"""
        guidelines = ["", "【避免事项】"]

        # 根据风格类型添加特定避免事项
        if style_type == "business_professional":
            guidelines.extend([
                "- 避免过于感性或主观的表达",
                "- 避免使用网络流行语或俚语",
                "- 避免冗长的修饰和华丽的辞藻"
            ])
        elif style_type == "academic_research":
            guidelines.extend([
                "- 避免口语化表达和非正式用词",
                "- 避免主观臆断，确保论证严谨",
                "- 避免过于绝对的表述"
            ])
        elif style_type == "concise_practical":
            guidelines.extend([
                "- 避免冗余和重复表达",
                "- 避免过于复杂的句式结构",
                "- 避免不必要的修饰词汇"
            ])

        # 通用避免事项
        guidelines.extend([
            "- 避免明显的AI生成痕迹",
            "- 避免逻辑不清或前后矛盾",
            "- 避免语法错误和用词不当"
        ])

        return guidelines
    
    def _generate_template_id(self, document_name: str, features: Dict[str, Any]) -> str:
        """生成文风模板ID"""
        content = f"{document_name}_{json.dumps(features, sort_keys=True)}"
        return hashlib.md5(content.encode('utf-8')).hexdigest()[:12]
    
    def save_style_template(self, analysis_result: Dict[str, Any]) -> Dict[str, Any]:
        """保存文风模板"""
        try:
            template_id = analysis_result.get("template_id")
            # 自动补全template_id
            if not template_id:
                template_id = self._generate_template_id(
                    analysis_result.get("document_name", "未命名文档"),
                    analysis_result.get("style_features", {})
                )
                analysis_result["template_id"] = template_id
            
            # 保存到JSON文件
            template_file = os.path.join(self.storage_path, f"{template_id}.json")
            with open(template_file, 'w', encoding='utf-8') as f:
                json.dump(analysis_result, f, ensure_ascii=False, indent=2)
            
            # 更新模板索引
            self._update_style_template_index(template_id, analysis_result)
            
            return {
                "success": True,
                "template_id": template_id,
                "template_name": analysis_result.get("document_name", "未命名模板"),
                "style_type": analysis_result.get("style_type", "未知风格"),
                "saved_path": template_file
            }
            
        except Exception as e:
            return {"error": f"保存文风模板失败: {str(e)}"}
    
    def load_style_template(self, template_id: str) -> Dict[str, Any]:
        """加载文风模板"""
        try:
            template_file = os.path.join(self.storage_path, f"{template_id}.json")
            if not os.path.exists(template_file):
                return {"error": f"文风模板不存在: {template_id}"}
            
            with open(template_file, 'r', encoding='utf-8') as f:
                template_data = json.load(f)
            
            return template_data
            
        except Exception as e:
            return {"error": f"加载文风模板失败: {str(e)}"}
    
    def list_style_templates(self) -> List[Dict[str, Any]]:
        """列出所有文风模板"""
        try:
            index_file = os.path.join(self.storage_path, "style_template_index.json")
            if not os.path.exists(index_file):
                return []
            
            with open(index_file, 'r', encoding='utf-8') as f:
                index_data = json.load(f)
            
            return index_data.get("templates", [])
            
        except Exception as e:
            print(f"加载文风模板索引失败: {str(e)}")
            return []
    
    def _update_style_template_index(self, template_id: str, template_data: Dict[str, Any]):
        """更新文风模板索引"""
        index_file = os.path.join(self.storage_path, "style_template_index.json")
        
        # 读取现有索引
        if os.path.exists(index_file):
            with open(index_file, 'r', encoding='utf-8') as f:
                index_data = json.load(f)
        else:
            index_data = {"templates": []}
        
        # 更新或添加模板信息
        template_info = {
            "template_id": template_id,
            "name": template_data.get("document_name", "未命名模板"),
            "style_type": template_data.get("style_type", "未知风格"),
            "style_name": self.style_types.get(template_data.get("style_type", ""), {}).get("name", "未知风格"),
            "confidence_score": template_data.get("confidence_score", 0.0),
            "created_time": template_data.get("analysis_time", datetime.now().isoformat()),
            "description": f"文风模板：{template_data.get('document_name', '未命名文档')}"
        }
        
        # 检查是否已存在
        existing_index = -1
        for i, template in enumerate(index_data["templates"]):
            if template["template_id"] == template_id:
                existing_index = i
                break
        
        if existing_index >= 0:
            index_data["templates"][existing_index] = template_info
        else:
            index_data["templates"].append(template_info)
        
        # 保存索引
        with open(index_file, 'w', encoding='utf-8') as f:
            json.dump(index_data, f, ensure_ascii=False, indent=2)

    def validate_style_application(self, original_content: str, generated_content: str,
                                 target_style: str, style_features: Dict[str, Any]) -> Dict[str, Any]:
        """验证风格应用效果"""
        try:
            # 分析生成内容的风格特征
            generated_features = self._analyze_style_features(generated_content)

            # 识别生成内容的风格类型
            identified_style, confidence = self._identify_style_type(generated_content, generated_features)

            # 计算风格一致性分数
            consistency_score = self._calculate_style_consistency(
                style_features, generated_features, target_style
            )

            # 计算质量指标
            quality_metrics = self._calculate_quality_metrics(
                original_content, generated_content, generated_features
            )

            # 生成改进建议
            improvement_suggestions = self._generate_improvement_suggestions(
                target_style, generated_features, consistency_score
            )

            return {
                "success": True,
                "target_style": target_style,
                "identified_style": identified_style,
                "style_confidence": confidence,
                "consistency_score": consistency_score,
                "quality_metrics": quality_metrics,
                "generated_features": generated_features,
                "improvement_suggestions": improvement_suggestions,
                "validation_time": datetime.now().isoformat()
            }

        except Exception as e:
            return {"error": f"风格验证失败: {str(e)}"}

    def _calculate_style_consistency(self, target_features: Dict[str, Any],
                                   generated_features: Dict[str, Any],
                                   target_style: str) -> float:
        """计算风格一致性分数"""
        try:
            total_score = 0.0
            total_weight = 0.0

            # 句式结构一致性 (权重: 0.25)
            sentence_score = self._compare_sentence_features(
                target_features.get("sentence_structure", {}),
                generated_features.get("sentence_structure", {})
            )
            total_score += sentence_score * 0.25
            total_weight += 0.25

            # 词汇选择一致性 (权重: 0.3)
            vocab_score = self._compare_vocabulary_features(
                target_features.get("vocabulary_choice", {}),
                generated_features.get("vocabulary_choice", {})
            )
            total_score += vocab_score * 0.3
            total_weight += 0.3

            # 表达方式一致性 (权重: 0.2)
            expression_score = self._compare_expression_features(
                target_features.get("expression_style", {}),
                generated_features.get("expression_style", {})
            )
            total_score += expression_score * 0.2
            total_weight += 0.2

            # 组织结构一致性 (权重: 0.15)
            org_score = self._compare_organization_features(
                target_features.get("text_organization", {}),
                generated_features.get("text_organization", {})
            )
            total_score += org_score * 0.15
            total_weight += 0.15

            # 语言习惯一致性 (权重: 0.1)
            habit_score = self._compare_habit_features(
                target_features.get("language_habits", {}),
                generated_features.get("language_habits", {})
            )
            total_score += habit_score * 0.1
            total_weight += 0.1

            return total_score / total_weight if total_weight > 0 else 0.0

        except Exception as e:
            print(f"计算风格一致性时出错: {str(e)}")
            return 0.0

    def _compare_sentence_features(self, target: Dict, generated: Dict) -> float:
        """比较句式特征"""
        score = 0.0
        comparisons = 0

        # 平均句长比较
        target_length = target.get("average_length", 15)
        generated_length = generated.get("average_length", 15)
        length_diff = abs(target_length - generated_length) / max(target_length, 1)
        score += max(0, 1 - length_diff)
        comparisons += 1

        # 复杂句比例比较
        target_complex = target.get("complex_ratio", 0.5)
        generated_complex = generated.get("complex_ratio", 0.5)
        complex_diff = abs(target_complex - generated_complex)
        score += max(0, 1 - complex_diff * 2)
        comparisons += 1

        return score / comparisons if comparisons > 0 else 0.0

    def _compare_vocabulary_features(self, target: Dict, generated: Dict) -> float:
        """比较词汇特征"""
        score = 0.0
        comparisons = 0

        # 正式度比较
        target_formality = target.get("formality_score", 0)
        generated_formality = generated.get("formality_score", 0)
        if target_formality > 0:
            formality_ratio = min(generated_formality / target_formality, target_formality / generated_formality)
            score += formality_ratio
            comparisons += 1

        # 动作动词比例比较
        target_action = target.get("action_verb_ratio", 0)
        generated_action = generated.get("action_verb_ratio", 0)
        if target_action > 0:
            action_ratio = min(generated_action / target_action, target_action / generated_action)
            score += action_ratio
            comparisons += 1

        return score / comparisons if comparisons > 0 else 0.5

    def _compare_expression_features(self, target: Dict, generated: Dict) -> float:
        """比较表达方式特征"""
        score = 0.0
        comparisons = 0

        # 被动语态比例比较
        target_passive = target.get("passive_active_ratio", 0)
        generated_passive = generated.get("passive_active_ratio", 0)
        passive_diff = abs(target_passive - generated_passive)
        score += max(0, 1 - passive_diff * 2)
        comparisons += 1

        return score / comparisons if comparisons > 0 else 0.0

    def _compare_organization_features(self, target: Dict, generated: Dict) -> float:
        """比较组织结构特征"""
        score = 0.0
        comparisons = 0

        # 连接词密度比较
        target_connector = target.get("connector_density", 0)
        generated_connector = generated.get("connector_density", 0)
        if target_connector > 0:
            connector_ratio = min(generated_connector / target_connector, target_connector / generated_connector)
            score += connector_ratio
            comparisons += 1

        return score / comparisons if comparisons > 0 else 0.5

    def _compare_habit_features(self, target: Dict, generated: Dict) -> float:
        """比较语言习惯特征"""
        score = 0.0
        comparisons = 0

        # 口语化程度比较
        target_colloquial = target.get("colloquial_level", 0)
        generated_colloquial = generated.get("colloquial_level", 0)
        colloquial_diff = abs(target_colloquial - generated_colloquial)
        score += max(0, 1 - colloquial_diff * 0.1)
        comparisons += 1

        return score / comparisons if comparisons > 0 else 0.0

    def _calculate_quality_metrics(self, original_content: str, generated_content: str,
                                 generated_features: Dict[str, Any]) -> Dict[str, Any]:
        """计算质量指标"""
        try:
            metrics = {}

            # 长度比较
            original_length = len(original_content)
            generated_length = len(generated_content)
            metrics["length_ratio"] = generated_length / original_length if original_length > 0 else 0

            # 句子数量比较
            original_sentences = generated_features.get("sentence_structure", {}).get("total_sentences", 0)
            metrics["sentence_count"] = original_sentences

            # 可读性评估（基于平均句长）
            avg_length = generated_features.get("sentence_structure", {}).get("average_length", 15)
            if 15 <= avg_length <= 25:
                metrics["readability_score"] = 1.0
            elif 10 <= avg_length < 15 or 25 < avg_length <= 35:
                metrics["readability_score"] = 0.8
            else:
                metrics["readability_score"] = 0.6

            # 词汇丰富度（基于修饰词使用）
            modifier_usage = generated_features.get("vocabulary_choice", {}).get("modifier_usage", 0)
            metrics["vocabulary_richness"] = min(modifier_usage / 10, 1.0)

            # 专业性评估
            professionalism = generated_features.get("professionalism", {})
            authority_score = professionalism.get("authority_indicators", 0)
            precision_score = professionalism.get("precision_level", 0)
            metrics["professionalism_score"] = min((authority_score + precision_score) / 20, 1.0)

            return metrics

        except Exception as e:
            print(f"计算质量指标时出错: {str(e)}")
            return {"error": str(e)}

    def _generate_improvement_suggestions(self, target_style: str, generated_features: Dict[str, Any],
                                        consistency_score: float) -> List[str]:
        """生成改进建议"""
        suggestions = []

        try:
            # 基于一致性分数给出总体建议
            if consistency_score < 0.6:
                suggestions.append("整体风格一致性较低，建议重新调整生成策略")
            elif consistency_score < 0.8:
                suggestions.append("风格基本符合要求，但仍有改进空间")
            else:
                suggestions.append("风格一致性良好，符合目标要求")

            # 基于具体特征给出建议
            sentence_features = generated_features.get("sentence_structure", {})
            vocab_features = generated_features.get("vocabulary_choice", {})
            expression_features = generated_features.get("expression_style", {})

            # 句式建议
            avg_length = sentence_features.get("average_length", 15)
            if target_style == "concise_practical" and avg_length > 20:
                suggestions.append("句子偏长，建议使用更简洁的表达")
            elif target_style == "academic_research" and avg_length < 20:
                suggestions.append("句子偏短，建议增加句式的复杂性和完整性")

            # 词汇建议
            formality = vocab_features.get("formality_score", 0)
            if target_style == "formal_official" and formality < 10:
                suggestions.append("正式度不够，建议使用更规范的书面语")
            elif target_style == "narrative_descriptive" and formality > 15:
                suggestions.append("过于正式，建议使用更生动自然的表达")

            # 语态建议
            passive_ratio = expression_features.get("passive_active_ratio", 0)
            if target_style == "business_professional" and passive_ratio > 0.4:
                suggestions.append("被动语态使用过多，建议多用主动语态")

            # 如果没有具体建议，给出通用建议
            if len(suggestions) <= 1:
                suggestions.append("继续保持当前的写作风格")
                suggestions.append("注意语言的准确性和逻辑性")

        except Exception as e:
            suggestions = [f"生成改进建议时出错: {str(e)}"]

        return suggestions

    def export_styled_document(self, session_id: str) -> Dict[str, Any]:
        """
        导出应用了文风变化的文档（增强：保留结构、自动封面、目录、批注、高亮diff、原格式、元数据、变更报告）
        """
        try:
            # 1. 读取会话数据
            session_file = os.path.join(self.semantic_behavior_dir, "profiles", f"{session_id}.json")
            if not os.path.exists(session_file):
                return {"error": "会话不存在或已过期"}
            with open(session_file, 'r', encoding='utf-8') as f:
                session_data = json.load(f)
            original_content = session_data.get("original_content", "")
            suggested_changes = session_data.get("suggested_changes", [])
            doc_name = session_data.get("document_name", f"session_{session_id}")
            template_id = session_data.get("style_template_id", "template")
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            # 2. 应用变更，保留结构，按diff顺序合成
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.enum.style import WD_STYLE_TYPE
            from docx.oxml import OxmlElement
            import re
            doc = Document()
            # 2.1 封面
            doc.add_section()
            cover = doc.add_paragraph()
            run = cover.add_run(f"文风统一导出文档\n\n原文件: {doc_name}\n模板: {template_id}\n导出时间: {timestamp}")
            run.font.size = Pt(20)
            run.bold = True
            cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()
            # 2.2 目录
            toc = doc.add_paragraph("目录", style='Heading 1')
            doc.add_paragraph("（请在Word中更新目录域以显示章节）")
            doc.add_page_break()
            # 2.3 正文（保留原结构，应用变更并高亮diff）
            paragraphs = original_content.split('\n')
            for para_text in paragraphs:
                if not para_text.strip():
                    doc.add_paragraph()
                    continue
                para = doc.add_paragraph()
                applied = False
                for change in suggested_changes:
                    if change.get("status") == "accepted" and change.get("original_text") in para_text:
                        before, match, after = para_text.partition(change["original_text"])
                        if before:
                            para.add_run(before)
                        run = para.add_run(change["suggested_text"])
                        # 简化高亮设置，避免导入问题
                        run.font.color.rgb = RGBColor(255, 255, 0)  # 黄色文字
                        comment = OxmlElement('w:commentRangeStart')
                        comment.set(qn('w:id'), '0')
                        para._p.append(comment)
                        run = para.add_run(f"（风格变更：{change.get('change_type','')}，置信度{change.get('confidence',0):.2f}）")
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(255,0,0)
                        comment_end = OxmlElement('w:commentRangeEnd')
                        comment_end.set(qn('w:id'), '0')
                        para._p.append(comment_end)
                        if after:
                            para.add_run(after)
                        applied = True
                        break
                if not applied:
                    para.add_run(para_text)
            # 2.4 统一样式
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            # 简化样式设置，避免复杂的字体操作
            for style in doc.styles:
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    style.font.name = '宋体'
                    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    style.font.size = Pt(12)
            # 2.5 元数据嵌入
            core_props = doc.core_properties
            core_props.title = f"文风统一导出-{doc_name}"
            core_props.subject = f"风格模板ID: {template_id}"
            core_props.author = "智能文档助手"
            core_props.comments = f"导出时间: {timestamp}；变更数: {len([c for c in suggested_changes if c.get('status') == 'accepted'])}"
            core_props.keywords = f"style_template_id:{template_id};export_time:{timestamp}"
            # 2.6 文档末尾添加风格调整报告
            doc.add_page_break()
            report = doc.add_paragraph("风格调整报告", style='Heading 1')
            report.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph(f"原文件名: {doc_name}")
            doc.add_paragraph(f"风格模板ID: {template_id}")
            doc.add_paragraph(f"导出时间: {timestamp}")
            doc.add_paragraph(f"总变更数: {len([c for c in suggested_changes if c.get('status') == 'accepted'])}")
            for idx, change in enumerate(suggested_changes, 1):
                if change.get("status") == "accepted":
                    para = doc.add_paragraph(f"[{idx}] 类型: {change.get('change_type','')} | 置信度: {change.get('confidence',0):.2f}", style='List Number')
                    para.add_run(f"\n原文: {change.get('original_text','')}")
                    para.add_run(f"\n建议: {change.get('suggested_text','')}")
            # 2.7 文件命名
            filename = f"{doc_name}_{template_id}_{timestamp}.docx"
            import tempfile
            temp_dir = tempfile.gettempdir()
            output_path = os.path.join(temp_dir, filename)
            doc.save(output_path)
            with open(output_path, 'rb') as f:
                docx_content = f.read()
            try:
                os.remove(output_path)
            except:
                pass
            return {
                "success": True,
                "docx_content": docx_content,
                "filename": filename,
                "content_length": len(paragraphs),
                "changes_applied": len([c for c in suggested_changes if c.get("status") == "accepted"])
            }
        except Exception as e:
            return {"error": f"导出文风调整文档失败: {str(e)}"}

    def generate_style_preview(self, analysis_result: dict, style_template_id: str) -> dict:
        """
        根据分析结果和风格模板ID，对目标文档进行风格迁移预览，并返回迁移结果、diff和一致性评分。
        
        Args:
            analysis_result: 文档分析结果
            style_template_id: 风格模板ID
            
        Returns:
            风格预览结果，包含预览文本、差异、一致性评分等
        """
        try:
            # 1. 获取目标文档内容
            document_content = analysis_result.get("document_content") or analysis_result.get("text") or ""
            document_name = analysis_result.get("document_name", "未命名文档")

            # 2. 加载风格模板（增强版，带回退机制）
            template = self._load_template_with_fallback(style_template_id, analysis_result)
            if not template:
                return {"error": "无法加载风格模板且无法创建默认模板", "success": False}

            target_style_type = template.get("style_type") or template.get("styleType") or "business_professional"

            # 3. 选择增强处理器或规则引擎进行风格迁移
            if self.use_enhanced_features and self.enhanced_processor and hasattr(self.enhanced_processor, "transfer_style"):
                # 使用增强风格迁移
                transfer_result = self.enhanced_processor.transfer_style(
                    document_content=document_content,
                    target_style_type=target_style_type,
                    template=template,
                    original_analysis=analysis_result
                )
            else:
                # 基础规则迁移（可自定义更复杂逻辑）
                transfer_result = {
                    "rewritten_text": document_content,  # 这里可调用基础迁移方法
                    "style_changes": [],
                    "success": True,
                    "note": "未启用增强风格迁移，返回原文"
                }

            # 4. 计算风格一致性分数
            consistency_score = 0.0
            if "rewritten_text" in transfer_result:
                generated_features = self._analyze_style_features(transfer_result["rewritten_text"])
                target_features = template.get("style_features", {})
                consistency_score = self._calculate_style_consistency(target_features, generated_features, target_style_type)

            # 5. 生成diff（可选：对比迁移前后文本差异）
            diff = []
            if "rewritten_text" in transfer_result and transfer_result["rewritten_text"] != document_content:
                import difflib
                diff = list(difflib.unified_diff(
                    document_content.splitlines(), 
                    transfer_result["rewritten_text"].splitlines(),
                    lineterm=""
                ))

            # 6. 返回结构化结果
            return {
                "success": True,
                "preview_text": transfer_result.get("rewritten_text", document_content),
                "diff": diff,
                "consistency_score": consistency_score,
                "style_changes": transfer_result.get("style_changes", []),
                "note": transfer_result.get("note", ""),
                "template_id": style_template_id,
                "target_style_type": target_style_type
            }
        except Exception as e:
            return {"error": f"风格迁移预览失败: {str(e)}", "success": False}

    def _load_template_with_fallback(self, template_id: str, analysis_result: dict = None) -> Optional[dict]:
        """
        带回退机制的模板加载
        
        Args:
            template_id: 模板ID
            analysis_result: 分析结果（用于创建默认模板）
            
        Returns:
            模板数据或None
        """
        # 1. 尝试加载指定模板
        template = self.load_style_template(template_id)
        if template and "error" not in template:
            return template
        
        # 2. 如果模板不存在且有分析结果，尝试保存分析结果作为模板
        if analysis_result:
            print(f"模板 {template_id} 不存在，尝试保存分析结果作为模板")
            save_result = self.save_style_template(analysis_result)
            if save_result.get("success"):
                print(f"成功保存模板: {save_result.get('template_id')}")
                return analysis_result  # 使用分析结果作为模板
        
        # 3. 尝试查找相似模板
        similar_template = self._find_similar_template(template_id)
        if similar_template:
            print(f"使用相似模板: {similar_template.get('template_id', 'unknown')}")
            return similar_template
        
        # 4. 使用默认模板
        default_template = self._get_default_template()
        if default_template:
            print("使用默认模板")
            return default_template
        
        # 5. 创建基础模板
        basic_template = self._create_basic_template()
        if basic_template:
            print("创建基础模板")
            return basic_template
        
        return None

    def _find_similar_template(self, template_id: str) -> Optional[dict]:
        """
        查找相似模板
        
        Args:
            template_id: 目标模板ID
            
        Returns:
            相似模板或None
        """
        try:
            templates = self.list_style_templates()
            if not templates:
                return None
            
            # 简单的相似性查找（基于ID前缀）
            template_prefix = template_id[:8] if len(template_id) >= 8 else template_id
            
            for template in templates:
                current_id = template.get("template_id", "")
                if current_id.startswith(template_prefix):
                    return self.load_style_template(current_id)
            
            # 如果没有找到相似ID，返回第一个可用模板
            if templates:
                first_template_id = templates[0].get("template_id")
                if first_template_id:
                    return self.load_style_template(first_template_id)
            
            return None
            
        except Exception as e:
            print(f"查找相似模板失败: {str(e)}")
            return None

    def _get_default_template(self) -> Optional[dict]:
        """
        获取默认模板
        
        Returns:
            默认模板或None
        """
        try:
            # 创建默认的商务专业风格模板
            default_template = {
                "template_id": "default_business_professional",
                "document_name": "默认商务专业模板",
                "style_type": "business_professional",
                "style_features": {
                    "formality": 0.8,
                    "technicality": 0.6,
                    "objectivity": 0.7,
                    "conciseness": 0.5,
                    "professionalism": 0.8
                },
                "style_prompt": """
请使用商务专业的写作风格，要求：
1. 语言正式、客观，避免主观色彩
2. 使用专业术语，但保持易懂
3. 结构清晰，逻辑严密
4. 避免口语化表达
5. 保持简洁明了
                """.strip(),
                "confidence_score": 0.9,
                "created_time": datetime.now().isoformat()
            }
            
            return default_template
            
        except Exception as e:
            print(f"获取默认模板失败: {str(e)}")
            return None

    def _create_basic_template(self) -> Optional[dict]:
        """
        创建基础模板
        
        Returns:
            基础模板或None
        """
        try:
            basic_template = {
                "template_id": "basic_template",
                "document_name": "基础模板",
                "style_type": "general",
                "style_features": {
                    "formality": 0.5,
                    "technicality": 0.3,
                    "objectivity": 0.6,
                    "conciseness": 0.5,
                    "professionalism": 0.5
                },
                "style_prompt": """
请使用通用的写作风格，要求：
1. 语言清晰易懂
2. 结构合理
3. 表达准确
4. 避免过于复杂的句式
                """.strip(),
                "confidence_score": 0.7,
                "created_time": datetime.now().isoformat()
            }
            
            return basic_template
            
        except Exception as e:
            print(f"创建基础模板失败: {str(e)}")
            return None

    def handle_batch_style_changes(self, session_id: str, action: str = "accept_all") -> dict:
        """
        批量处理风格变化（如全部接受/全部拒绝）- 真实实现
        """
        try:
            # 1. 从session文件中读取真实的原始文档内容和风格模板
            session_file = os.path.join(self.semantic_behavior_dir, "profiles", f"{session_id}.json")
            if not os.path.exists(session_file):
                return {
                    "success": False,
                    "error": f"会话文件不存在: {session_id}"
                }
            
            with open(session_file, 'r', encoding='utf-8') as f:
                session_data = json.load(f)
            
            # 获取真实的原始文档内容
            original_content = session_data.get("original_content", "")
            document_name = session_data.get("document_name", f"session_{session_id}")
            style_template_id = session_data.get("style_template_id", "")
            
            if not original_content:
                return {
                    "success": False,
                    "error": "session文件中没有找到原始文档内容"
                }
            
            if not style_template_id:
                return {
                    "success": False,
                    "error": "session文件中没有找到风格模板ID"
                }
            
            # 2. 加载风格模板
            template = self.load_style_template(style_template_id)
            if not template:
                template = {
                    "style_type": "business_professional",
                    "style_features": {
                        "formality": 0.8,
                        "technicality": 0.6,
                        "objectivity": 0.7,
                        "conciseness": 0.5
                    }
                }
            
            # 3. 根据action执行真实的风格迁移
            if action == "accept_all":
                # 使用LLM进行真实的风格迁移
                migrated_content = self._perform_real_style_migration(
                    original_content, template, style_template_id
                )
                
                # 生成真实的变更记录
                suggested_changes = self._generate_real_changes(
                    original_content, migrated_content, template
                )
                
                # 标记所有变更为已接受
                for change in suggested_changes:
                    change["status"] = "accepted"
                    
            elif action == "reject_all":
                # 拒绝所有变更，保持原文
                migrated_content = original_content
                suggested_changes = []
                
            else:
                return {
                    "success": False,
                    "error": f"不支持的操作: {action}"
                }
            
            # 4. 更新session数据
            session_data.update({
                "action": action,
                "migrated_content": migrated_content,
                "suggested_changes": suggested_changes,
                "target_style": template.get("style_type", "business_professional"),
                "last_updated": datetime.now().isoformat()
            })
            
            # 保存更新后的session文件
            with open(session_file, 'w', encoding='utf-8') as f:
                json.dump(session_data, f, ensure_ascii=False, indent=2)
            
            print(f"会话数据已保存: {session_file}")
            
            return {
                "success": True,
                "session_id": session_id,
                "action": action,
                "message": f"批量{action}风格变化已处理",
                "changes_count": len(suggested_changes),
                "accepted_count": len([c for c in suggested_changes if c.get("status") == "accepted"]),
                "migrated_content_length": len(migrated_content)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"批量处理风格变化失败: {str(e)}"
            }
    
    def _perform_real_style_migration(self, original_content: str, template: dict, template_id: str) -> str:
        """
        执行真实的风格迁移
        """
        try:
            # 1. 构建风格迁移提示词
            target_style = template.get("style_type", "business_professional")
            style_features = template.get("style_features", {})
            
            prompt = self._build_style_migration_prompt(
                original_content, target_style, style_features
            )
            
            # 2. 调用LLM进行风格迁移
            if self.llm_client:
                response = self.llm_client.generate_text(prompt, max_tokens=2000)
                if response and "content" in response:
                    migrated_content = response["content"].strip()
                    # 清理可能的markdown格式
                    if migrated_content.startswith("```"):
                        lines = migrated_content.split('\n')
                        if len(lines) > 2:
                            migrated_content = '\n'.join(lines[1:-1])
                    return migrated_content
            
            # 3. 如果LLM调用失败，使用规则基础迁移
            return self._rule_based_style_migration(original_content, target_style, style_features)
            
        except Exception as e:
            print(f"风格迁移失败: {e}")
            return original_content
    
    def _build_style_migration_prompt(self, content: str, target_style: str, style_features: dict) -> str:
        """
        构建风格迁移的LLM提示词
        """
        formality = style_features.get("formality", 0.5)
        technicality = style_features.get("technicality", 0.5)
        objectivity = style_features.get("objectivity", 0.5)
        conciseness = style_features.get("conciseness", 0.5)
        
        prompt = f"""
请将以下文档的风格调整为{target_style}风格，要求：

风格特征：
- 正式程度：{formality:.1f}（0-1，越高越正式）
- 技术性：{technicality:.1f}（0-1，越高越技术化）
- 客观性：{objectivity:.1f}（0-1，越高越客观）
- 简洁性：{conciseness:.1f}（0-1，越高越简洁）

调整要求：
1. 保持原文的核心信息和逻辑结构
2. 调整词汇选择，使其符合目标风格
3. 优化句式结构，提高表达的专业性
4. 确保语言的一致性和连贯性

原文：
{content}

请直接返回调整后的文本，不要添加任何解释或标记。
"""
        return prompt
    
    def _rule_based_style_migration(self, content: str, target_style: str, style_features: dict) -> str:
        """
        基于规则的风格迁移（LLM不可用时的备选方案）
        """
        migrated_content = content
        
        # 根据目标风格应用不同的规则
        if target_style == "business_professional":
            # 商务专业风格调整
            replacements = {
                "我觉得": "我认为",
                "挺好的": "较为理想",
                "应该可以": "能够",
                "解决问题": "解决相关问题",
                "用了": "采用了",
                "算了一下": "进行了分析",
                "总的来说": "综上所述",
                "不错": "良好",
                "应该能用": "具备可行性"
            }
            
            for old, new in replacements.items():
                migrated_content = migrated_content.replace(old, new)
        
        elif target_style == "academic":
            # 学术风格调整
            replacements = {
                "我觉得": "研究表明",
                "挺好的": "具有积极效果",
                "应该可以": "能够有效",
                "解决问题": "解决相关问题",
                "用了": "采用了",
                "算了一下": "进行了统计分析",
                "总的来说": "综上所述",
                "不错": "表现良好",
                "应该能用": "具备应用价值"
            }
            
            for old, new in replacements.items():
                migrated_content = migrated_content.replace(old, new)
        
        return migrated_content
    
    def _generate_real_changes(self, original_content: str, migrated_content: str, template: dict) -> list:
        """
        生成真实的变更记录
        """
        changes = []
        
        # 使用difflib生成差异
        import difflib
        
        # 简单的词级别差异检测
        original_words = original_content.split()
        migrated_words = migrated_content.split()
        
        matcher = difflib.SequenceMatcher(None, original_words, migrated_words)
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'replace':
                original_text = ' '.join(original_words[i1:i2])
                suggested_text = ' '.join(migrated_words[j1:j2])
                
                if original_text.strip() and suggested_text.strip():
                    change = {
                        "original_text": original_text,
                        "suggested_text": suggested_text,
                        "status": "accepted",
                        "change_type": self._classify_change_type(original_text, suggested_text),
                        "confidence": 0.85,
                        "position": i1
                    }
                    changes.append(change)
        
        return changes
    
    def _classify_change_type(self, original: str, suggested: str) -> str:
        """
        分类变更类型
        """
        if len(original) < len(suggested):
            return "vocabulary_improvement"
        elif len(original) > len(suggested):
            return "conciseness_improvement"
        else:
            return "style_alignment"

    def handle_style_change(self, session_id: str, change_id: str, action: str) -> dict:
        """
        处理单个风格变化（接受/拒绝）
        
        Args:
            session_id: 会话ID
            change_id: 变化ID
            action: 操作类型 ('accept' 或 'reject')
            
        Returns:
            Dict: 处理结果
        """
        try:
            # 1. 验证参数
            if action not in ['accept', 'reject']:
                return {
                    "success": False,
                    "error": f"不支持的操作: {action}，必须是 'accept' 或 'reject'"
                }
            
            # 2. 从session文件中读取数据
            session_file = os.path.join(self.semantic_behavior_dir, "profiles", f"{session_id}.json")
            if not os.path.exists(session_file):
                return {
                    "success": False,
                    "error": f"会话文件不存在: {session_id}"
                }
            
            with open(session_file, 'r', encoding='utf-8') as f:
                session_data = json.load(f)
            
            # 3. 获取原始内容和建议的变化
            original_content = session_data.get("original_content", "")
            suggested_changes = session_data.get("suggested_changes", [])
            
            if not original_content:
                return {
                    "success": False,
                    "error": "session文件中没有找到原始文档内容"
                }
            
            # 4. 查找指定的变化
            target_change = None
            for change in suggested_changes:
                if change.get("change_id") == change_id:
                    target_change = change
                    break
            
            if not target_change:
                return {
                    "success": False,
                    "error": f"未找到指定的变化: {change_id}"
                }
            
            # 5. 更新变化状态
            target_change["status"] = action
            target_change["action_time"] = datetime.now().isoformat()
            
            # 6. 生成更新后的预览内容
            updated_preview = self._generate_updated_preview(original_content, suggested_changes)
            
            # 7. 更新session数据
            session_data.update({
                "suggested_changes": suggested_changes,
                "updated_preview": updated_preview,
                "last_updated": datetime.now().isoformat()
            })
            
            # 8. 保存更新后的session文件
            with open(session_file, 'w', encoding='utf-8') as f:
                json.dump(session_data, f, ensure_ascii=False, indent=2)
            
            print(f"单个风格变化已处理: {change_id} -> {action}")
            
            return {
                "success": True,
                "change_id": change_id,
                "action": action,
                "message": f"变化 {change_id} 已{action}",
                "updated_preview": {
                    "content": updated_preview,
                    "accepted_changes": len([c for c in suggested_changes if c.get("status") == "accepted"]),
                    "rejected_changes": len([c for c in suggested_changes if c.get("status") == "rejected"]),
                    "pending_changes": len([c for c in suggested_changes if c.get("status") == "pending"])
                }
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"处理单个风格变化失败: {str(e)}"
            }
    
    def _generate_updated_preview(self, original_content: str, suggested_changes: list) -> str:
        """
        根据已接受的变化生成更新后的预览内容
        
        Args:
            original_content: 原始内容
            suggested_changes: 建议的变化列表
            
        Returns:
            str: 更新后的预览内容
        """
        try:
            # 按位置排序变化，确保按顺序应用
            accepted_changes = [c for c in suggested_changes if c.get("status") == "accepted"]
            accepted_changes.sort(key=lambda x: x.get("position", {}).get("start", 0))
            
            # 从后往前应用变化，避免位置偏移
            updated_content = original_content
            offset = 0
            
            for change in accepted_changes:
                position = change.get("position", {})
                start = position.get("start", 0) + offset
                end = position.get("end", 0) + offset
                suggested_text = change.get("suggested_text", "")
                
                # 应用变化
                if start < len(updated_content) and end <= len(updated_content):
                    updated_content = updated_content[:start] + suggested_text + updated_content[end:]
                    # 更新偏移量
                    offset += len(suggested_text) - (end - start)
            
            return updated_content
            
        except Exception as e:
            print(f"生成更新预览失败: {e}")
            return original_content

    def apply_style_changes(self, session_id: str, changes: List[Dict[str, Any]]) -> dict:
        """
        应用风格变化到文档
        
        Args:
            session_id: 会话ID
            changes: 要应用的变化列表
            
        Returns:
            Dict: 应用结果
        """
        try:
            # 1. 验证参数
            if not changes:
                return {
                    "success": False,
                    "error": "没有提供要应用的变化"
                }
            
            # 2. 从session文件中读取数据
            session_file = os.path.join(self.semantic_behavior_dir, "profiles", f"{session_id}.json")
            if not os.path.exists(session_file):
                return {
                    "success": False,
                    "error": f"会话文件不存在: {session_id}"
                }
            
            with open(session_file, 'r', encoding='utf-8') as f:
                session_data = json.load(f)
            
            # 3. 获取原始文档内容
            original_content = session_data.get("original_content", "")
            if not original_content:
                return {
                    "success": False,
                    "error": "会话中没有原始文档内容"
                }
            
            # 4. 应用变化
            updated_content = original_content
            applied_changes = []
            failed_changes = []
            
            for change in changes:
                change_id = change.get("id")
                change_type = change.get("type")
                change_data = change.get("data", {})
                
                try:
                    if change_type == "text_replacement":
                        result = self._apply_text_replacement(updated_content, change_data)
                        if result["success"]:
                            updated_content = result["updated_content"]
                            applied_changes.append({
                                "id": change_id,
                                "type": change_type,
                                "status": "applied"
                            })
                        else:
                            failed_changes.append({
                                "id": change_id,
                                "type": change_type,
                                "error": result["error"]
                            })
                    
                    elif change_type == "format_change":
                        result = self._apply_format_change(updated_content, change_data)
                        if result["success"]:
                            updated_content = result["updated_content"]
                            applied_changes.append({
                                "id": change_id,
                                "type": change_type,
                                "status": "applied"
                            })
                        else:
                            failed_changes.append({
                                "id": change_id,
                                "type": change_type,
                                "error": result["error"]
                            })
                    
                    elif change_type == "style_adjustment":
                        result = self._apply_style_adjustment(updated_content, change_data)
                        if result["success"]:
                            updated_content = result["updated_content"]
                            applied_changes.append({
                                "id": change_id,
                                "type": change_type,
                                "status": "applied"
                            })
                        else:
                            failed_changes.append({
                                "id": change_id,
                                "type": change_type,
                                "error": result["error"]
                            })
                    
                    else:
                        failed_changes.append({
                            "id": change_id,
                            "type": change_type,
                            "error": f"不支持的变化类型: {change_type}"
                        })
                
                except Exception as e:
                    failed_changes.append({
                        "id": change_id,
                        "type": change_type,
                        "error": f"应用变化时发生错误: {str(e)}"
                    })
            
            # 5. 更新session数据
            session_data["updated_content"] = updated_content
            session_data["applied_changes"] = applied_changes
            session_data["failed_changes"] = failed_changes
            session_data["last_updated"] = datetime.now().isoformat()
            
            # 6. 保存更新后的session
            with open(session_file, 'w', encoding='utf-8') as f:
                json.dump(session_data, f, ensure_ascii=False, indent=2)
            
            # 7. 生成应用报告
            application_report = self._generate_application_report(applied_changes, failed_changes, changes)
            
            return {
                "success": True,
                "session_id": session_id,
                "applied_changes_count": len(applied_changes),
                "failed_changes_count": len(failed_changes),
                "total_changes_count": len(changes),
                "application_report": application_report,
                "updated_content_preview": updated_content[:500] + "..." if len(updated_content) > 500 else updated_content,
                "applied_at": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"应用风格变化失败: {str(e)}"
            }
    
    def _apply_text_replacement(self, content: str, change_data: Dict[str, Any]) -> Dict[str, Any]:
        """应用文本替换"""
        try:
            old_text = change_data.get("old_text", "")
            new_text = change_data.get("new_text", "")
            
            if not old_text:
                return {
                    "success": False,
                    "error": "缺少要替换的文本"
                }
            
            if old_text not in content:
                return {
                    "success": False,
                    "error": f"在文档中找不到文本: {old_text[:50]}..."
                }
            
            updated_content = content.replace(old_text, new_text)
            
            return {
                "success": True,
                "updated_content": updated_content,
                "replacement_count": content.count(old_text)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"文本替换失败: {str(e)}"
            }
    
    def _apply_format_change(self, content: str, change_data: Dict[str, Any]) -> Dict[str, Any]:
        """应用格式变化"""
        try:
            format_type = change_data.get("format_type")
            target_text = change_data.get("target_text", "")
            new_format = change_data.get("new_format", {})
            
            if not target_text or not format_type:
                return {
                    "success": False,
                    "error": "缺少格式变化的目标文本或格式类型"
                }
            
            if target_text not in content:
                return {
                    "success": False,
                    "error": f"在文档中找不到目标文本: {target_text[:50]}..."
                }
            
            # 根据格式类型应用不同的格式
            if format_type == "bold":
                updated_content = content.replace(target_text, f"**{target_text}**")
            elif format_type == "italic":
                updated_content = content.replace(target_text, f"*{target_text}*")
            elif format_type == "underline":
                updated_content = content.replace(target_text, f"__{target_text}__")
            elif format_type == "highlight":
                updated_content = content.replace(target_text, f"=={target_text}==")
            else:
                return {
                    "success": False,
                    "error": f"不支持的格式类型: {format_type}"
                }
            
            return {
                "success": True,
                "updated_content": updated_content,
                "format_type": format_type
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"格式变化失败: {str(e)}"
            }
    
    def _apply_style_adjustment(self, content: str, change_data: Dict[str, Any]) -> Dict[str, Any]:
        """应用风格调整"""
        try:
            adjustment_type = change_data.get("adjustment_type")
            target_section = change_data.get("target_section", "")
            adjustment_data = change_data.get("adjustment_data", {})
            
            if not adjustment_type:
                return {
                    "success": False,
                    "error": "缺少风格调整类型"
                }
            
            # 根据调整类型应用不同的风格调整
            if adjustment_type == "tone_adjustment":
                # 调整语气
                tone = adjustment_data.get("tone", "neutral")
                updated_content = self._adjust_tone(content, tone)
            elif adjustment_type == "complexity_adjustment":
                # 调整复杂度
                complexity = adjustment_data.get("complexity", "medium")
                updated_content = self._adjust_complexity(content, complexity)
            elif adjustment_type == "formality_adjustment":
                # 调整正式程度
                formality = adjustment_data.get("formality", "neutral")
                updated_content = self._adjust_formality(content, formality)
            else:
                return {
                    "success": False,
                    "error": f"不支持的风格调整类型: {adjustment_type}"
                }
            
            return {
                "success": True,
                "updated_content": updated_content,
                "adjustment_type": adjustment_type
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"风格调整失败: {str(e)}"
            }
    
    def _adjust_tone(self, content: str, tone: str) -> str:
        """调整语气"""
        # 简化的语气调整实现
        if tone == "formal":
            # 增加正式性
            content = content.replace("我们", "本机构")
            content = content.replace("你们", "贵方")
        elif tone == "casual":
            # 增加随意性
            content = content.replace("本机构", "我们")
            content = content.replace("贵方", "你们")
        
        return content
    
    def _adjust_complexity(self, content: str, complexity: str) -> str:
        """调整复杂度"""
        # 简化的复杂度调整实现
        if complexity == "simple":
            # 简化表达
            content = content.replace("因此", "所以")
            content = content.replace("然而", "但是")
        elif complexity == "complex":
            # 增加复杂度
            content = content.replace("所以", "因此")
            content = content.replace("但是", "然而")
        
        return content
    
    def _adjust_formality(self, content: str, formality: str) -> str:
        """调整正式程度"""
        # 简化的正式程度调整实现
        if formality == "formal":
            # 增加正式性
            content = content.replace("这个", "该")
            content = content.replace("那个", "该")
        elif formality == "informal":
            # 减少正式性
            content = content.replace("该", "这个")
        
        return content
    
    def _generate_application_report(self, applied_changes: List[Dict[str, Any]], 
                                   failed_changes: List[Dict[str, Any]], 
                                   total_changes: List[Dict[str, Any]]) -> Dict[str, Any]:
        """生成应用报告"""
        try:
            # 统计信息
            total_count = len(total_changes)
            applied_count = len(applied_changes)
            failed_count = len(failed_changes)
            success_rate = applied_count / total_count if total_count > 0 else 0
            
            # 按类型统计
            type_stats = {}
            for change in total_changes:
                change_type = change.get("type", "unknown")
                if change_type not in type_stats:
                    type_stats[change_type] = {"total": 0, "applied": 0, "failed": 0}
                type_stats[change_type]["total"] += 1
            
            for change in applied_changes:
                change_type = change.get("type", "unknown")
                if change_type in type_stats:
                    type_stats[change_type]["applied"] += 1
            
            for change in failed_changes:
                change_type = change.get("type", "unknown")
                if change_type in type_stats:
                    type_stats[change_type]["failed"] += 1
            
            return {
                "summary": {
                    "total_changes": total_count,
                    "applied_changes": applied_count,
                    "failed_changes": failed_count,
                    "success_rate": success_rate
                },
                "type_statistics": type_stats,
                "applied_changes": applied_changes,
                "failed_changes": failed_changes,
                "recommendations": self._generate_application_recommendations(applied_changes, failed_changes)
            }
            
        except Exception as e:
            return {"error": f"生成应用报告失败: {str(e)}"}
    
    def _generate_application_recommendations(self, applied_changes: List[Dict[str, Any]], 
                                            failed_changes: List[Dict[str, Any]]) -> List[str]:
        """生成应用建议"""
        recommendations = []
        
        if failed_changes:
            recommendations.append(f"有 {len(failed_changes)} 个变化应用失败，建议检查失败原因")
        
        if applied_changes:
            recommendations.append(f"成功应用了 {len(applied_changes)} 个变化")
        
        # 检查特定类型的失败
        text_replacement_failures = [c for c in failed_changes if c.get("type") == "text_replacement"]
        if text_replacement_failures:
            recommendations.append("文本替换失败较多，建议检查目标文本是否存在")
        
        format_failures = [c for c in failed_changes if c.get("type") == "format_change"]
        if format_failures:
            recommendations.append("格式变化失败较多，建议检查格式类型是否支持")
        
        if not recommendations:
            recommendations.append("所有变化都已成功应用")
        
        return recommendations

    def handle_batch_style_changes(self, session_id: str, changes: List[Dict[str, Any]]) -> dict:
        """
        批量处理风格变化（如全部接受/全部拒绝）- 真实实现
        """
        try:
            # 1. 从session文件中读取真实的原始文档内容和风格模板
            session_file = os.path.join(self.semantic_behavior_dir, "profiles", f"{session_id}.json")
            if not os.path.exists(session_file):
                return {
                    "success": False,
                    "error": f"会话文件不存在: {session_id}"
                }
            
            with open(session_file, 'r', encoding='utf-8') as f:
                session_data = json.load(f)
            
            # 获取真实的原始文档内容
            original_content = session_data.get("original_content", "")
            document_name = session_data.get("document_name", f"session_{session_id}")
            style_template_id = session_data.get("style_template_id", "")
            
            if not original_content:
                return {
                    "success": False,
                    "error": "session文件中没有找到原始文档内容"
                }
            
            if not style_template_id:
                return {
                    "success": False,
                    "error": "session文件中没有找到风格模板ID"
                }
            
            # 2. 加载风格模板
            template = self.load_style_template(style_template_id)
            if not template:
                template = {
                    "style_type": "business_professional",
                    "style_features": {
                        "formality": 0.8,
                        "technicality": 0.6,
                        "objectivity": 0.7,
                        "conciseness": 0.5
                    }
                }
            
            # 3. 根据action执行真实的风格迁移
            if action == "accept_all":
                # 使用LLM进行真实的风格迁移
                migrated_content = self._perform_real_style_migration(
                    original_content, template, style_template_id
                )
                
                # 生成真实的变更记录
                suggested_changes = self._generate_real_changes(
                    original_content, migrated_content, template
                )
                
                # 标记所有变更为已接受
                for change in suggested_changes:
                    change["status"] = "accepted"
                    
            elif action == "reject_all":
                # 拒绝所有变更，保持原文
                migrated_content = original_content
                suggested_changes = []
                
            else:
                return {
                    "success": False,
                    "error": f"不支持的操作: {action}"
                }
            
            # 4. 更新session数据
            session_data.update({
                "action": action,
                "migrated_content": migrated_content,
                "suggested_changes": suggested_changes,
                "target_style": template.get("style_type", "business_professional"),
                "last_updated": datetime.now().isoformat()
            })
            
            # 保存更新后的session文件
            with open(session_file, 'w', encoding='utf-8') as f:
                json.dump(session_data, f, ensure_ascii=False, indent=2)
            
            print(f"会话数据已保存: {session_file}")
            
            return {
                "success": True,
                "session_id": session_id,
                "action": action,
                "message": f"批量{action}风格变化已处理",
                "changes_count": len(suggested_changes),
                "accepted_count": len([c for c in suggested_changes if c.get("status") == "accepted"]),
                "migrated_content_length": len(migrated_content)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"批量处理风格变化失败: {str(e)}"
            }
    
    def _perform_real_style_migration(self, original_content: str, template: dict, template_id: str) -> str:
        """
        执行真实的风格迁移
        """
        try:
            # 1. 构建风格迁移提示词
            target_style = template.get("style_type", "business_professional")
            style_features = template.get("style_features", {})
            
            prompt = self._build_style_migration_prompt(
                original_content, target_style, style_features
            )
            
            # 2. 调用LLM进行风格迁移
            if self.llm_client:
                response = self.llm_client.generate_text(prompt, max_tokens=2000)
                if response and "content" in response:
                    migrated_content = response["content"].strip()
                    # 清理可能的markdown格式
                    if migrated_content.startswith("```"):
                        lines = migrated_content.split('\n')
                        if len(lines) > 2:
                            migrated_content = '\n'.join(lines[1:-1])
                    return migrated_content
            
            # 3. 如果LLM调用失败，使用规则基础迁移
            return self._rule_based_style_migration(original_content, target_style, style_features)
            
        except Exception as e:
            print(f"风格迁移失败: {e}")
            return original_content
    
    def _build_style_migration_prompt(self, content: str, target_style: str, style_features: dict) -> str:
        """
        构建风格迁移的LLM提示词
        """
        formality = style_features.get("formality", 0.5)
        technicality = style_features.get("technicality", 0.5)
        objectivity = style_features.get("objectivity", 0.5)
        conciseness = style_features.get("conciseness", 0.5)
        
        prompt = f"""
请将以下文档的风格调整为{target_style}风格，要求：

风格特征：
- 正式程度：{formality:.1f}（0-1，越高越正式）
- 技术性：{technicality:.1f}（0-1，越高越技术化）
- 客观性：{objectivity:.1f}（0-1，越高越客观）
- 简洁性：{conciseness:.1f}（0-1，越高越简洁）

调整要求：
1. 保持原文的核心信息和逻辑结构
2. 调整词汇选择，使其符合目标风格
3. 优化句式结构，提高表达的专业性
4. 确保语言的一致性和连贯性

原文：
{content}

请直接返回调整后的文本，不要添加任何解释或标记。
"""
        return prompt
    
    def _rule_based_style_migration(self, content: str, target_style: str, style_features: dict) -> str:
        """
        基于规则的风格迁移（LLM不可用时的备选方案）
        """
        migrated_content = content
        
        # 根据目标风格应用不同的规则
        if target_style == "business_professional":
            # 商务专业风格调整
            replacements = {
                "我觉得": "我认为",
                "挺好的": "较为理想",
                "应该可以": "能够",
                "解决问题": "解决相关问题",
                "用了": "采用了",
                "算了一下": "进行了分析",
                "总的来说": "综上所述",
                "不错": "良好",
                "应该能用": "具备可行性"
            }
            
            for old, new in replacements.items():
                migrated_content = migrated_content.replace(old, new)
        
        elif target_style == "academic":
            # 学术风格调整
            replacements = {
                "我觉得": "研究表明",
                "挺好的": "具有积极效果",
                "应该可以": "能够有效",
                "解决问题": "解决相关问题",
                "用了": "采用了",
                "算了一下": "进行了统计分析",
                "总的来说": "综上所述",
                "不错": "表现良好",
                "应该能用": "具备应用价值"
            }
            
            for old, new in replacements.items():
                migrated_content = migrated_content.replace(old, new)
        
        return migrated_content
    
    def _generate_real_changes(self, original_content: str, migrated_content: str, template: dict) -> list:
        """
        生成真实的变更记录
        """
        changes = []
        
        # 使用difflib生成差异
        import difflib
        
        # 简单的词级别差异检测
        original_words = original_content.split()
        migrated_words = migrated_content.split()
        
        matcher = difflib.SequenceMatcher(None, original_words, migrated_words)
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'replace':
                original_text = ' '.join(original_words[i1:i2])
                suggested_text = ' '.join(migrated_words[j1:j2])
                
                if original_text.strip() and suggested_text.strip():
                    change = {
                        "original_text": original_text,
                        "suggested_text": suggested_text,
                        "status": "accepted",
                        "change_type": self._classify_change_type(original_text, suggested_text),
                        "confidence": 0.85,
                        "position": i1
                    }
                    changes.append(change)
        
        return changes
    
    def _classify_change_type(self, original: str, suggested: str) -> str:
        """
        分类变更类型
        """
        if len(original) < len(suggested):
            return "vocabulary_improvement"
        elif len(original) > len(suggested):
            return "conciseness_improvement"
        else:
            return "style_alignment"

    def handle_style_change(self, session_id: str, change_id: str, action: str) -> dict:
        """
        处理单个风格变化（接受/拒绝）
        
        Args:
            session_id: 会话ID
            change_id: 变化ID
            action: 操作类型 ('accept' 或 'reject')
            
        Returns:
            Dict: 处理结果
        """
        try:
            # 1. 验证参数
            if action not in ['accept', 'reject']:
                return {
                    "success": False,
                    "error": f"不支持的操作: {action}，必须是 'accept' 或 'reject'"
                }
            
            # 2. 从session文件中读取数据
            session_file = os.path.join(self.semantic_behavior_dir, "profiles", f"{session_id}.json")
            if not os.path.exists(session_file):
                return {
                    "success": False,
                    "error": f"会话文件不存在: {session_id}"
                }
            
            with open(session_file, 'r', encoding='utf-8') as f:
                session_data = json.load(f)
            
            # 3. 获取原始内容和建议的变化
            original_content = session_data.get("original_content", "")
            suggested_changes = session_data.get("suggested_changes", [])
            
            if not original_content:
                return {
                    "success": False,
                    "error": "session文件中没有找到原始文档内容"
                }
            
            # 4. 查找指定的变化
            target_change = None
            for change in suggested_changes:
                if change.get("change_id") == change_id:
                    target_change = change
                    break
            
            if not target_change:
                return {
                    "success": False,
                    "error": f"未找到指定的变化: {change_id}"
                }
            
            # 5. 更新变化状态
            target_change["status"] = action
            target_change["action_time"] = datetime.now().isoformat()
            
            # 6. 生成更新后的预览内容
            updated_preview = self._generate_updated_preview(original_content, suggested_changes)
            
            # 7. 更新session数据
            session_data.update({
                "suggested_changes": suggested_changes,
                "updated_preview": updated_preview,
                "last_updated": datetime.now().isoformat()
            })
            
            # 8. 保存更新后的session文件
            with open(session_file, 'w', encoding='utf-8') as f:
                json.dump(session_data, f, ensure_ascii=False, indent=2)
            
            print(f"单个风格变化已处理: {change_id} -> {action}")
            
            return {
                "success": True,
                "change_id": change_id,
                "action": action,
                "message": f"变化 {change_id} 已{action}",
                "updated_preview": {
                    "content": updated_preview,
                    "accepted_changes": len([c for c in suggested_changes if c.get("status") == "accepted"]),
                    "rejected_changes": len([c for c in suggested_changes if c.get("status") == "rejected"]),
                    "pending_changes": len([c for c in suggested_changes if c.get("status") == "pending"])
                }
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"处理单个风格变化失败: {str(e)}"
            }
    
    def _generate_updated_preview(self, original_content: str, suggested_changes: list) -> str:
        """
        根据已接受的变化生成更新后的预览内容
        
        Args:
            original_content: 原始内容
            suggested_changes: 建议的变化列表
            
        Returns:
            str: 更新后的预览内容
        """
        try:
            # 按位置排序变化，确保按顺序应用
            accepted_changes = [c for c in suggested_changes if c.get("status") == "accepted"]
            accepted_changes.sort(key=lambda x: x.get("position", {}).get("start", 0))
            
            # 从后往前应用变化，避免位置偏移
            updated_content = original_content
            offset = 0
            
            for change in accepted_changes:
                position = change.get("position", {})
                start = position.get("start", 0) + offset
                end = position.get("end", 0) + offset
                suggested_text = change.get("suggested_text", "")
                
                # 应用变化
                if start < len(updated_content) and end <= len(updated_content):
                    updated_content = updated_content[:start] + suggested_text + updated_content[end:]
                    # 更新偏移量
                    offset += len(suggested_text) - (end - start)
            
            return updated_content
            
        except Exception as e:
            print(f"生成更新预览失败: {e}")
            return original_content
