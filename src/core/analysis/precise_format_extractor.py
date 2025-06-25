"""
精确格式提取器 - 完整提取Word/PDF/Excel/PPT的所有格式信息

核心功能：
1. 提取所有样式信息（字体、字号、颜色、对齐方式等）
2. 提取布局信息（页边距、行间距、段间距等）
3. 提取表格格式（边框、单元格样式、列宽等）
4. 保存为可重用的格式模板
"""

import json
import os
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, asdict
from datetime import datetime
import hashlib

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Warning: python-docx not installed. Word format extraction will be limited.")

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, PatternFill
except ImportError:
    print("Warning: openpyxl not installed. Excel format extraction will be limited.")


@dataclass
class FontStyle:
    """字体样式"""
    name: str = "宋体"
    size: float = 12.0
    bold: bool = False
    italic: bool = False
    underline: bool = False
    color: str = "#000000"


@dataclass
class ParagraphStyle:
    """段落样式"""
    alignment: str = "left"  # left, center, right, justify
    line_spacing: float = 1.0
    space_before: float = 0.0
    space_after: float = 0.0
    first_line_indent: float = 0.0
    left_indent: float = 0.0
    right_indent: float = 0.0


@dataclass
class TableStyle:
    """表格样式"""
    table_alignment: str = "left"
    border_style: str = "single"
    border_width: float = 0.5
    border_color: str = "#000000"
    cell_padding: float = 0.0
    column_widths: List[float] = None
    row_heights: List[float] = None
    
    def __post_init__(self):
        if self.column_widths is None:
            self.column_widths = []
        if self.row_heights is None:
            self.row_heights = []


@dataclass
class PageStyle:
    """页面样式"""
    width: float = 8.5  # 英寸
    height: float = 11.0  # 英寸
    margin_top: float = 1.0
    margin_bottom: float = 1.0
    margin_left: float = 1.0
    margin_right: float = 1.0
    orientation: str = "portrait"  # portrait, landscape


@dataclass
class NumberingStyle:
    """标号样式"""
    level_1: str = "一、"  # 一级标号
    level_2: str = "（一）"  # 二级标号
    level_3: str = "1."     # 三级标号
    level_4: str = "（1）"  # 四级标号
    indent_per_level: float = 0.5  # 每级缩进（英寸）


@dataclass
class PageNumberStyle:
    """页码样式"""
    position: str = "footer_center"  # footer_center, footer_right, header_center
    format: str = "arabic"  # arabic(1,2,3), roman(i,ii,iii), chinese(一,二,三)
    start_number: int = 1
    show_total: bool = False  # 是否显示"第X页 共Y页"


@dataclass
class DocumentFormatTemplate:
    """精简的文档格式模板 - 专注核心格式要素"""
    template_id: str
    source_document: str
    document_type: str  # docx, pdf, xlsx, pptx
    extraction_time: str

    # 页面设置
    page_style: PageStyle

    # 核心样式定义
    title_style: Dict[str, Any]
    heading_styles: Dict[int, Dict[str, Any]]  # 各级标题样式
    paragraph_style: Dict[str, Any]
    table_styles: List[Dict[str, Any]]  # 表格样式（取第一个作为标准）

    # 标号和页码
    numbering_style: NumberingStyle
    page_number_style: PageNumberStyle

    # 应用规则
    style_mapping_rules: Dict[str, str]

    # 生成的模板文件
    template_files: Dict[str, str]


class PreciseFormatExtractor:
    """精确格式提取器"""
    
    def __init__(self, templates_dir: str = "src/core/knowledge_base/format_templates"):
        self.templates_dir = templates_dir
        os.makedirs(templates_dir, exist_ok=True)
        
        # 支持的文档类型
        self.supported_types = {
            '.docx': self._extract_word_format,
            '.pdf': self._extract_pdf_format,
            '.xlsx': self._extract_excel_format,
            '.pptx': self._extract_ppt_format
        }
    
    def extract_complete_format(self, file_path: str, template_name: str = None) -> Dict[str, Any]:
        """
        完整提取文档格式
        
        Args:
            file_path: 文档文件路径
            template_name: 模板名称（可选）
            
        Returns:
            {
                "success": bool,
                "template_id": str,
                "template_data": DocumentFormatTemplate,
                "template_files": Dict[str, str]
            }
        """
        try:
            if not os.path.exists(file_path):
                return {"error": f"文件不存在: {file_path}"}
            
            # 确定文档类型
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext not in self.supported_types:
                return {"error": f"不支持的文档类型: {file_ext}"}
            
            # 生成模板ID
            template_id = self._generate_template_id(file_path, template_name)
            
            # 提取格式信息
            extractor_func = self.supported_types[file_ext]
            format_data = extractor_func(file_path)
            
            if "error" in format_data:
                return format_data
            
            # 创建完整的格式模板
            template = DocumentFormatTemplate(
                template_id=template_id,
                source_document=os.path.basename(file_path),
                document_type=file_ext[1:],  # 去掉点号
                extraction_time=datetime.now().isoformat(),
                **format_data
            )
            
            # 保存模板
            saved_files = self._save_format_template(template)
            
            return {
                "success": True,
                "template_id": template_id,
                "template_data": template,
                "template_files": saved_files
            }
            
        except Exception as e:
            return {"error": f"格式提取失败: {str(e)}"}
    
    def _extract_word_format(self, file_path: str) -> Dict[str, Any]:
        """提取Word文档格式"""
        try:
            doc = Document(file_path)
            
            # 提取页面设置
            section = doc.sections[0]
            page_style = PageStyle(
                width=section.page_width.inches,
                height=section.page_height.inches,
                margin_top=section.top_margin.inches,
                margin_bottom=section.bottom_margin.inches,
                margin_left=section.left_margin.inches,
                margin_right=section.right_margin.inches,
                orientation="landscape" if section.page_width > section.page_height else "portrait"
            )
            
            # 提取样式信息
            styles_info = self._extract_word_styles(doc)
            
            # 提取表格格式
            table_styles = self._extract_word_table_styles(doc)
            
            # 提取标号样式
            numbering_style = self._extract_numbering_style(doc)

            # 提取页码样式
            page_number_style = self._extract_page_number_style(doc)

            return {
                "page_style": page_style,
                "title_style": styles_info.get("title", {}),
                "heading_styles": styles_info.get("headings", {}),
                "paragraph_style": styles_info.get("paragraph", {}),
                "table_styles": table_styles[:1] if table_styles else [],  # 只保留第一个表格样式作为标准
                "numbering_style": numbering_style,
                "page_number_style": page_number_style,
                "style_mapping_rules": self._generate_word_mapping_rules(styles_info),
                "template_files": {}
            }
            
        except Exception as e:
            return {"error": f"Word格式提取失败: {str(e)}"}
    
    def _extract_word_styles(self, doc: Document) -> Dict[str, Any]:
        """提取Word文档的样式信息"""
        styles_info = {
            "title": {},
            "headings": {},
            "paragraph": {},
            "list": {},
            "quote": {}
        }
        
        # 分析文档中的实际样式使用
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
            
            # 获取段落样式
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # 提取字体信息
            font_info = self._extract_font_info(paragraph)
            
            # 提取段落格式信息
            para_info = self._extract_paragraph_info(paragraph)
            
            # 根据样式名称和内容特征分类
            if "Title" in style_name or self._is_title(paragraph.text):
                styles_info["title"] = {
                    "font": font_info,
                    "paragraph": para_info,
                    "style_name": style_name
                }
            elif "Heading" in style_name or self._is_heading(paragraph.text):
                level = self._extract_heading_level(style_name, paragraph.text)
                styles_info["headings"][level] = {
                    "font": font_info,
                    "paragraph": para_info,
                    "style_name": style_name
                }
            elif not styles_info["paragraph"]:  # 使用第一个普通段落作为默认样式
                styles_info["paragraph"] = {
                    "font": font_info,
                    "paragraph": para_info,
                    "style_name": style_name
                }
        
        return styles_info
    
    def _extract_font_info(self, paragraph) -> Dict[str, Any]:
        """提取字体信息"""
        try:
            run = paragraph.runs[0] if paragraph.runs else None
            if not run:
                return {"name": "宋体", "size": 12.0, "bold": False, "italic": False}
            
            font = run.font
            return {
                "name": font.name or "宋体",
                "size": float(font.size.pt) if font.size else 12.0,
                "bold": font.bold or False,
                "italic": font.italic or False,
                "underline": font.underline or False,
                "color": str(font.color.rgb) if font.color and font.color.rgb else "#000000"
            }
        except:
            return {"name": "宋体", "size": 12.0, "bold": False, "italic": False}
    
    def _extract_paragraph_info(self, paragraph) -> Dict[str, Any]:
        """提取段落格式信息"""
        try:
            fmt = paragraph.paragraph_format
            
            # 对齐方式映射
            alignment_map = {
                WD_ALIGN_PARAGRAPH.LEFT: "left",
                WD_ALIGN_PARAGRAPH.CENTER: "center",
                WD_ALIGN_PARAGRAPH.RIGHT: "right",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
            }
            
            return {
                "alignment": alignment_map.get(fmt.alignment, "left"),
                "line_spacing": float(fmt.line_spacing) if fmt.line_spacing else 1.0,
                "space_before": float(fmt.space_before.pt) if fmt.space_before else 0.0,
                "space_after": float(fmt.space_after.pt) if fmt.space_after else 0.0,
                "first_line_indent": float(fmt.first_line_indent.pt) if fmt.first_line_indent else 0.0,
                "left_indent": float(fmt.left_indent.pt) if fmt.left_indent else 0.0,
                "right_indent": float(fmt.right_indent.pt) if fmt.right_indent else 0.0
            }
        except:
            return {"alignment": "left", "line_spacing": 1.0}
    
    def _extract_word_table_styles(self, doc: Document) -> List[Dict[str, Any]]:
        """提取Word表格样式"""
        table_styles = []
        
        for table in doc.tables:
            try:
                # 提取表格基本信息
                table_info = {
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "style_name": table.style.name if table.style else "Table Grid",
                    "alignment": "left",  # 默认值
                    "column_widths": [],
                    "row_heights": [],
                    "border_style": "single",
                    "cell_styles": []
                }
                
                # 提取列宽
                for col_idx in range(len(table.columns)):
                    try:
                        width = table.columns[col_idx].width.inches
                        table_info["column_widths"].append(width)
                    except:
                        table_info["column_widths"].append(1.0)  # 默认宽度
                
                # 提取行高
                for row in table.rows:
                    try:
                        height = row.height.inches if row.height else 0.3
                        table_info["row_heights"].append(height)
                    except:
                        table_info["row_heights"].append(0.3)  # 默认高度
                
                # 提取单元格样式（取第一个单元格作为样本）
                if table.rows and table.rows[0].cells:
                    cell = table.rows[0].cells[0]
                    if cell.paragraphs:
                        cell_font = self._extract_font_info(cell.paragraphs[0])
                        cell_para = self._extract_paragraph_info(cell.paragraphs[0])
                        table_info["cell_styles"] = {
                            "font": cell_font,
                            "paragraph": cell_para
                        }
                
                table_styles.append(table_info)
                
            except Exception as e:
                # 如果提取失败，添加默认样式
                table_styles.append({
                    "rows": 1,
                    "columns": 1,
                    "style_name": "Table Grid",
                    "alignment": "left",
                    "column_widths": [1.0],
                    "row_heights": [0.3],
                    "border_style": "single"
                })
        
        return table_styles
    
    def _is_title(self, text: str) -> bool:
        """判断是否为标题"""
        return len(text) < 50 and not text.endswith('。') and not text.endswith('：')
    
    def _is_heading(self, text: str) -> bool:
        """判断是否为标题"""
        import re
        heading_patterns = [r'^[一二三四五六七八九十]+[、.]', r'^\d+[、.]', r'^[（(]\d+[）)]']
        return any(re.match(pattern, text.strip()) for pattern in heading_patterns)
    
    def _extract_heading_level(self, style_name: str, text: str) -> int:
        """提取标题级别"""
        import re
        
        # 从样式名称提取
        if "Heading" in style_name:
            match = re.search(r'Heading (\d+)', style_name)
            if match:
                return int(match.group(1))
        
        # 从文本内容推断
        if re.match(r'^[一二三四五六七八九十]+[、.]', text.strip()):
            return 1
        elif re.match(r'^\d+[、.]', text.strip()):
            return 2
        elif re.match(r'^[（(]\d+[）)]', text.strip()):
            return 3
        
        return 1  # 默认一级标题
    
    def _generate_word_mapping_rules(self, styles_info: Dict[str, Any]) -> Dict[str, str]:
        """生成Word样式映射规则"""
        return {
            "title": "title_style",
            "heading_1": "heading_1_style",
            "heading_2": "heading_2_style", 
            "heading_3": "heading_3_style",
            "paragraph": "paragraph_style",
            "table": "table_style_0"  # 使用第一个表格样式
        }
    
    def _extract_pdf_format(self, file_path: str) -> Dict[str, Any]:
        """提取PDF格式（简化实现）"""
        # PDF格式提取较复杂，这里提供基础框架
        return {
            "page_style": PageStyle(),
            "title_style": {},
            "heading_styles": {},
            "paragraph_style": {},
            "table_styles": [],
            "list_styles": {},
            "quote_styles": {},
            "style_mapping_rules": {},
            "template_files": {}
        }
    
    def _extract_excel_format(self, file_path: str) -> Dict[str, Any]:
        """提取Excel格式（简化实现）"""
        # Excel格式提取，主要关注单元格样式
        return {
            "page_style": PageStyle(),
            "title_style": {},
            "heading_styles": {},
            "paragraph_style": {},
            "table_styles": [],
            "list_styles": {},
            "quote_styles": {},
            "style_mapping_rules": {},
            "template_files": {}
        }
    
    def _extract_ppt_format(self, file_path: str) -> Dict[str, Any]:
        """提取PowerPoint格式（简化实现）"""
        # PPT格式提取，主要关注幻灯片布局和文本样式
        return {
            "page_style": PageStyle(),
            "title_style": {},
            "heading_styles": {},
            "paragraph_style": {},
            "table_styles": [],
            "list_styles": {},
            "quote_styles": {},
            "style_mapping_rules": {},
            "template_files": {}
        }
    
    def _generate_template_id(self, file_path: str, template_name: str = None) -> str:
        """生成模板ID"""
        if template_name:
            base_string = template_name
        else:
            base_string = os.path.basename(file_path) + str(os.path.getmtime(file_path))
        
        return hashlib.md5(base_string.encode()).hexdigest()[:12]
    
    def _save_format_template(self, template: DocumentFormatTemplate) -> Dict[str, str]:
        """保存格式模板"""
        template_dir = os.path.join(self.templates_dir, template.template_id)
        os.makedirs(template_dir, exist_ok=True)
        
        # 保存JSON格式的模板数据
        json_file = os.path.join(template_dir, "template.json")
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(asdict(template), f, ensure_ascii=False, indent=2, default=str)
        
        # 生成Word模板文件
        docx_template = self._generate_word_template(template)
        docx_file = os.path.join(template_dir, "template.docx")
        docx_template.save(docx_file)
        
        return {
            "json": json_file,
            "docx": docx_file,
            "template_dir": template_dir
        }
    
    def _generate_word_template(self, template: DocumentFormatTemplate) -> Document:
        """生成Word模板文件"""
        doc = Document()
        
        # 设置页面格式
        section = doc.sections[0]
        section.page_width = Inches(template.page_style.width)
        section.page_height = Inches(template.page_style.height)
        section.top_margin = Inches(template.page_style.margin_top)
        section.bottom_margin = Inches(template.page_style.margin_bottom)
        section.left_margin = Inches(template.page_style.margin_left)
        section.right_margin = Inches(template.page_style.margin_right)
        
        # 添加样式示例
        if template.title_style:
            title = doc.add_paragraph("标题样式示例")
            self._apply_word_style(title, template.title_style)
        
        for level, style in template.heading_styles.items():
            heading = doc.add_paragraph(f"标题{level}样式示例")
            self._apply_word_style(heading, style)
        
        if template.paragraph_style:
            para = doc.add_paragraph("段落样式示例。这是一个示例段落，用于展示段落的格式设置。")
            self._apply_word_style(para, template.paragraph_style)
        
        return doc
    
    def _apply_word_style(self, paragraph, style_info: Dict[str, Any]):
        """应用Word样式到段落"""
        try:
            # 应用字体样式
            if "font" in style_info and paragraph.runs:
                font_info = style_info["font"]
                for run in paragraph.runs:
                    if "name" in font_info:
                        run.font.name = font_info["name"]
                    if "size" in font_info:
                        run.font.size = Pt(font_info["size"])
                    if "bold" in font_info:
                        run.font.bold = font_info["bold"]
                    if "italic" in font_info:
                        run.font.italic = font_info["italic"]
            
            # 应用段落格式
            if "paragraph" in style_info:
                para_info = style_info["paragraph"]
                fmt = paragraph.paragraph_format
                
                if "alignment" in para_info:
                    alignment_map = {
                        "left": WD_ALIGN_PARAGRAPH.LEFT,
                        "center": WD_ALIGN_PARAGRAPH.CENTER,
                        "right": WD_ALIGN_PARAGRAPH.RIGHT,
                        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    fmt.alignment = alignment_map.get(para_info["alignment"])
                
                if "line_spacing" in para_info:
                    fmt.line_spacing = para_info["line_spacing"]
                
                if "space_before" in para_info:
                    fmt.space_before = Pt(para_info["space_before"])
                
                if "space_after" in para_info:
                    fmt.space_after = Pt(para_info["space_after"])
        
        except Exception as e:
            print(f"应用样式失败: {e}")
    
    def load_format_template(self, template_id: str) -> Dict[str, Any]:
        """加载格式模板"""
        template_dir = os.path.join(self.templates_dir, template_id)
        json_file = os.path.join(template_dir, "template.json")
        
        if not os.path.exists(json_file):
            return {"error": f"模板不存在: {template_id}"}
        
        try:
            with open(json_file, 'r', encoding='utf-8') as f:
                template_data = json.load(f)
            
            return {
                "success": True,
                "template_data": template_data,
                "template_files": {
                    "json": json_file,
                    "docx": os.path.join(template_dir, "template.docx"),
                    "template_dir": template_dir
                }
            }
        except Exception as e:
            return {"error": f"加载模板失败: {str(e)}"}
