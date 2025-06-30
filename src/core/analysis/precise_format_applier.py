#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
精确格式应用器

Author: AI Assistant (Claude)
Created: 2025-01-28
Last Modified: 2025-01-28
Modified By: AI Assistant (Claude)
AI Assisted: 是 - Claude 3.5 Sonnet
Version: v1.0
License: MIT
"""





import os
import json
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("Warning: python-docx not installed. Word format application will be limited.")

from .precise_format_extractor import PreciseFormatExtractor, DocumentFormatTemplate


@dataclass
class ContentElement:
    
    def __init__(self, templates_dir: str = "src/core/knowledge_base/format_templates"):
        self.templates_dir = templates_dir
        self.format_extractor = PreciseFormatExtractor(templates_dir)
    
    def apply_format_precisely(self, source_content: str, target_template_id: str, 
                              output_path: str = None) -> Dict[str, Any]:
        try:
            # 1. 加载目标格式模板
            template_result = self.format_extractor.load_format_template(target_template_id)
            if "error" in template_result:
                return template_result
            
            template_data = template_result["template_data"]
            
            # 2. 解析源文档内容结构
            if os.path.exists(source_content):
                content_elements = self._parse_document_file(source_content)
            else:
                content_elements = self._parse_text_content(source_content)
            
            # 3. 根据模板类型生成对应格式的文档
            doc_type = template_data.get("document_type", "docx")
            
            if doc_type == "docx":
                result = self._generate_word_document(content_elements, template_data, output_path)
            elif doc_type == "pdf":
                result = self._generate_pdf_document(content_elements, template_data, output_path)
            elif doc_type == "xlsx":
                result = self._generate_excel_document(content_elements, template_data, output_path)
            elif doc_type == "pptx":
                result = self._generate_ppt_document(content_elements, template_data, output_path)
            else:
                return {"error": f"不支持的文档类型: {doc_type}"}
            
            return result
            
        except Exception as e:
            return {"error": f"格式应用失败: {str(e)}"}
    
    def _parse_document_file(self, file_path: str) -> List[ContentElement]:
        try:
            doc = Document(file_path)
            elements = []
            
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    continue
                
                # 识别元素类型
                element_type, level = self._identify_content_type(paragraph.text, paragraph.style.name if paragraph.style else "Normal")
                
                elements.append(ContentElement(
                    type=element_type,
                    content=paragraph.text.strip(),
                    level=level
                ))
            
            # 处理表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                elements.append(ContentElement(
                    type="table",
                    content="",
                    table_data=table_data
                ))
            
            return elements
            
        except Exception as e:
            return [ContentElement(type="paragraph", content=f"解析失败: {str(e)}")]
    
    def _parse_text_content(self, content: str) -> List[ContentElement]:
        text = text.strip()
        
        # 基于样式名称判断
        if "Title" in style_name:
            return "title", 0
        elif "Heading" in style_name:
            level = 1
            if "Heading 2" in style_name:
                level = 2
            elif "Heading 3" in style_name:
                level = 3
            return "heading", level
        
        # 基于文本模式判断
        if len(text) < 50 and not text.endswith('。') and not text.endswith('：'):
            # 可能是标题
            if re.match(r'^[一二三四五六七八九十]+[、.]', text):
                return "heading", 1
            elif re.match(r'^\d+[、.]', text):
                return "heading", 2
            elif re.match(r'^[（(]\d+[）)]', text):
                return "heading", 3
            else:
                return "title", 0
        
        # 列表项
        if re.match(r'^\s*[-*+]\s', text) or re.match(r'^\s*\d+\.\s', text):
            return "list", 1
        
        # 默认为段落
        return "paragraph", 0
    
    def _generate_word_document(self, content_elements: List[ContentElement], 
                               template_data: Dict[str, Any], output_path: str = None) -> Dict[str, Any]:
        try:
            section = doc.sections[0]
            
            if "width" in page_style:
                section.page_width = Inches(page_style["width"])
            if "height" in page_style:
                section.page_height = Inches(page_style["height"])
            if "margin_top" in page_style:
                section.top_margin = Inches(page_style["margin_top"])
            if "margin_bottom" in page_style:
                section.bottom_margin = Inches(page_style["margin_bottom"])
            if "margin_left" in page_style:
                section.left_margin = Inches(page_style["margin_left"])
            if "margin_right" in page_style:
                section.right_margin = Inches(page_style["margin_right"])
                
        except Exception as e:
            print(f"应用页面设置失败: {e}")
    
    def _apply_word_style(self, paragraph, style_info: Dict[str, Any]):
        try:
            if not table_data:
                return None
            
            rows = len(table_data)
            cols = max(len(row) for row in table_data) if table_data else 1
            
            table = doc.add_table(rows=rows, cols=cols)
            
            # 填充表格数据
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = str(cell_data)
                        
                        # 应用单元格样式
                        if cell.paragraphs:
                            self._apply_table_cell_style(cell.paragraphs[0], template_data, style_index)
            
            # 应用表格样式
            self._apply_table_style(table, template_data, style_index)
            
            return table
            
        except Exception as e:
            print(f"创建表格失败: {e}")
            return None
    
    def _apply_table_style(self, table, template_data: Dict[str, Any], style_index: int):
        try:
            table_styles = template_data.get("table_styles", [])
            if style_index < len(table_styles):
                table_style = table_styles[style_index]
                if "cell_styles" in table_style:
                    self._apply_word_style(paragraph, table_style["cell_styles"])
        except Exception as e:
            print(f"应用单元格样式失败: {e}")
    
    def _generate_pdf_document(self, content_elements: List[ContentElement], 
                              template_data: Dict[str, Any], output_path: str = None) -> Dict[str, Any]:
        try:
            # TODO: MVP仅占位，后续完善 - 当前只支持纯文本PDF导出
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import inch
            except ImportError:
                return {"error": "MVP: PDF生成需要安装reportlab库: pip install reportlab"}
            
            # 确定输出路径
            if not output_path:
                import time
                output_path = f"output/document_{int(time.time())}.pdf"
            
            # 确保输出目录存在
            import os
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # 创建PDF文档
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # 处理内容元素 - MVP只处理文本类型
            for element in content_elements:
                if element.type == 'title':
                    # 标题样式
                    title_style = styles['Heading1']
                    story.append(Paragraph(element.content, title_style))
                    story.append(Spacer(1, 12))
                    
                elif element.type == 'heading':
                    # 标题样式
                    heading_style = styles['Heading2']
                    story.append(Paragraph(element.content, heading_style))
                    story.append(Spacer(1, 8))
                    
                elif element.type == 'paragraph':
                    # 段落内容
                    para_style = styles['Normal']
                    story.append(Paragraph(element.content, para_style))
                    story.append(Spacer(1, 6))
                    
                elif element.type == 'table':
                    # TODO: MVP仅占位，后续完善 - 表格功能暂不支持
                    story.append(Paragraph(f"[表格内容: {len(element.table_data) if element.table_data else 0} 行]", styles['Normal']))
                    story.append(Spacer(1, 6))
                    
                elif element.type == 'list':
                    # TODO: MVP仅占位，后续完善 - 列表功能暂不支持
                    story.append(Paragraph(f"[列表内容: {element.content}]", styles['Normal']))
                    story.append(Spacer(1, 6))
            
            # 生成PDF
            doc.build(story)
            
            return {
                "success": True,
                "output_path": output_path,
                "file_size": os.path.getsize(output_path),
                "pages": 1,  # MVP简化：假设只有1页
                "note": "MVP: 仅支持纯文本导出，表格和复杂样式后续完善"
            }
            
        except Exception as e:
            return {"error": f"MVP: PDF生成失败: {str(e)}"}
    
    def _generate_excel_document(self, content_elements: List[ContentElement], 
                                template_data: Dict[str, Any], output_path: str = None) -> Dict[str, Any]:
        # TODO: MVP仅占位，后续完善 - Excel生成功能待实现
        return {"error": "MVP: Excel生成功能待实现，后续使用openpyxl库完善"}
    
    def _generate_ppt_document(self, content_elements: List[ContentElement], 
                              template_data: Dict[str, Any], output_path: str = None) -> Dict[str, Any]:
        # TODO: MVP仅占位，后续完善 - PowerPoint生成功能待实现
        return {"error": "MVP: PowerPoint生成功能待实现，后续使用python-pptx库完善"}
    
    def batch_apply_format(self, source_files: List[str], target_template_id: str, 
                          output_dir: str = "output/batch_formatted") -> Dict[str, Any]:
        try:
            os.makedirs(output_dir, exist_ok=True)
            
            results = {}
            for file_path in source_files:
                file_name = os.path.basename(file_path)
                output_path = os.path.join(output_dir, file_name)
                
                result = self.apply_format_precisely(file_path, target_template_id, output_path)
                results[file_name] = result
            
            return {
                "success": True,
                "results": results
            }
        
        except Exception as e:
            return {"error": f"批量处理失败: {str(e)}"}
    
    def _identify_content_type(self, text: str, style_name: str) -> Tuple[str, int]:
        text = text.strip()
        
        # 基于样式名称判断
        if "Title" in style_name:
            return "title", 0
        elif "Heading" in style_name:
            level = 1
            if "Heading 2" in style_name:
                level = 2
            elif "Heading 3" in style_name:
                level = 3
            return "heading", level
        
        # 基于文本模式判断
        if len(text) < 50 and not text.endswith('。') and not text.endswith('：'):
            # 可能是标题
            if re.match(r'^[一二三四五六七八九十]+[、.]', text):
                return "heading", 1
            elif re.match(r'^\d+[、.]', text):
                return "heading", 2
            elif re.match(r'^[（(]\d+[）)]', text):
                return "heading", 3
            else:
                return "title", 0
        
        # 列表项
        if re.match(r'^\s*[-*+]\s', text) or re.match(r'^\s*\d+\.\s', text):
            return "list", 1
        
        # 默认为段落
        return "paragraph", 0
    
    def _apply_table_cell_style(self, paragraph, template_data: Dict[str, Any], style_index: int):
        pass
    
    def _apply_paragraph_style(self, paragraph, style_info: Dict[str, Any]):
        pass
    
    def _apply_character_style(self, run, style_info: Dict[str, Any]):
        pass
    
    def _apply_list_style(self, paragraph, style_info: Dict[str, Any]):
        pass
    
    def _apply_table_style_v2(self, table, template_data: Dict[str, Any], style_index: int):
        pass