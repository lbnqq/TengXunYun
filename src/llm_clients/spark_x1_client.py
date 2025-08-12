#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
星火X1大模型客户端 - 智能填报模块专用

Author: AI Assistant (Claude)
Created: 2025-07-08
Last Modified: 2025-07-08
Modified By: AI Assistant (Claude)
AI Assisted: 是 - Claude 3.5 Sonnet
Version: v1.0 - 集成星火X1 HTTP接口
License: MIT
"""

import os
import json
import logging
import requests
import re
import tempfile
import uuid
from typing import Dict, Any, Optional, List, Iterator
from datetime import datetime
from docx import Document

class SparkX1Client:
    """星火X1大模型HTTP客户端 - 专为智能填报设计"""
    
    def __init__(self, api_password: str = None):
        """
        初始化星火X1客户端
        
        Args:
            api_password: 星火X1的APIPassword (格式: AK:SK)
        """
        self.api_password = api_password or os.getenv('SPARK_X1_API_PASSWORD')
        if not self.api_password:
            raise ValueError("缺少星火X1 APIPassword配置")
        
        self.base_url = "https://spark-api-open.xf-yun.com/v2/chat/completions"
        self.model = "x1"
        self.logger = logging.getLogger(__name__)

        # 验证APIPassword格式
        if ':' not in self.api_password:
            raise ValueError("APIPassword格式错误，应为 AK:SK 格式")

        # 多轮对话会话管理
        self.conversations = {}  # 存储会话历史
    
    def _call_api(self, messages: List[Dict[str, str]], stream: bool = False, **kwargs) -> Dict[str, Any]:
        """
        调用星火X1 API
        
        Args:
            messages: 对话消息列表
            stream: 是否流式返回
            **kwargs: 其他参数
            
        Returns:
            API响应结果
        """
        try:
            headers = {
                "Authorization": f"Bearer {self.api_password}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": self.model,
                "user": kwargs.get("user", "smart_form_filler"),
                "messages": messages,
                "stream": stream,
                "temperature": kwargs.get("temperature", 0.7),
                "max_tokens": kwargs.get("max_tokens", 4096)
            }
            
            self.logger.info(f"调用星火X1 API，消息数量: {len(messages)}")
            
            response = requests.post(
                self.base_url,
                headers=headers,
                json=payload,
                timeout=kwargs.get("timeout", 180)
            )
            
            if response.status_code == 200:
                result = response.json()
                if result.get("code") == 0:
                    return result
                else:
                    raise Exception(f"API返回错误: {result.get('message', '未知错误')}")
            else:
                raise Exception(f"HTTP请求失败: {response.status_code}")
                
        except Exception as e:
            self.logger.error(f"星火X1 API调用失败: {e}")
            raise
    
    def generate_summary(self, work_content: str) -> Dict[str, Any]:
        """
        生成年度总结
        
        Args:
            work_content: 工作内容描述
            
        Returns:
            生成结果 {'success': bool, 'content': str, 'file_path': str, 'filename': str}
        """
        try:
            prompt = f"""
请根据以下工作内容，生成一份专业的年度工作总结，要求：

1. 总结格式要正式、专业
2. 内容要条理清晰，逻辑性强
3. 突出工作成果和个人成长
4. 字数控制在800-1200字之间
5. 包含工作回顾、主要成就、经验总结、不足反思、未来规划等部分

工作内容：
{work_content}

请直接输出年度总结内容，不需要额外的说明。
"""
            
            messages = [{"role": "user", "content": prompt}]
            result = self._call_api(messages, timeout=120)
            
            # 提取生成的内容
            content = result["choices"][0]["message"]["content"]
            
            # 生成Word文档
            doc_result = self._create_summary_document(content)
            
            return {
                "success": True,
                "content": content,
                "file_path": doc_result["file_path"],
                "filename": doc_result["filename"],
                "usage": result.get("usage", {})
            }
            
        except Exception as e:
            self.logger.error(f"生成年度总结失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "content": None,
                "file_path": None,
                "filename": None
            }
    
    def generate_resume(self, personal_info: str) -> Dict[str, Any]:
        """
        生成个人简历
        
        Args:
            personal_info: 个人信息描述
            
        Returns:
            生成结果 {'success': bool, 'data': dict, 'file_path': str, 'filename': str}
        """
        try:
            prompt = f"""
请分析以下个人信息，提取关键信息并按JSON格式输出，用于填充简历模板：

个人信息：
{personal_info}

请严格按照以下JSON格式输出，如果某项信息缺失，请填写"暂无相关内容"：

{{
    "姓名": "提取的姓名",
    "电话": "提取的电话号码",
    "邮箱": "提取的邮箱地址",
    "个人陈述": "根据信息生成的个人陈述，突出优势和特点",
    "教育背景": [
        {{"学校": "学校名称", "专业": "专业名称", "学位": "学位", "时间": "时间段"}},
        {{"学校": "暂无相关内容", "专业": "暂无相关内容", "学位": "暂无相关内容", "时间": "暂无相关内容"}},
        {{"学校": "暂无相关内容", "专业": "暂无相关内容", "学位": "暂无相关内容", "时间": "暂无相关内容"}}
    ],
    "工作经验": [
        {{"公司": "公司名称", "职位": "职位名称", "时间": "时间段", "职责": "主要职责描述"}},
        {{"公司": "暂无相关内容", "职位": "暂无相关内容", "时间": "暂无相关内容", "职责": "暂无相关内容"}},
        {{"公司": "暂无相关内容", "职位": "暂无相关内容", "时间": "暂无相关内容", "职责": "暂无相关内容"}}
    ],
    "项目经验": [
        {{"名称": "项目名称", "角色": "担任角色", "时间": "时间段", "描述": "项目描述和成果"}},
        {{"名称": "暂无相关内容", "角色": "暂无相关内容", "时间": "暂无相关内容", "描述": "暂无相关内容"}},
        {{"名称": "暂无相关内容", "角色": "暂无相关内容", "时间": "暂无相关内容", "描述": "暂无相关内容"}}
    ],
    "专业技能": "整理后的技能列表，包括编程语言、框架、工具等"
}}

只输出JSON数据，不要包含任何其他文字说明。
"""
            
            messages = [{"role": "user", "content": prompt}]
            result = self._call_api(messages, timeout=180)
            
            # 提取生成的内容
            ai_response = result["choices"][0]["message"]["content"]
            
            # 解析JSON数据
            json_match = re.search(r'\{.*\}', ai_response, re.DOTALL)
            if not json_match:
                raise Exception("AI响应中未找到有效的JSON数据")
            
            json_str = json_match.group()
            resume_data = json.loads(json_str)
            
            # 生成Word文档
            doc_result = self._create_resume_document(resume_data)
            
            return {
                "success": True,
                "data": resume_data,
                "file_path": doc_result["file_path"],
                "filename": doc_result["filename"],
                "usage": result.get("usage", {})
            }
            
        except Exception as e:
            self.logger.error(f"生成简历失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "data": None,
                "file_path": None,
                "filename": None
            }
    
    def _create_summary_document(self, content: str) -> Dict[str, str]:
        """创建年度总结Word文档"""
        try:
            doc = Document()

            # 添加标题
            title = doc.add_heading('年度工作总结', 0)
            title.alignment = 1  # 居中对齐

            # 添加日期
            date_para = doc.add_paragraph(f'日期：{datetime.now().strftime("%Y年%m月%d日")}')
            date_para.alignment = 2  # 右对齐

            # 添加内容
            doc.add_paragraph(content)

            # 保存文档
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"年度工作总结_{timestamp}.docx"

            # 创建临时文件
            temp_dir = tempfile.gettempdir()
            file_path = os.path.join(temp_dir, filename)
            doc.save(file_path)

            return {
                "file_path": file_path,
                "filename": filename
            }

        except Exception as e:
            self.logger.error(f"创建年度总结文档失败: {e}")
            raise

    def _create_resume_document(self, resume_data: Dict[str, Any]) -> Dict[str, str]:
        """创建简历Word文档"""
        try:
            doc = Document()

            # 添加标题
            title = doc.add_heading('个人简历', 0)
            title.alignment = 1  # 居中对齐

            # 基本信息
            doc.add_heading('基本信息', level=1)
            basic_info = doc.add_paragraph()
            basic_info.add_run(f"姓名：{resume_data.get('姓名', '暂无')}\n")
            basic_info.add_run(f"电话：{resume_data.get('电话', '暂无')}\n")
            basic_info.add_run(f"邮箱：{resume_data.get('邮箱', '暂无')}")

            # 个人陈述
            if resume_data.get('个人陈述') and resume_data['个人陈述'] != '暂无相关内容':
                doc.add_heading('个人陈述', level=1)
                doc.add_paragraph(resume_data['个人陈述'])

            # 教育背景
            doc.add_heading('教育背景', level=1)
            education_list = resume_data.get('教育背景', [])
            for edu in education_list:
                if edu.get('学校') and edu['学校'] != '暂无相关内容':
                    edu_para = doc.add_paragraph()
                    edu_para.add_run(f"• {edu.get('时间', '')} - {edu.get('学校', '')}\n")
                    edu_para.add_run(f"  专业：{edu.get('专业', '')} | 学位：{edu.get('学位', '')}")

            # 工作经验
            doc.add_heading('工作经验', level=1)
            work_list = resume_data.get('工作经验', [])
            for work in work_list:
                if work.get('公司') and work['公司'] != '暂无相关内容':
                    work_para = doc.add_paragraph()
                    work_para.add_run(f"• {work.get('时间', '')} - {work.get('公司', '')}\n")
                    work_para.add_run(f"  职位：{work.get('职位', '')}\n")
                    work_para.add_run(f"  职责：{work.get('职责', '')}")

            # 项目经验
            doc.add_heading('项目经验', level=1)
            project_list = resume_data.get('项目经验', [])
            for project in project_list:
                if project.get('名称') and project['名称'] != '暂无相关内容':
                    proj_para = doc.add_paragraph()
                    proj_para.add_run(f"• {project.get('名称', '')} ({project.get('时间', '')})\n")
                    proj_para.add_run(f"  角色：{project.get('角色', '')}\n")
                    proj_para.add_run(f"  描述：{project.get('描述', '')}")

            # 专业技能
            if resume_data.get('专业技能') and resume_data['专业技能'] != '暂无相关内容':
                doc.add_heading('专业技能', level=1)
                doc.add_paragraph(resume_data['专业技能'])

            # 保存文档
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            name = resume_data.get('姓名', '未知')
            filename = f"个人简历_{name}_{timestamp}.docx"

            # 创建临时文件
            temp_dir = tempfile.gettempdir()
            file_path = os.path.join(temp_dir, filename)
            doc.save(file_path)

            return {
                "file_path": file_path,
                "filename": filename
            }

        except Exception as e:
            self.logger.error(f"创建简历文档失败: {e}")
            raise

    def format_text(self, instruction: str, content: str, **kwargs) -> str:
        """
        格式化文本内容

        Args:
            instruction: 格式化指令
            content: 要格式化的内容
            **kwargs: 其他参数 (temperature, max_tokens等)

        Returns:
            格式化后的文本
        """
        try:
            messages = [
                {"role": "user", "content": f"{instruction}\n\n原文内容：\n{content}"}
            ]

            result = self._call_api(messages, **kwargs)

            # 从API响应中提取内容
            if result.get("choices") and len(result["choices"]) > 0:
                formatted_content = result["choices"][0]["message"]["content"]
                return formatted_content
            else:
                raise Exception("API响应格式错误，未找到choices字段")

        except Exception as e:
            self.logger.error(f"格式化文本失败: {str(e)}")
            raise Exception(f"格式化文本失败: {str(e)}")

    def format_text_stream(self, instruction: str, content: str, **kwargs) -> Iterator[str]:
        """
        流式格式化文本内容

        Args:
            instruction: 格式化指令
            content: 要格式化的内容
            **kwargs: 其他参数

        Yields:
            格式化后的文本片段
        """
        try:
            messages = [
                {"role": "user", "content": f"{instruction}\n\n原文内容：\n{content}"}
            ]

            result = self._call_api(messages, stream=True, **kwargs)

            if result.get('success'):
                yield result['content']
            else:
                raise Exception(result.get('error', '流式格式化失败'))

        except Exception as e:
            self.logger.error(f"流式格式化失败: {str(e)}")
            raise Exception(f"流式格式化失败: {str(e)}")

    def multi_turn_chat(self, messages: List[Dict], **kwargs) -> str:
        """
        多轮对话

        Args:
            messages: 对话消息列表
            **kwargs: 其他参数

        Returns:
            AI回复内容
        """
        try:
            result = self._call_api(messages, **kwargs)

            # 从API响应中提取内容
            if result.get("choices") and len(result["choices"]) > 0:
                response_content = result["choices"][0]["message"]["content"]
                return response_content
            else:
                raise Exception("API响应格式错误，未找到choices字段")

        except Exception as e:
            self.logger.error(f"多轮对话失败: {str(e)}")
            raise Exception(f"多轮对话失败: {str(e)}")

    def start_conversation(self, initial_message: str) -> str:
        """
        开始新的对话会话

        Args:
            initial_message: 初始消息

        Returns:
            会话ID
        """
        conversation_id = str(uuid.uuid4())
        self.conversations[conversation_id] = [
            {"role": "user", "content": initial_message}
        ]
        return conversation_id

    def continue_conversation(self, conversation_id: str, message: str, **kwargs) -> str:
        """
        继续对话会话

        Args:
            conversation_id: 会话ID
            message: 用户消息
            **kwargs: 其他参数

        Returns:
            AI回复内容
        """
        try:
            if conversation_id not in self.conversations:
                raise ValueError(f"会话ID {conversation_id} 不存在")

            # 添加用户消息到会话历史
            self.conversations[conversation_id].append({"role": "user", "content": message})

            # 调用API
            result = self._call_api(self.conversations[conversation_id], **kwargs)

            # 从API响应中提取内容
            if result.get("choices") and len(result["choices"]) > 0:
                response_content = result["choices"][0]["message"]["content"]
                # 添加AI回复到会话历史
                self.conversations[conversation_id].append({"role": "assistant", "content": response_content})
                return response_content
            else:
                raise Exception("API响应格式错误，未找到choices字段")

        except Exception as e:
            self.logger.error(f"继续对话失败: {str(e)}")
            raise Exception(f"继续对话失败: {str(e)}")

    def get_conversation_history(self, conversation_id: str) -> List[Dict]:
        """
        获取对话历史

        Args:
            conversation_id: 会话ID

        Returns:
            对话历史列表
        """
        return self.conversations.get(conversation_id, [])

    def clear_conversation(self, conversation_id: str) -> bool:
        """
        清除对话会话

        Args:
            conversation_id: 会话ID

        Returns:
            是否成功清除
        """
        if conversation_id in self.conversations:
            del self.conversations[conversation_id]
            return True
        return False

    def is_available(self) -> bool:
        """检查服务是否可用"""
        try:
            # 简单的健康检查
            test_messages = [{"role": "user", "content": "你好"}]
            result = self._call_api(test_messages, timeout=30)
            return result.get("code") == 0
        except Exception as e:
            self.logger.error(f"健康检查失败: {e}")
            return False

    # ==================== 文风统一专用方法 ====================

    def generate_with_style(self, prompt_instruction: str, user_id: str = "user_style_transfer",
                          temperature: float = 0.7, max_tokens: int = 3000, timeout: int = 120) -> str:
        """
        【AI文风统一】调用星火大模型 X1，生成指定风格文本

        Args:
            prompt_instruction: 包含风格描述和内容的Prompt
            user_id: 用户ID（可选，默认"user_style_transfer"）
            temperature: 生成文本的随机性（0.0-1.0，默认0.7）
            max_tokens: 最大生成token数
            timeout: 超时时间（秒）

        Returns:
            生成的文本内容或错误信息
        """
        try:
            self.logger.info(f"🎨 开始文风统一生成，用户: {user_id}")

            headers = {
                "Authorization": f"Bearer {self.api_password}",
                "Content-Type": "application/json"
            }

            data = {
                "model": self.model,
                "user": user_id,
                "messages": [
                    {
                        "role": "user",
                        "content": prompt_instruction
                    }
                ],
                "temperature": temperature,
                "max_tokens": max_tokens
            }

            response = requests.post(self.base_url, headers=headers,
                                   data=json.dumps(data), timeout=timeout)
            response.raise_for_status()
            response_json = response.json()

            # 解析官方接口返回格式
            if response_json and response_json.get("choices"):
                # 兼容openai风格，取第一个choice的message内容
                result = response_json["choices"][0]["message"]["content"].strip()
                self.logger.info(f"✅ 文风统一生成成功，长度: {len(result)}")
                return result
            elif response_json and response_json.get("error"):
                error_msg = f"API 返回错误: {response_json['error']}"
                self.logger.error(error_msg)
                return f"[错误] {error_msg}"
            else:
                error_msg = "无法解析的 API 响应格式"
                self.logger.error(error_msg)
                return f"[错误] {error_msg}"

        except requests.exceptions.RequestException as e:
            error_msg = f"API 请求失败: {e}"
            self.logger.error(error_msg)
            return f"[错误] {error_msg}"
        except json.JSONDecodeError:
            error_msg = "API 响应不是有效的 JSON 格式"
            self.logger.error(error_msg)
            return f"[错误] {error_msg}"
        except Exception as e:
            error_msg = f"文风统一生成异常: {e}"
            self.logger.error(error_msg)
            return f"[错误] {error_msg}"

    def few_shot_generate(self, examples: List[str], content: str, style_description: str = "",
                         user_id: str = "user_few_shot", temperature: float = 0.7,
                         max_tokens: int = 3000, timeout: int = 120) -> str:
        """
        Few-Shot风格生成

        Args:
            examples: 风格示例列表
            content: 要转换的内容
            style_description: 风格描述
            user_id: 用户ID
            temperature: 温度参数
            max_tokens: 最大token数
            timeout: 超时时间

        Returns:
            生成的文本内容
        """
        try:
            self.logger.info(f"🔄 开始Few-Shot风格生成，示例数: {len(examples)}")

            # 构建Few-Shot提示词
            prompt = "请参考以下示例的写作风格，将给定内容转换为相同的风格。\n\n"

            if style_description:
                prompt += f"风格描述：{style_description}\n\n"

            prompt += "参考示例：\n"
            for i, example in enumerate(examples, 1):
                prompt += f"示例{i}：{example}\n"

            prompt += f"\n请将以下内容转换为上述示例的风格：\n{content}"

            # 调用生成方法
            return self.generate_with_style(
                prompt_instruction=prompt,
                user_id=user_id,
                temperature=temperature,
                max_tokens=max_tokens,
                timeout=timeout
            )

        except Exception as e:
            error_msg = f"Few-Shot生成异常: {e}"
            self.logger.error(error_msg)
            return f"[错误] {error_msg}"

    def analyze_style(self, document_content: str, user_id: str = "user_style_analysis",
                     temperature: float = 0.3, timeout: int = 60) -> str:
        """
        分析文档风格特征

        Args:
            document_content: 文档内容
            user_id: 用户ID
            temperature: 温度参数（使用较低值确保分析一致性）
            timeout: 超时时间

        Returns:
            风格分析结果
        """
        try:
            self.logger.info(f"🔍 开始文档风格分析，内容长度: {len(document_content)}")

            analysis_prompt = """请分析以下文档的写作风格特征，包括：
1. 语言风格（正式/非正式、严谨/轻松等）
2. 句式特点（长短句比例、复杂度等）
3. 词汇选择（专业术语、口语化程度等）
4. 表达方式（直接/委婉、客观/主观等）
5. 整体语调和情感色彩

请提供详细的风格分析报告：

文档内容：
""" + document_content

            return self.generate_with_style(
                prompt_instruction=analysis_prompt,
                user_id=user_id,
                temperature=temperature,
                max_tokens=2000,
                timeout=timeout
            )

        except Exception as e:
            error_msg = f"风格分析异常: {e}"
            self.logger.error(error_msg)
            return f"[错误] {error_msg}"
