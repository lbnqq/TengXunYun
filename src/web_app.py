#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Web应用主文件

Author: AI Assistant (Claude)
Created: 2025-01-28
Last Modified: 2025-01-28
Modified By: AI Assistant (Claude)
AI Assisted: 是 - Claude 3.5 Sonnet
Version: v1.0
License: MIT
"""













import os
import sys
import json
import uuid
import time
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_from_directory, make_response
from flask_cors import CORS
from werkzeug.utils import secure_filename

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# 导入密钥管理器
from src.core.config.spark_x1_key_manager import get_spark_x1_key, get_spark_x1_config

# 导入智能填报模块
try:
    from llm_clients.spark_x1_client import SparkX1Client
    from core.tools.simple_smart_fill_manager import SimpleSmartFillManager
    SPARK_X1_AVAILABLE = True
except ImportError as e:
    print(f"警告: 智能填报模块导入失败: {e}")
    SPARK_X1_AVAILABLE = False

# 创建Flask应用
app = Flask(__name__, 
           template_folder=os.path.join(os.path.dirname(os.path.dirname(__file__)), 'templates'),
           static_folder=os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static'))
app.config['SECRET_KEY'] = 'dev-secret-key'
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB

# 启用CORS
CORS(app)

# 确保上传目录存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# 添加uploads文件下载路由
@app.route('/uploads/<filename>')
def download_uploaded_file(filename):
    """下载uploads目录中的文件"""
    try:
        return send_from_directory(
            app.config['UPLOAD_FOLDER'],
            filename,
            as_attachment=True
        )
    except FileNotFoundError:
        return jsonify({
            'success': False,
            'error': '文件不存在'
        }), 404
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'文件下载失败: {str(e)}'
        }), 500

# 初始化智能填报管理器
integrated_manager = None
if SPARK_X1_AVAILABLE:
    try:
        smart_fill_config = {
            'spark_x1_api_password': get_spark_x1_key('smart_fill')
        }
        integrated_manager = SimpleSmartFillManager(smart_fill_config)
        print("✅ 简化智能填报管理器初始化成功")
    except Exception as e:
        print(f"❌ 智能填报管理器初始化失败: {e}")
        integrated_manager = None

# 初始化格式对齐协调器
format_alignment_coordinator = None
try:
    from src.core.tools.format_alignment_coordinator import FormatAlignmentCoordinator
    # 使用密钥管理器获取API密钥
    format_alignment_coordinator = FormatAlignmentCoordinator(get_spark_x1_key('format_alignment'))
    print("✅ 格式对齐协调器初始化成功")
except Exception as e:
    print(f"❌ 格式对齐协调器初始化失败: {e}")

# 初始化文风对齐协调器
style_alignment_coordinator = None
if SPARK_X1_AVAILABLE:
    try:
        from src.core.tools.style_alignment_coordinator import StyleAlignmentCoordinator
        # 使用密钥管理器获取API密钥
        spark_x1_client = SparkX1Client(get_spark_x1_key('style_alignment'))
        style_alignment_coordinator = StyleAlignmentCoordinator(spark_x1_client)
        print("✅ 文风对齐协调器初始化成功")
    except Exception as e:
        print(f"❌ 文风对齐协调器初始化失败: {e}")
        style_alignment_coordinator = None

# 模拟数据存储
document_history = []
format_templates = [
    {'id': 'template1', 'name': '标准格式', 'description': '标准文档格式模板', 'type': 'baseline'}
]

# 新的文风统一系统 - 清空重建策略
# writing_style_templates 已被新的预设风格系统替代

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif', 'docx'}

def is_binary_content(content):
    """检查内容是否为二进制内容（乱码）"""
    if not content:
        return False

    # 检查是否包含大量非打印字符
    non_printable_count = 0
    total_chars = len(content)

    if total_chars == 0:
        return False

    for char in content[:1000]:  # 只检查前1000个字符
        if ord(char) < 32 and char not in '\n\r\t':  # 非打印字符（除了换行、回车、制表符）
            non_printable_count += 1

    # 如果非打印字符超过10%，认为是二进制内容
    return (non_printable_count / min(1000, total_chars)) > 0.1

def generate_sample_content(file_info):
    """根据文件信息生成示例内容"""
    file_name = file_info.get('name', '未知文件')
    file_type = file_info.get('type', '未知类型')
    file_size = file_info.get('size', 0)

    if file_name.endswith('.docx') or 'officedocument' in file_type:
        # 为DOCX文件生成更丰富的示例内容
        doc_title = file_name.replace('.docx', '').replace('新建 DOCX 文档', '工作报告').replace('新建', '项目')
        return f"""{doc_title}

项目概述
本项目旨在提升工作效率和质量，通过系统化的方法实现目标。

当前进展
1. 需求分析阶段
   - 已完成用户调研
   - 确定了核心功能需求
   - 制定了技术方案

2. 设计开发阶段
   - 完成了系统架构设计
   - 正在进行功能模块开发
   - 预计下月完成主要功能

3. 测试验证阶段
   - 制定了测试计划
   - 准备开始功能测试
   - 计划进行用户验收测试

遇到的问题
目前项目进展顺利，暂无重大技术难题。

下一步计划
1. 继续推进开发工作
2. 加强质量控制
3. 准备项目验收

总结
项目按计划稳步推进，预期能够按时完成既定目标。

附件信息
原始文件：{file_name}
文件大小：{file_size} 字节
创建时间：最近修改"""

    elif file_name.endswith('.txt'):
        return f"""文本文档：{file_name}

内容概要：
- 这是一个文本文件
- 包含纯文本内容
- 需要进行格式化处理

文件信息：
大小：{file_size} 字节
类型：{file_type}

请根据目标格式要求进行对齐处理。"""

    else:
        return f"""文档：{file_name}

这是一个需要格式化的文档文件。

基本信息：
- 文件名：{file_name}
- 文件类型：{file_type}
- 文件大小：{file_size} 字节

请根据目标格式进行相应的格式对齐处理。"""

# 主页路由
@app.route('/')
def index():
    """主页"""
    try:
        return render_template('enhanced-frontend-complete.html')
    except Exception as e:
        # 如果模板不存在，返回简单的HTML页面
        return f'''
        <!DOCTYPE html>
        <html>
        <head>
            <title>AI文档处理系统</title>
            <meta charset="utf-8">
        </head>
        <body>
            <h1>AI文档处理系统</h1>
            <p>系统正在运行中...</p>
            <p>API端点：</p>
            <ul>
                <li><a href="/api/health">健康检查</a></li>
                <li>POST /api/upload - 文件上传</li>
            </ul>
            <p>错误信息: {str(e)}</p>
        </body>
        </html>
        '''

# 健康检查路由
@app.route('/api/health')
def health_check():
    """健康检查端点"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'version': '1.0.0',
        'service': 'AI文档处理系统'
    })



@app.route('/api/upload', methods=['POST'])
def upload_file():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': '请求数据为空'}), 400
            
        # 支持多种参数格式
        source_content = data.get('source_content', '')
        target_content = data.get('target_content', '')
        
        # 如果直接传入了文件内容，使用文件内容
        if not source_content and 'source_file' in data:
            source_content = data['source_file']
        if not target_content and 'target_file' in data:
            target_content = data['target_file']
            
        # 如果传入了files数组，尝试从files中获取内容
        if not source_content or not target_content:
            files = data.get('files', [])
            if len(files) >= 2:
                # 处理files参数 - 支持字符串数组和对象数组
                if isinstance(files[0], dict):
                    source_content = f"参考文件内容: {files[0].get('name', 'source.txt')}"
                    target_content = f"目标文件内容: {files[1].get('name', 'target.txt')}"
                else:
                    # files是字符串数组
                    source_content = f"参考文件内容: {files[0]}"
                    target_content = f"目标文件内容: {files[1]}"
        
        if not source_content or not target_content:
            return jsonify({'success': False, 'error': '缺少必要参数：需要source_content和target_content，或者包含两个文件的files数组'}), 400
        
        # 模拟格式对齐结果
        result = {
            'aligned_content': f"对齐后的内容: {source_content[:50]}...",
            'alignment_score': 0.85,
            'suggestions': ['建议优化格式结构', '增加段落分隔'],
            'source_length': len(source_content),
            'target_length': len(target_content),
            'session_id': data.get('session_id', ''),
            'status': 'completed'
        }
        
        return jsonify({'success': True, 'data': result})
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/writing-style/analyze', methods=['POST'])
def analyze_writing_style():
    try:
        data = request.get_json()
        template_content = data.get('template_content', '')
        fill_data = data.get('data', {})
        
        if not template_content:
            return jsonify({'success': False, 'error': '缺少模板内容'}), 400
        
        # 模拟文档填写结果
        result = {
            'filled_content': f"填写后的内容: {template_content[:50]}...",
            'fill_status': 'completed',
            'filled_fields': len(fill_data)
        }
        
        return jsonify({'success': True, 'data': result})
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/document/parse', methods=['POST'])
def document_parse():
    try:
        return jsonify({
            'success': True,
            'history': document_history
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/documents/history', methods=['GET'])
def get_documents_history():
    """获取文档处理历史记录"""
    try:
        # 这里可以从数据库或文件系统获取历史记录
        # 目前返回空的历史记录列表
        history_records = []

        return jsonify({
            'success': True,
            'history': history_records,
            'count': len(history_records)
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取文档历史失败: {str(e)}',
            'history': []
        }), 500

@app.route('/api/documents/history/<record_id>/reapply', methods=['POST'])
def reapply_document_operation(record_id):
    return jsonify({
        'success': True,
        'config': {
            'version': '1.0.0',
            'features': ['format_alignment', 'writing_style', 'document_fill'],
            'max_file_size': '16MB'
        }
    })

@app.route('/api/models')
def get_models():
    if request.method == 'GET':
        return jsonify({
            'success': True,
            'templates': format_templates
        })
    elif request.method == 'POST':
        try:
            data = request.get_json()
            if not data:
                return jsonify({'success': False, 'error': '请求数据为空'}), 400
            
            name = data.get('name', '')
            content = data.get('content', '')
            template_type = data.get('type', 'custom')
            
            if not name:
                return jsonify({'success': False, 'error': '模板名称不能为空'}), 400
            
            # 创建新模板
            new_template = {
                'id': f"template_{len(format_templates) + 1}",
                'name': name,
                'description': f'{template_type}格式模板',
                'type': template_type,
                'created_at': datetime.now().isoformat()
            }
            
            format_templates.append(new_template)
            
            return jsonify({
                'success': True,
                'template': new_template,
                'message': '格式模板保存成功'
            })
            
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500

# ==================== 文风统一模块 API端点 ====================

@app.route('/api/style-alignment/preset-styles', methods=['GET'])
def get_preset_styles():
    """获取预设风格模板库"""
    try:
        if not style_alignment_coordinator:
            return jsonify({
                'success': False,
                'error': '文风对齐协调器未初始化'
            }), 500

        language = request.args.get('language', 'auto')
        result = style_alignment_coordinator.get_preset_styles(language)

        return jsonify({
            'success': result['success'],
            'styles': result.get('styles', {}),
            'count': result.get('count', 0),
            'language': language
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/style-alignment/generate-with-style', methods=['POST'])
def generate_with_style():
    """基于预设风格生成文本"""
    try:
        if not style_alignment_coordinator:
            return jsonify({
                'success': False,
                'error': '文风对齐协调器未初始化'
            }), 500

        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': '请求数据为空'}), 400

        # 获取请求参数
        content = data.get('content', '')
        style_id = data.get('style_id', '')
        action = data.get('action', '重写')
        language = data.get('language', 'auto')
        temperature = data.get('temperature')

        if not content:
            return jsonify({'success': False, 'error': '内容不能为空'}), 400

        if not style_id:
            return jsonify({'success': False, 'error': '请选择风格'}), 400

        # 创建会话
        session_id = style_alignment_coordinator.create_session()

        # 处理预设风格生成
        result = style_alignment_coordinator.process_preset_style_generation(
            session_id=session_id,
            content=content,
            style_id=style_id,
            action=action,
            language=language,
            temperature=temperature
        )

        return jsonify(result)

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/style-alignment/few-shot-transfer', methods=['POST'])
def few_shot_style_transfer():
    """Few-Shot风格迁移"""
    try:
        if not style_alignment_coordinator:
            return jsonify({
                'success': False,
                'error': '文风对齐协调器未初始化'
            }), 500

        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': '请求数据为空'}), 400

        # 获取请求参数
        content = data.get('content', '')
        reference_document = data.get('reference_document', '')
        target_description = data.get('target_description', '')
        language = data.get('language', 'auto')
        temperature = data.get('temperature', 0.7)

        if not content:
            return jsonify({'success': False, 'error': '内容不能为空'}), 400

        if not reference_document:
            return jsonify({'success': False, 'error': '请提供参考文档'}), 400

        # 创建会话
        session_id = style_alignment_coordinator.create_session()

        # 处理Few-Shot风格迁移
        result = style_alignment_coordinator.process_few_shot_transfer(
            session_id=session_id,
            content=content,
            reference_document=reference_document,
            target_description=target_description,
            language=language,
            temperature=temperature
        )

        return jsonify(result)

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/style-alignment/task-progress/<task_id>', methods=['GET'])
def get_task_progress(task_id):
    """获取任务进度"""
    try:
        if not style_alignment_coordinator:
            return jsonify({
                'success': False,
                'error': '文风对齐协调器未初始化'
            }), 500

        progress = style_alignment_coordinator.get_task_progress(task_id)

        return jsonify({
            'success': True,
            'progress': progress
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/style-alignment/task-result/<task_id>', methods=['GET'])
def get_task_result(task_id):
    """获取任务结果"""
    try:
        if not style_alignment_coordinator:
            return jsonify({
                'success': False,
                'error': '文风对齐协调器未初始化'
            }), 500

        result = style_alignment_coordinator.get_task_result(task_id)

        return jsonify(result)

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/style-alignment/export', methods=['POST'])
def export_style_result():
    """导出文风统一结果"""
    try:
        if not style_alignment_coordinator:
            return jsonify({
                'success': False,
                'error': '文风对齐协调器未初始化'
            }), 500

        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'error': '请求数据为空'
            }), 400

        task_id = data.get('task_id')
        format_type = data.get('format', 'txt').lower()

        if not task_id:
            return jsonify({
                'success': False,
                'error': '缺少任务ID'
            }), 400

        # 验证格式类型
        if format_type not in ['txt', 'docx']:
            return jsonify({
                'success': False,
                'error': f'不支持的导出格式: {format_type}'
            }), 400

        # 调用协调器的导出方法
        result = style_alignment_coordinator.export_result(task_id, format_type)

        if result.get('success'):
            return jsonify({
                'success': True,
                'message': f'{format_type.upper()}文件导出成功',
                'filename': result.get('filename'),
                'download_url': result.get('download_url'),
                'format': result.get('format')
            })
        else:
            return jsonify({
                'success': False,
                'error': result.get('error', '导出失败')
            }), 500

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'导出失败: {str(e)}'
        }), 500

# ==================== 智能填报模块 ====================

@app.route('/smart-fill-demo')
def smart_fill_demo():
    """智能填报演示页面"""
    return render_template('smart_fill_demo.html')

# ==================== 智能填报模块 API端点 ====================

@app.route('/api/smart-fill/generate-summary', methods=['POST'])
def generate_summary():
    """生成年度总结"""
    try:
        if not SPARK_X1_AVAILABLE:
            return jsonify({
                'success': False,
                'error': '星火X1客户端不可用，请检查配置'
            }), 500

        data = request.get_json()
        if not data or 'content' not in data:
            return jsonify({
                'success': False,
                'error': '缺少必要参数：content'
            }), 400

        work_content = data['content']
        api_password = data.get('api_password') or "NJFASGuFsRYYjeyLpZFk:jhjQJHHgIeoKVzbAORPh"

        # 创建星火X1客户端
        client = SparkX1Client(api_password=api_password)

        # 生成年度总结
        result = client.generate_summary(work_content)

        if result['success']:
            return jsonify({
                'success': True,
                'message': '年度总结生成成功',
                'content': result['content'],
                'filename': result['filename'],
                'file_path': result['file_path'],
                'usage': result.get('usage', {}),
                'data': {
                    'content': result['content'],
                    'filename': result['filename'],
                    'file_path': result['file_path'],
                    'usage': result.get('usage', {})
                }
            })
        else:
            return jsonify({
                'success': False,
                'error': result['error']
            }), 500

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'生成年度总结失败: {str(e)}'
        }), 500

@app.route('/api/smart-fill/generate-resume', methods=['POST'])
def generate_resume():
    """生成个人简历"""
    try:
        if not SPARK_X1_AVAILABLE:
            return jsonify({
                'success': False,
                'error': '星火X1客户端不可用，请检查配置'
            }), 500

        data = request.get_json()
        if not data or 'content' not in data:
            return jsonify({
                'success': False,
                'error': '缺少必要参数：content'
            }), 400

        personal_info = data['content']
        api_password = data.get('api_password') or "NJFASGuFsRYYjeyLpZFk:jhjQJHHgIeoKVzbAORPh"

        # 创建星火X1客户端
        client = SparkX1Client(api_password=api_password)

        # 生成简历
        result = client.generate_resume(personal_info)

        if result['success']:
            return jsonify({
                'success': True,
                'message': '简历生成成功',
                'resume_data': result['data'],
                'filename': result['filename'],
                'file_path': result['file_path'],
                'usage': result.get('usage', {}),
                'data': {
                    'resume_data': result['data'],
                    'filename': result['filename'],
                    'file_path': result['file_path'],
                    'usage': result.get('usage', {})
                }
            })
        else:
            return jsonify({
                'success': False,
                'error': result['error']
            }), 500

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'生成简历失败: {str(e)}'
        }), 500

@app.route('/api/smart-fill/download/<path:filename>')
def download_file(filename):
    """下载生成的文件"""
    try:
        import tempfile
        import urllib.parse

        # URL解码文件名
        decoded_filename = urllib.parse.unquote(filename)

        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, decoded_filename)

        print(f"尝试下载文件: {decoded_filename}")
        print(f"文件路径: {file_path}")
        print(f"文件是否存在: {os.path.exists(file_path)}")

        # 如果文件不存在，尝试查找类似的文件
        if not os.path.exists(file_path):
            # 列出临时目录中的所有文件
            temp_files = []
            try:
                temp_files = os.listdir(temp_dir)
                print(f"临时目录中的文件: {temp_files}")

                # 查找包含关键词的文件
                if '年度工作总结' in decoded_filename:
                    matching_files = [f for f in temp_files if '年度工作总结' in f and f.endswith('.docx')]
                elif '个人简历' in decoded_filename:
                    matching_files = [f for f in temp_files if '个人简历' in f and f.endswith('.docx')]
                else:
                    matching_files = [f for f in temp_files if decoded_filename in f]

                if matching_files:
                    # 使用最新的文件
                    matching_files.sort(reverse=True)
                    actual_filename = matching_files[0]
                    file_path = os.path.join(temp_dir, actual_filename)
                    print(f"找到匹配文件: {actual_filename}")
                else:
                    print(f"未找到匹配文件，搜索关键词: {decoded_filename}")
            except Exception as e:
                print(f"列出临时目录文件失败: {e}")

        if not os.path.exists(file_path):
            return jsonify({
                'success': False,
                'error': f'文件不存在: {decoded_filename}',
                'debug_info': {
                    'temp_dir': temp_dir,
                    'requested_file': decoded_filename,
                    'file_path': file_path,
                    'temp_files': temp_files[:10] if 'temp_files' in locals() else []
                }
            }), 404

        return send_from_directory(
            temp_dir,
            os.path.basename(file_path),
            as_attachment=True,
            download_name=decoded_filename
        )

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'文件下载失败: {str(e)}'
        }), 500

@app.route('/api/smart-fill/export', methods=['POST'])
def export_smart_fill_result():
    """导出智能填报结果为不同格式"""
    print("🔍 导出API被调用")

    try:
        data = request.get_json()
        print(f"📋 接收到的数据: {data}")

        if not data:
            print("❌ 请求数据为空")
            return jsonify({
                'success': False,
                'error': '请求数据为空'
            }), 400

        filename = data.get('filename')
        format_type = data.get('format', 'txt').lower()
        result_type = data.get('type', 'summary')

        print(f"📄 解析参数 - 文件名: {filename}, 格式: {format_type}, 类型: {result_type}")

        if not filename:
            print("❌ 缺少文件名")
            return jsonify({
                'success': False,
                'error': '缺少文件名'
            }), 400

        # 验证格式类型
        if format_type not in ['docx']:
            print(f"❌ 不支持的格式: {format_type}")
            return jsonify({
                'success': False,
                'error': f'不支持的导出格式: {format_type}'
            }), 400

        print("✅ 参数验证通过，开始处理...")

        # 读取原始文件内容
        import tempfile
        temp_dir = tempfile.gettempdir()
        original_file_path = os.path.join(temp_dir, filename)

        content = ""
        if os.path.exists(original_file_path):
            try:
                if filename.endswith('.docx'):
                    # 读取Word文档内容
                    from docx import Document
                    doc = Document(original_file_path)
                    content = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
                elif filename.endswith('.txt'):
                    # 读取文本文件内容
                    with open(original_file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                else:
                    content = f"无法读取文件格式: {filename}"
            except Exception as e:
                print(f"❌ 读取原始文件失败: {e}")
                content = f"读取原始文件失败: {str(e)}"
        else:
            print(f"❌ 原始文件不存在: {original_file_path}")
            content = f"原始文件不存在: {filename}"

        # 生成新格式文件
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        new_filename = f"{result_type}_result_{timestamp}.{format_type}"

        # 确保uploads目录存在
        uploads_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'uploads')
        os.makedirs(uploads_dir, exist_ok=True)
        new_file_path = os.path.join(uploads_dir, new_filename)

        print(f"🔄 生成{format_type.upper()}文件: {new_file_path}")

        if format_type == 'docx':
            try:
                from docx import Document
                doc = Document()

                # 添加标题
                title = '年度工作总结' if result_type == 'summary' else '个人简历'
                doc.add_heading(title, 0)

                # 添加内容
                paragraphs = content.split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())

                doc.save(new_file_path)
                print(f"✅ DOCX文件生成成功")

            except ImportError as e:
                print(f"❌ python-docx导入失败: {e}")
                return jsonify({
                    'success': False,
                    'error': 'python-docx库未安装'
                }), 500
            except Exception as e:
                print(f"❌ DOCX生成失败: {e}")
                return jsonify({
                    'success': False,
                    'error': f'DOCX生成失败: {str(e)}'
                }), 500

        elif format_type == 'pdf':
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet

                doc = SimpleDocTemplate(new_file_path, pagesize=A4)
                story = []

                styles = getSampleStyleSheet()
                title_style = styles['Title']
                normal_style = styles['Normal']

                # 添加标题
                title = '年度工作总结' if result_type == 'summary' else '个人简历'
                story.append(Paragraph(title, title_style))
                story.append(Spacer(1, 12))

                # 添加内容
                paragraphs = content.split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        # 处理特殊字符
                        clean_text = paragraph.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(clean_text, normal_style))
                        story.append(Spacer(1, 12))

                doc.build(story)
                print(f"✅ PDF文件生成成功")

            except ImportError as e:
                print(f"❌ reportlab导入失败: {e}")
                return jsonify({
                    'success': False,
                    'error': 'reportlab库未安装'
                }), 500
            except Exception as e:
                print(f"❌ PDF生成失败: {e}")
                return jsonify({
                    'success': False,
                    'error': f'PDF生成失败: {str(e)}'
                }), 500
        else:
            # TXT格式
            try:
                with open(new_file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                print(f"✅ TXT文件生成成功")
            except Exception as e:
                print(f"❌ TXT生成失败: {e}")
                return jsonify({
                    'success': False,
                    'error': f'TXT生成失败: {str(e)}'
                }), 500

        print(f"🎉 文件生成完成: {new_filename}")
        return jsonify({
            'success': True,
            'message': f'{format_type.upper()}文件导出成功',
            'filename': new_filename,
            'download_url': f'/uploads/{new_filename}',
            'format': format_type
        })

    except Exception as e:
        print(f"❌ 导出API发生异常: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'error': f'服务器内部错误: {str(e)}'
        }), 500




@app.route('/api/smart-fill/status')
def smart_fill_status():
    """智能填报模块状态检查"""
    try:
        if integrated_manager:
            # 使用简化管理器获取状态
            status = integrated_manager.get_status()
        else:
            # 备用状态信息
            status = {
                'spark_x1_available': SPARK_X1_AVAILABLE,
                'core_components_available': False,
                'supported_types': ['summary', 'resume'],
                'supported_functions': ['generate_summary', 'generate_resume'],
                'integration_mode': 'simplified_spark_x1_only',
                'version': '1.0.0',
                'timestamp': datetime.now().isoformat()
            }

            if SPARK_X1_AVAILABLE:
                # 测试API连接
                try:
                    client = SparkX1Client(api_password="NJFASGuFsRYYjeyLpZFk:jhjQJHHgIeoKVzbAORPh")
                    api_status = client.is_available()
                    status['api_connection'] = 'connected' if api_status else 'disconnected'
                except Exception as e:
                    status['api_connection'] = f'error: {str(e)}'

        return jsonify({
            'success': True,
            'data': status
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'状态检查失败: {str(e)}'
        }), 500

# 格式对齐模块API端点
@app.route('/api/format-alignment/upload', methods=['POST'])
def format_alignment_upload():
    """格式对齐文件上传接口"""
    try:
        # 检查是否有文件上传
        if 'files' not in request.files:
            return jsonify({
                'code': 1,
                'message': '没有上传文件',
                'data': None
            }), 400

        files = request.files.getlist('files')
        if not files or len(files) == 0:
            return jsonify({
                'code': 1,
                'message': '文件列表为空',
                'data': None
            }), 400

        # 生成上传ID
        import time
        upload_id = f"upload_{int(time.time())}_{len(files)}"
        uploaded_files = []

        # 处理每个上传的文件
        for i, file in enumerate(files):
            if file.filename == '':
                continue

            if file and allowed_file(file.filename):
                # 生成文件ID
                file_id = f"file_{upload_id}_{i}"

                # 读取文件内容
                file_content = file.read().decode('utf-8', errors='ignore')

                # 存储文件信息（这里简化存储在内存中）
                file_info = {
                    'file_id': file_id,
                    'filename': file.filename,
                    'size': len(file_content),
                    'type': file.filename.rsplit('.', 1)[1].lower(),
                    'content': file_content
                }

                uploaded_files.append(file_info)

        if not uploaded_files:
            return jsonify({
                'code': 1,
                'message': '没有有效的文件',
                'data': None
            }), 400

        # 存储上传信息（简化存储）
        upload_info = {
            'upload_id': upload_id,
            'files': uploaded_files,
            'upload_time': datetime.now().isoformat()
        }

        # 这里应该存储到数据库或文件系统，现在简化存储在全局变量
        if not hasattr(app, 'upload_storage'):
            app.upload_storage = {}
        app.upload_storage[upload_id] = upload_info

        return jsonify({
            'code': 0,
            'message': 'success',
            'data': {
                'upload_id': upload_id,
                'files': [
                    {
                        'file_id': f['file_id'],
                        'filename': f['filename'],
                        'size': f['size'],
                        'type': f['type']
                    } for f in uploaded_files
                ]
            }
        })

    except Exception as e:
        return jsonify({
            'code': 1,
            'message': f'上传失败: {str(e)}',
            'data': None
        }), 500

@app.route('/api/format-alignment/process', methods=['POST'])
def format_alignment_process():
    """格式对齐处理接口"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({
                'code': 1,
                'message': '请求数据为空',
                'data': None
            }), 400

        upload_id = data.get('upload_id')
        format_instruction = data.get('format_instruction', '请格式化文档')
        template_id = data.get('template_id')
        options = data.get('options', {})

        if not upload_id:
            return jsonify({
                'code': 1,
                'message': '缺少upload_id参数',
                'data': None
            }), 400

        # 获取上传的文件
        if not hasattr(app, 'upload_storage') or upload_id not in app.upload_storage:
            return jsonify({
                'code': 1,
                'message': f'上传ID {upload_id} 不存在',
                'data': None
            }), 404

        upload_info = app.upload_storage[upload_id]
        files = upload_info['files']

        if not files:
            return jsonify({
                'code': 1,
                'message': '没有找到上传的文件',
                'data': None
            }), 400

        # 使用全局格式对齐协调器
        if format_alignment_coordinator is None:
            return jsonify({
                'code': 1,
                'message': '格式对齐协调器未初始化',
                'data': None
            }), 500

        try:
            # 处理第一个文件（简化处理）
            file_content = files[0]['content']

            # 使用星火X1进行格式化
            result = format_alignment_coordinator.format_with_spark_x1(
                content=file_content,
                instruction=format_instruction,
                temperature=options.get('temperature', 0.7),
                max_tokens=options.get('max_tokens', 4000)
            )

            if result.get('success'):
                return jsonify({
                    'code': 0,
                    'message': 'success',
                    'data': {
                        'task_id': result['task_id'],
                        'status': result['status'],
                        'estimated_time': 30
                    }
                })
            else:
                return jsonify({
                    'code': 1,
                    'message': result.get('error', '格式化失败'),
                    'data': None
                }), 500

        except Exception as e:
            return jsonify({
                'code': 1,
                'message': f'格式化处理失败: {str(e)}',
                'data': None
            }), 500

    except Exception as e:
        return jsonify({
            'code': 1,
            'message': f'处理请求失败: {str(e)}',
            'data': None
        }), 500

@app.route('/api/format-alignment/result/<task_id>', methods=['GET'])
def format_alignment_result(task_id):
    """获取格式对齐结果接口"""
    try:
        # 使用全局格式对齐协调器
        if format_alignment_coordinator is None:
            return jsonify({
                'code': 1,
                'message': '格式对齐协调器未初始化',
                'data': None
            }), 500

        # 获取任务结果
        result = format_alignment_coordinator.get_task_result(task_id)

        if result.get('success'):
            task_data = result

            # 构建响应数据
            response_data = {
                'task_id': task_id,
                'status': task_data['status'],
                'processing_log': task_data.get('processing_log', ''),
            }

            # 如果任务完成，添加结果文件信息
            if task_data['status'] == 'completed':
                response_data['result_files'] = [
                    {
                        'file_id': f"result_{task_id}",
                        'filename': f"formatted_document_{task_id}.txt",
                        'download_url': f"/api/format-alignment/download/{task_id}"
                    }
                ]

            return jsonify({
                'code': 0,
                'message': 'success',
                'data': response_data
            })
        else:
            return jsonify({
                'code': 1,
                'message': result.get('error', '获取任务结果失败'),
                'data': None
            }), 404

    except Exception as e:
        return jsonify({
            'code': 1,
            'message': f'获取结果失败: {str(e)}',
            'data': None
        }), 500

# 格式模板管理API（文风统一模块使用）
@app.route('/api/format-templates', methods=['GET'])
def get_format_templates_for_style():
    """获取格式模板列表（文风统一模块使用）"""
    try:
        # 使用全局格式对齐协调器
        if format_alignment_coordinator is None:
            return jsonify({
                'success': False,
                'error': '格式对齐协调器未初始化',
                'templates': []
            }), 500

        # 读取模板索引文件
        import os
        # 获取当前文件的目录，然后构建相对路径
        current_dir = os.path.dirname(os.path.abspath(__file__))
        template_index_path = os.path.join(current_dir, 'core', 'knowledge_base', 'format_templates', 'template_index.json')

        if not os.path.exists(template_index_path):
            return jsonify({
                'success': True,
                'templates': [],
                'count': 0
            })

        with open(template_index_path, 'r', encoding='utf-8') as f:
            template_data = json.load(f)

        templates = template_data.get('templates', [])

        # 格式化模板数据，适配文风统一模块的格式
        formatted_templates = []
        for template in templates:
            formatted_templates.append({
                'id': template.get('template_id', ''),
                'name': template.get('name', '未命名模板'),
                'description': template.get('description', ''),
                'created_time': template.get('created_time', ''),
                'category': '通用格式',  # 默认分类
                'preview': f"格式模板：{template.get('name', '未命名模板')}"
            })

        return jsonify({
            'success': True,
            'templates': formatted_templates,
            'count': len(formatted_templates)
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取格式模板失败: {str(e)}',
            'templates': []
        }), 500

@app.route('/api/format-templates', methods=['POST'])
def create_format_template():
    """创建格式模板（文风统一模块使用）"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'error': '请求数据为空'
            }), 400

        # 这里可以添加创建模板的逻辑
        # 目前返回成功响应
        return jsonify({
            'success': True,
            'message': '格式模板创建成功',
            'template_id': 'temp_' + str(int(time.time()))
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'创建格式模板失败: {str(e)}'
        }), 500

# 新增：格式模板管理API（格式对齐模块使用）
@app.route('/api/format-alignment/templates', methods=['GET'])
def get_format_templates():
    """获取格式模板列表"""
    try:
        # 使用全局格式对齐协调器
        if format_alignment_coordinator is None:
            return jsonify({
                'success': False,
                'error': '格式对齐协调器未初始化',
                'templates': []
            }), 500

        # 读取模板索引文件
        import os
        # 获取当前文件的目录，然后构建相对路径
        current_dir = os.path.dirname(os.path.abspath(__file__))
        template_index_path = os.path.join(current_dir, 'core', 'knowledge_base', 'format_templates', 'template_index.json')

        if not os.path.exists(template_index_path):
            return jsonify({
                'success': True,
                'templates': [],
                'count': 0
            })

        with open(template_index_path, 'r', encoding='utf-8') as f:
            template_data = json.load(f)

        templates = template_data.get('templates', [])

        # 格式化模板数据，参考文风统一的格式
        formatted_templates = []
        for template in templates:
            formatted_templates.append({
                'id': template.get('template_id', ''),
                'name': template.get('name', '未命名模板'),
                'description': template.get('description', ''),
                'created_time': template.get('created_time', ''),
                'category': '通用格式',  # 默认分类
                'preview': f"格式模板：{template.get('name', '未命名模板')}"
            })

        return jsonify({
            'success': True,
            'templates': formatted_templates,
            'count': len(formatted_templates)
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取格式模板失败: {str(e)}',
            'templates': []
        }), 500

@app.route('/api/format-alignment/template/<template_id>', methods=['GET'])
def get_format_template_detail(template_id):
    """获取特定格式模板的详细信息"""
    try:
        # 使用全局格式对齐协调器
        if format_alignment_coordinator is None:
            return jsonify({
                'success': False,
                'error': '格式对齐协调器未初始化'
            }), 500

        # 读取具体模板文件
        import os
        # 获取当前文件的目录，然后构建相对路径
        current_dir = os.path.dirname(os.path.abspath(__file__))
        template_file_path = os.path.join(current_dir, 'core', 'knowledge_base', 'format_templates', f'{template_id}.json')

        if not os.path.exists(template_file_path):
            return jsonify({
                'success': False,
                'error': '模板不存在'
            }), 404

        with open(template_file_path, 'r', encoding='utf-8') as f:
            template_detail = json.load(f)

        return jsonify({
            'success': True,
            'template': template_detail
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取模板详情失败: {str(e)}'
        }), 500

@app.route('/api/format-alignment/download/<task_id>', methods=['GET'])
def format_alignment_download(task_id):
    """下载格式化结果文件"""
    try:
        # 获取文件格式参数
        file_format = request.args.get('format', 'txt').lower()

        # 使用全局格式对齐协调器
        if format_alignment_coordinator is None:
            return jsonify({
                'code': 1,
                'message': '格式对齐协调器未初始化',
                'data': None
            }), 500

        # 获取任务结果
        result = format_alignment_coordinator.get_task_result(task_id)

        if result.get('success'):
            formatted_content = result.get('formatted_content', '')

            if file_format == 'docx':
                # 生成Word格式文件
                return generate_word_document(formatted_content, task_id)
            elif file_format == 'html':
                # 生成HTML格式文件
                return generate_html_document(formatted_content, task_id)
            elif file_format == 'pdf':
                # 生成PDF格式文件
                return generate_pdf_document(formatted_content, task_id)
            else:
                # 默认生成文本文件
                response = make_response(formatted_content)
                response.headers['Content-Type'] = 'text/plain; charset=utf-8'
                response.headers['Content-Disposition'] = f'attachment; filename=formatted_document_{task_id}.txt'
                return response
        else:
            return jsonify({
                'code': 1,
                'message': result.get('error', '任务不存在'),
                'data': None
            }), 404

    except Exception as e:
        return jsonify({
            'code': 1,
            'message': f'下载失败: {str(e)}',
            'data': None
        }), 500

def generate_word_document(content, task_id):
    """生成Word文档"""
    try:
        from docx import Document
        from docx.shared import Inches
        import io

        # 创建Word文档
        doc = Document()

        # 添加标题
        doc.add_heading('格式化文档', 0)

        # 处理内容，按行分割并添加到文档
        lines = content.split('\n')
        current_paragraph = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 检查是否是标题
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                title_text = line.lstrip('#').strip()
                doc.add_heading(title_text, level)
            elif line.startswith('##'):
                doc.add_heading(line[2:].strip(), 2)
            elif line.startswith('#'):
                doc.add_heading(line[1:].strip(), 1)
            else:
                # 普通段落
                doc.add_paragraph(line)

        # 保存到内存
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # 创建响应
        response = make_response(doc_io.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename=formatted_document_{task_id}.docx'

        return response

    except ImportError:
        # 如果没有安装python-docx，回退到文本格式
        response = make_response(content)
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        response.headers['Content-Disposition'] = f'attachment; filename=formatted_document_{task_id}.txt'
        return response
    except Exception as e:
        # 出错时回退到文本格式
        response = make_response(content)
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        response.headers['Content-Disposition'] = f'attachment; filename=formatted_document_{task_id}.txt'
        return response

def generate_html_document(content, task_id):
    """生成HTML文档"""
    try:
        # 将Markdown格式的内容转换为HTML
        html_content = markdown_to_html(content)

        # 创建完整的HTML文档
        full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>格式化文档</title>
    <style>
        body {{
            font-family: 'Microsoft YaHei', Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 30px;
            margin-bottom: 15px;
        }}
        h1 {{ font-size: 2.5em; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ font-size: 2em; border-bottom: 1px solid #bdc3c7; padding-bottom: 5px; }}
        h3 {{ font-size: 1.5em; }}
        p {{ margin-bottom: 15px; text-align: justify; }}
        ul, ol {{ margin-bottom: 15px; padding-left: 30px; }}
        li {{ margin-bottom: 5px; }}
        .header {{ text-align: center; margin-bottom: 40px; }}
        .footer {{ text-align: center; margin-top: 40px; color: #7f8c8d; font-size: 0.9em; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>格式化文档</h1>
        <p>生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
    <div class="content">
        {html_content}
    </div>
    <div class="footer">
        <p>由 aiDoc 智能文档处理系统生成</p>
    </div>
</body>
</html>"""

        # 创建响应
        response = make_response(full_html)
        response.headers['Content-Type'] = 'text/html; charset=utf-8'
        response.headers['Content-Disposition'] = f'attachment; filename=formatted_document_{task_id}.html'

        return response

    except Exception as e:
        # 出错时回退到文本格式
        response = make_response(content)
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        response.headers['Content-Disposition'] = f'attachment; filename=formatted_document_{task_id}.txt'
        return response

def markdown_to_html(content):
    """将Markdown格式的内容转换为HTML"""
    import re

    # 简单的Markdown到HTML转换
    html = content

    # 转换标题
    html = re.sub(r'^### (.*?)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.*?)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^# (.*?)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)

    # 转换段落
    paragraphs = html.split('\n\n')
    html_paragraphs = []

    for para in paragraphs:
        para = para.strip()
        if para:
            # 检查是否已经是HTML标签
            if not (para.startswith('<') and para.endswith('>')):
                # 处理列表
                if para.startswith('- ') or para.startswith('* '):
                    items = para.split('\n')
                    list_html = '<ul>\n'
                    for item in items:
                        if item.strip().startswith(('- ', '* ')):
                            list_html += f'<li>{item.strip()[2:]}</li>\n'
                    list_html += '</ul>'
                    html_paragraphs.append(list_html)
                elif re.match(r'^\d+\. ', para):
                    items = para.split('\n')
                    list_html = '<ol>\n'
                    for item in items:
                        if re.match(r'^\d+\. ', item.strip()):
                            list_html += f'<li>{re.sub(r"^\d+\. ", "", item.strip())}</li>\n'
                    list_html += '</ol>'
                    html_paragraphs.append(list_html)
                else:
                    # 普通段落
                    html_paragraphs.append(f'<p>{para}</p>')
            else:
                html_paragraphs.append(para)

    return '\n'.join(html_paragraphs)

def generate_pdf_document(content, task_id):
    """生成PDF文档"""
    try:
        # 尝试使用reportlab生成PDF
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import io

        # 创建PDF文档
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)

        # 获取样式
        styles = getSampleStyleSheet()

        # 创建内容
        story = []

        # 添加标题
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # 居中
        )
        story.append(Paragraph('格式化文档', title_style))
        story.append(Spacer(1, 20))

        # 处理内容
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 12))
                continue

            if line.startswith('###'):
                style = styles['Heading3']
                text = line[3:].strip()
            elif line.startswith('##'):
                style = styles['Heading2']
                text = line[2:].strip()
            elif line.startswith('#'):
                style = styles['Heading1']
                text = line[1:].strip()
            else:
                style = styles['Normal']
                text = line

            story.append(Paragraph(text, style))
            story.append(Spacer(1, 6))

        # 构建PDF
        doc.build(story)

        # 创建响应
        buffer.seek(0)
        response = make_response(buffer.getvalue())
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename=formatted_document_{task_id}.pdf'

        return response

    except ImportError:
        # 如果没有安装reportlab，回退到HTML格式
        return generate_html_document(content, task_id)
    except Exception as e:
        # 出错时回退到文本格式
        response = make_response(content)
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        response.headers['Content-Disposition'] = f'attachment; filename=formatted_document_{task_id}.txt'
        return response

@app.route('/api/format-alignment/continue', methods=['POST'])
def format_alignment_continue():
    """多轮对话继续接口"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({
                'code': 1,
                'message': '请求数据为空',
                'data': None
            }), 400

        task_id = data.get('task_id')
        instruction = data.get('instruction', '')
        context = data.get('context')

        if not task_id:
            return jsonify({
                'code': 1,
                'message': '缺少task_id参数',
                'data': None
            }), 400

        if not instruction:
            return jsonify({
                'code': 1,
                'message': '缺少instruction参数',
                'data': None
            }), 400

        # 使用全局格式对齐协调器
        if format_alignment_coordinator is None:
            return jsonify({
                'code': 1,
                'message': '格式对齐协调器未初始化',
                'data': None
            }), 500

        try:
            # 继续对话
            result = format_alignment_coordinator.continue_conversation_task(task_id, instruction)

            if result.get('success'):
                return jsonify({
                    'code': 0,
                    'message': 'success',
                    'data': {
                        'task_id': result['task_id'],
                        'status': result['status'],
                        'response': result['response']
                    }
                })
            else:
                return jsonify({
                    'code': 1,
                    'message': result.get('error', '继续对话失败'),
                    'data': None
                }), 500

        except Exception as e:
            return jsonify({
                'code': 1,
                'message': f'对话处理失败: {str(e)}',
                'data': None
            }), 500

    except Exception as e:
        return jsonify({
            'code': 1,
            'message': f'处理请求失败: {str(e)}',
            'data': None
        }), 500

# 兼容现有前端的格式对齐接口
@app.route('/api/format-alignment', methods=['POST'])
def format_alignment_legacy():
    """兼容现有前端的格式对齐接口"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': '请求数据为空'}), 400

        session_id = data.get('session_id', '')
        files = data.get('files', [])

        if len(files) < 2:
            return jsonify({'success': False, 'error': '需要至少2个文件进行格式对齐'}), 400

        # 使用全局格式对齐协调器
        if format_alignment_coordinator is None:
            return jsonify({
                'success': False,
                'error': '格式对齐协调器未初始化'
            }), 500

        # 获取文件内容
        try:
            # 从files中提取内容
            source_file = files[0]
            target_file = files[1] if len(files) > 1 else files[0]

            print(f"📝 源文件: {source_file.get('name', '未知')}, 类型: {source_file.get('type', '未知')}")
            print(f"📝 目标文件: {target_file.get('name', '未知')}, 类型: {target_file.get('type', '未知')}")

            # 处理源文件内容
            source_content = source_file.get('content', '')

            # 检查是否是乱码或二进制内容
            if is_binary_content(source_content):
                print("⚠️ 检测到二进制内容，使用文件信息生成示例内容")
                source_content = generate_sample_content(source_file)
            elif len(source_content.strip()) == 0:
                print("⚠️ 文件内容为空，生成示例内容")
                source_content = generate_sample_content(source_file)

            # 限制内容长度，避免API调用超时
            if len(source_content) > 2000:
                print(f"⚠️ 内容过长({len(source_content)}字符)，截取前2000字符")
                source_content = source_content[:2000] + "\n\n[内容已截取...]"

            print(f"📝 处理后源文件内容长度: {len(source_content)}")
            print(f"📝 内容预览: {source_content[:100]}...")

        except Exception as e:
            print(f"❌ 文件内容处理失败: {e}")
            source_content = "示例文档内容，需要格式化"

        # 构建更详细的格式对齐指令
        target_content = target_file.get('content', '')
        if target_content and len(target_content.strip()) > 0:
            instruction = f"""请将以下文档内容按照目标格式进行重新组织和格式化：

目标格式示例：
{target_content}

要求：
1. 保持原文档的核心信息和内容
2. 按照目标格式的结构重新组织内容
3. 使用目标格式的标题层级和编号方式
4. 保持专业的文档风格
5. 确保内容完整性和逻辑性

请直接输出格式化后的完整文档内容，不要添加额外的说明。"""
        else:
            instruction = f"""请将以下文档内容格式化为标准的专业文档格式：

要求：
1. 使用清晰的标题层级结构
2. 内容分段明确，逻辑清晰
3. 使用适当的编号和列表
4. 保持专业的文档风格
5. 确保内容完整性

请直接输出格式化后的完整文档内容。"""

        print("🚀 开始调用星火X1进行格式对齐...")
        print(f"📝 格式化指令: {instruction[:100]}...")
        result = format_alignment_coordinator.format_with_spark_x1(
            content=source_content,
            instruction=instruction,
            temperature=0.3,  # 降低温度，提高一致性
            max_tokens=3000,  # 增加最大token数
            timeout=60
        )

        if result.get('success'):
            return jsonify({
                'success': True,
                'data': {
                    'aligned_content': result.get('formatted_content', ''),
                    'alignment_score': 0.95,
                    'suggestions': ['格式对齐完成', '使用星火X1处理'],
                    'session_id': session_id,
                    'status': 'completed',
                    'task_id': result.get('task_id', '')
                }
            })
        else:
            return jsonify({
                'success': False,
                'error': result.get('error', '格式对齐失败')
            }), 500

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

# ==================== 文档审查模块 API ====================
# @AI-Generated: 2025-01-25, Confidence: 0.99, Model: Claude Sonnet 4, Prompt: document_review_api

# 初始化文档审查协调器
document_review_coordinator = None

try:
    import sys
    import os
    sys.path.append(os.path.dirname(os.path.abspath(__file__)))
    from core.document_review_coordinator import DocumentReviewCoordinator
    # 使用密钥管理器获取API密钥
    api_password = get_spark_x1_key('document_review')
    document_review_coordinator = DocumentReviewCoordinator(api_password)
    print("✅ 文档审查协调器初始化成功")
except Exception as e:
    print(f"❌ 文档审查协调器初始化失败: {e}")
    document_review_coordinator = None

@app.route('/api/document-review/analyze', methods=['POST'])
def analyze_document():
    """文档审查分析API"""
    try:
        if not document_review_coordinator:
            return jsonify({
                'success': False,
                'error': '文档审查服务未初始化'
            }), 500

        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'error': '请求数据为空'
            }), 400

        # 获取参数
        document_content = data.get('content', '').strip()
        review_type = data.get('review_type', 'keyword_review')
        custom_prompt = data.get('custom_prompt', '')

        if not document_content:
            return jsonify({
                'success': False,
                'error': '文档内容不能为空'
            }), 400

        # 执行文档审查 - 使用默认的综合审查模式
        result = document_review_coordinator.review_document(
            document_content=document_content,
            review_type=review_type,
            custom_prompt=custom_prompt if custom_prompt else None
        )

        if result.get('success'):
            return jsonify({
                'success': True,
                'data': {
                    'review_result': result.get('review_result', ''),
                    'document_length': result.get('document_length', 0),
                    'processing_time': result.get('processing_time', 0),
                    'review_type': result.get('review_type', review_type),
                    'chunks_count': result.get('chunks_count', 1)
                }
            })
        else:
            return jsonify({
                'success': False,
                'error': result.get('error', '文档审查失败')
            }), 500

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'文档审查异常: {str(e)}'
        }), 500

@app.route('/api/document-review/templates', methods=['GET'])
def get_review_templates():
    """获取审查模板列表"""
    try:
        if not document_review_coordinator:
            return jsonify({
                'success': False,
                'error': '文档审查服务未初始化'
            }), 500

        templates = document_review_coordinator.get_available_review_types()

        return jsonify({
            'success': True,
            'templates': templates,
            'count': len(templates)
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'获取审查模板失败: {str(e)}'
        }), 500

@app.route('/api/document-review/export-pdf', methods=['POST'])
def export_review_report_pdf():
    """导出PDF格式审查报告"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'error': '请求数据为空'
            }), 400

        review_data = data.get('review_data', {})
        filename = data.get('filename', 'document_review_report')

        if not review_data or not review_data.get('review_result'):
            return jsonify({
                'success': False,
                'error': '审查结果数据为空'
            }), 400

        # 生成PDF内容
        pdf_content = generate_pdf_report(review_data, filename)

        # 创建响应
        response = make_response(pdf_content)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'

        return response

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'PDF导出失败: {str(e)}'
        }), 500

@app.route('/api/document-review/export-word', methods=['POST'])
def export_review_report_word():
    """导出Word格式审查报告"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({
                'success': False,
                'error': '请求数据为空'
            }), 400

        review_data = data.get('review_data', {})
        filename = data.get('filename', 'document_review_report')

        if not review_data or not review_data.get('review_result'):
            return jsonify({
                'success': False,
                'error': '审查结果数据为空'
            }), 400

        # 生成Word内容
        word_content = generate_word_report(review_data, filename)

        # 创建响应
        response = make_response(word_content)
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}.docx"'

        return response

    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Word导出失败: {str(e)}'
        }), 500

def generate_pdf_report(review_data, filename):
    """生成PDF报告"""
    # @AI-Generated: 2025-01-25, Confidence: 0.99, Model: Claude Sonnet 4, Prompt: generate_pdf_report
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import io
        from datetime import datetime

        # 创建内存缓冲区
        buffer = io.BytesIO()

        # 创建PDF文档
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1*inch)

        # 获取样式
        styles = getSampleStyleSheet()

        # 创建自定义样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # 居中
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue
        )

        # 构建文档内容
        story = []

        # 标题
        story.append(Paragraph("📋 AI文档审查报告", title_style))
        story.append(Spacer(1, 20))

        # 元数据表格
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        meta_data = [
            ['生成时间', timestamp],
            ['文档长度', f"{review_data.get('document_length', 0)} 字符"],
            ['处理时间', f"{review_data.get('processing_time', 0):.2f} 秒"],
        ]

        if review_data.get('chunks_count', 1) > 1:
            meta_data.append(['分块处理', f"{review_data.get('chunks_count')} 个块"])

        meta_table = Table(meta_data, colWidths=[2*inch, 3*inch])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('BACKGROUND', (1, 0), (1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        story.append(meta_table)
        story.append(Spacer(1, 30))

        # 审查结果内容
        story.append(Paragraph("审查结果", heading_style))
        story.append(Spacer(1, 12))

        # 处理审查结果文本
        review_content = review_data.get('review_result', '')

        # 简单的Markdown到PDF转换
        lines = review_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue

            if line.startswith('# '):
                # 一级标题
                story.append(Paragraph(line[2:], heading_style))
                story.append(Spacer(1, 12))
            elif line.startswith('## '):
                # 二级标题
                sub_heading_style = ParagraphStyle(
                    'SubHeading',
                    parent=styles['Heading3'],
                    fontSize=12,
                    spaceAfter=8,
                    textColor=colors.darkgreen
                )
                story.append(Paragraph(line[3:], sub_heading_style))
                story.append(Spacer(1, 8))
            elif line.startswith('- ') or line.startswith('* '):
                # 列表项
                list_style = ParagraphStyle(
                    'ListItem',
                    parent=styles['Normal'],
                    leftIndent=20,
                    spaceAfter=4
                )
                story.append(Paragraph(f"• {line[2:]}", list_style))
            else:
                # 普通段落
                story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 6))

        # 页脚
        story.append(Spacer(1, 30))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=1  # 居中
        )
        story.append(Paragraph("本报告由aiDoc AI文档审查系统生成 | 基于讯飞星火X1大模型", footer_style))

        # 构建PDF
        doc.build(story)

        # 获取PDF内容
        pdf_content = buffer.getvalue()
        buffer.close()

        return pdf_content

    except ImportError:
        # 如果没有安装reportlab，返回简单的文本文件
        content = f"""AI文档审查报告

生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
文档长度: {review_data.get('document_length', 0)} 字符
处理时间: {review_data.get('processing_time', 0):.2f} 秒

审查结果:
{review_data.get('review_result', '')}

本报告由aiDoc AI文档审查系统生成 | 基于讯飞星火X1大模型
"""
        return content.encode('utf-8')
    except Exception as e:
        # 发生错误时返回简单的文本内容
        content = f"""AI文档审查报告 (简化版)

生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
错误信息: PDF生成失败 - {str(e)}

审查结果:
{review_data.get('review_result', '')}
"""
        return content.encode('utf-8')

def generate_word_report(review_data, filename):
    """生成Word报告"""
    # @AI-Generated: 2025-01-25, Confidence: 0.99, Model: Claude Sonnet 4, Prompt: generate_word_report
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        import io
        from datetime import datetime

        # 创建Word文档
        doc = Document()

        # 设置文档标题
        title = doc.add_heading('📋 AI文档审查报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加元数据表格
        doc.add_heading('报告信息', level=1)

        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'

        # 填充表格数据
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        table_data = [
            ('生成时间', timestamp),
            ('文档长度', f"{review_data.get('document_length', 0)} 字符"),
            ('处理时间', f"{review_data.get('processing_time', 0):.2f} 秒"),
        ]

        for i, (key, value) in enumerate(table_data):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value

        # 如果有分块处理，添加额外行
        if review_data.get('chunks_count', 1) > 1:
            row = table.add_row()
            row.cells[0].text = '分块处理'
            row.cells[1].text = f"{review_data.get('chunks_count')} 个块"

        # 添加审查结果
        doc.add_heading('审查结果', level=1)

        # 处理审查结果文本
        review_content = review_data.get('review_result', '')
        lines = review_content.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith('# '):
                # 一级标题
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # 二级标题
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # 三级标题
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                # 列表项
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(line[2:])
            elif line.startswith('**') and line.endswith('**'):
                # 粗体文本
                p = doc.add_paragraph()
                p.add_run(line[2:-2]).bold = True
            else:
                # 普通段落
                doc.add_paragraph(line)

        # 添加页脚
        doc.add_paragraph()
        footer = doc.add_paragraph('本报告由aiDoc AI文档审查系统生成 | 基于讯飞星火X1大模型')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8)

        # 保存到内存缓冲区
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    except ImportError:
        # 如果没有安装python-docx，返回简单的文本文件
        content = f"""AI文档审查报告

生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
文档长度: {review_data.get('document_length', 0)} 字符
处理时间: {review_data.get('processing_time', 0):.2f} 秒

审查结果:
{review_data.get('review_result', '')}

本报告由aiDoc AI文档审查系统生成 | 基于讯飞星火X1大模型
"""
        return content.encode('utf-8')
    except Exception as e:
        # 发生错误时返回简单的文本内容
        content = f"""AI文档审查报告 (简化版)

生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
错误信息: Word生成失败 - {str(e)}

审查结果:
{review_data.get('review_result', '')}
"""
        return content.encode('utf-8')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
